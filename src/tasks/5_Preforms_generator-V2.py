

import os, sys, json, shutil, tempfile, datetime, re, io, time, traceback
import urllib.parse, html, random
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from Config import Global

# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIGURACIÓN GLOBAL
# ═══════════════════════════════════════════════════════════════════════════════

MYSQL_CONFIG = {
    "host":     Global.DB_HOST,
    "user":     Global.DB_USER,
    "password": Global.DB_PASSWORD,
    "database": Global.DATABASE,
    "connection_timeout": 20,
}


GEMINI_CREDENTIALS_PATH = Global.GEMINI_CREDENTIALS
BUCKET_NAME             = Global.BUCKET_NAME
BUCKET_FOLDER           = "Documentos de Contratación"
AI_MODEL                = "gemini-2.5-pro"

SCRIPT_DIR    = Path(__file__).parent
TEMPLATE_DOCX = SCRIPT_DIR / "FICHA TECNICA MICROFONO Y MEMORIA.docx"
TEMPLATE_XLSX = SCRIPT_DIR / "FORMATO DE PROFORMA RECREADO VACÍO.xlsm"



USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)
DEFAULT_HEADERS = {
    "User-Agent":      USER_AGENT,
    "Accept-Language": "es-ES,es;q=0.9,en;q=0.7",
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
}

PROVEEDORES_NACIONALES = [
    "https://mibodega.ec/",
    "https://bodeguitadelahorro.com/",
    "https://comercialvaca.ec/",
    "https://kissu.com.ec/",
    "https://altatenalmacen.com.ec/",
    "https://www.tventas.com/",
    "https://store.intcomex.com/es-XPE/Home",
    "https://lavictoria.ec/",
    "https://www.marcimex.com/",
    "https://almacenesespana.ec/",
    "https://www.novicompu.com/",
    "https://technoprimec.com/",
    "https://point.com.ec/",
    "https://www.computron.com.ec/",
    "https://nomadaware.com.ec/",
    "https://www.idcmayoristas.com/",
    "https://tecnit.com.ec/",
    "https://tecnomegastore.ec/",
    "https://www.artefacta.com/",
    "https://mundotek.com.ec/",
    "https://eljuri.store/",
    "https://www.kywi.com.ec/",
    "https://tecnocostoec.com/",
    "https://www.almacenesjapon.com/",
    "https://electromegaecuador.com/",
    "https://www.compraecuador.com/",
    "https://granhogar.com.ec/",
    "https://miamihome-ec.com/",
    # Para carpas y eventos
    "https://www.carpas-casa.com.ec/",
    "https://www.eventoshop.com.ec/",
    "https://www.toldosya.com.ec/",
    "https://www.feriaventos.com.ec/",
    "https://www.carpasycarros.com.ec/",
    
    # Para equipos industriales
    "https://www.industriasys.com.ec/",
    "https://www.equipindust.com.ec/",
    "https://www.maquinariasecuador.com/",
]

PROVEEDORES_EXTRANJEROS_AMERICA = [
    "https://www.amazon.com/",        # EE.UU.
    "https://www.mercadolibre.com.mx/",
    "https://www.mercadolibre.com.co/",
    "https://www.mercadolibre.com.ar/",
    "https://www.mercadolibre.cl/",
    "https://www.mercadolibre.com.pe/",
    "https://www.ebay.com/",          # EE.UU.
    "https://www.walmart.com/",
    "https://www.bestbuy.com/",
    "https://www.homedepot.com/",
    "https://www.costco.com/",
]

# ═══════════════════════════════════════════════════════════════════════════════
#  MOTOR DE BÚSQUEDA MEJORADO CON SCORING DE RELEVANCIA
# ═══════════════════════════════════════════════════════════════════════════════

class ProductSearchEngine:
    """Motor de búsqueda mejorado para productos con scoring de relevancia"""
    
    def __init__(self, region="ec-es", timeout=25):
        self.region = region
        self.timeout = timeout
        
        # Palabras clave que indican páginas no relevantes
        self.invalid_url_patterns = [
            'categoria', 'category', 'categoría', 'blog', 'foro', 'forum',
            'noticia', 'news', 'opinion', 'review', 'comparativa', 'pregunta',
            'respuesta', 'comentario', 'qué es', 'cómo funciona', 'derecho', 'abogado'
        ]
        
        # Patrones de URL de producto
        self.product_url_patterns = [
            r'/p/', r'/product/', r'/item/', r'/dp/', r'/producto/',
            r'/itm/', r'/detail/', r'/oferta/', r'id=\d+', r'cod=\d+', 
            r'sku=\w+', r'product=\d+'
        ]
    
    def buscar_con_scoring(self, nombre_articulo: str, marca: str = "", modelo: str = "",
                            num_resultados: int = 15,
                            dominios_permitidos: List[str] = None) -> List[Dict]:
        query = self._construir_query(nombre_articulo, marca, modelo)
        log(f"  🔍 Buscando con scoring: {query}")
        resultados_raw = duckduckgo_search_simple(query, num_results=num_resultados * 2, region=self.region)

        resultados = []
        for result in resultados_raw:
            if dominios_permitidos:
                dom = _dominio(result.get('href', ''))
                if dom not in dominios_permitidos:
                    continue

            producto = self._analizar_resultado(result, nombre_articulo, marca, modelo)
            if producto['relevancia'] >= 50:
                html_text = fetch_html(producto['url'], timeout=self.timeout)
                if html_text:
                    if not producto.get('precio'):
                        producto['precio'] = extraer_precio_de_html(html_text)
                    if not producto.get('imagen_url'):
                        producto['imagen_url'] = extraer_imagen_de_html(html_text, producto['url'])
                    if producto['imagen_url']:
                        producto['imagen_validada'] = validar_url_imagen(producto['imagen_url'])
                resultados.append(producto)
            time.sleep(0.3)

        resultados.sort(key=lambda x: (-x['relevancia'], x.get('precio', 999999) or 999999))
        return resultados[:5]

    def _construir_query(self, nombre: str, marca: str, modelo: str) -> str:
        """Construye query optimizada para búsqueda"""
        partes = []
        
        if marca and marca.lower() not in nombre.lower():
            partes.append(marca)
        if modelo and modelo.lower() not in nombre.lower():
            partes.append(modelo)
        
        partes.append(nombre)
        partes.append("comprar")
        partes.append("precio")
        
        return " ".join(partes[:6])
    
    def _extraer_keywords_importantes(self, texto: str) -> List[str]:
        """Extrae palabras clave importantes del texto"""
        stop_words = {'el', 'la', 'los', 'las', 'un', 'una', 'unos', 'unas',
                     'de', 'del', 'y', 'o', 'para', 'por', 'con', 'sin', 'en',
                     'a', 'ante', 'bajo', 'cabe', 'contra', 'de', 'desde',
                     'durante', 'entre', 'hacia', 'hasta', 'mediante', 'segun'}
        
        palabras = re.findall(r'\b\w+\b|\d+[x]\d+|\d+\.\d+|\d+', texto.lower())
        keywords = [p for p in palabras if p not in stop_words and (len(p) > 2 or p[0].isdigit())]
        
        return keywords
    
    def _analizar_resultado(self, result: Dict, nombre_articulo: str, marca: str, modelo: str) -> Dict:
        """Analiza y calcula relevancia de un resultado"""
        titulo = result.get('title', '').lower()
        cuerpo = result.get('body', '').lower()
        url = result.get('url', '')
        
        # Extraer palabras clave
        keywords = self._extraer_keywords_importantes(nombre_articulo)
        
        score = 0
        
        # 1. Coincidencia en título (40 puntos máx)
        titulo_score = sum(1 for kw in keywords if kw in titulo)
        score += (titulo_score / max(len(keywords), 1)) * 40
        
        # 2. Coincidencia en descripción (20 puntos máx)
        cuerpo_score = sum(1 for kw in keywords if kw in cuerpo)
        score += (cuerpo_score / max(len(keywords), 1)) * 20
        
        # 3. Números específicos como "6x3" (20 puntos)
        numeros_producto = re.findall(r'\d+[x]\d+|\d+\.\d+|\d+', nombre_articulo)
        if numeros_producto:
            for numero in numeros_producto:
                if numero in titulo or numero in cuerpo:
                    score += 20
                    break
        
        # 4. Marca (10 puntos)
        if marca and marca.lower() in titulo:
            score += 10
        
        # 5. Modelo (10 puntos)
        if modelo and modelo.lower() in titulo:
            score += 10
        
        # 6. Validar que sea URL de producto
        if self._es_url_producto(url):
            score += 15
        else:
            score -= 25
        
        # 7. Penalizar URLs inválidas
        if self._es_url_invalida(url):
            score -= 35
        
        return {
            'url': url,
            'titulo': result.get('title', ''),
            'descripcion': result.get('body', ''),
            'relevancia': max(0, min(100, score)),
            'precio': None,
            'imagen_url': None,
            'imagen_validada': False
        }
    
    def _es_url_producto(self, url: str) -> bool:
        """Verifica si la URL apunta a página de producto específico"""
        url_lower = url.lower()
        for pattern in self.product_url_patterns:
            if re.search(pattern, url_lower, re.IGNORECASE):
                return True
        return False
    
    def _es_url_invalida(self, url: str) -> bool:
        """Verifica si la URL es de categoría, blog, etc."""
        url_lower = url.lower()
        for pattern in self.invalid_url_patterns:
            if pattern in url_lower:
                return True
        return False


def _dominio(url):
    return urllib.parse.urlparse(url).netloc.lower().lstrip("www.")

DOMINIOS_NACIONALES  = [_dominio(u) for u in PROVEEDORES_NACIONALES]
DOMINIOS_EXTRANJEROS_AMERICA = [_dominio(u) for u in PROVEEDORES_EXTRANJEROS_AMERICA]
DOMINIOS_TODOS       = DOMINIOS_NACIONALES + DOMINIOS_EXTRANJEROS_AMERICA


# ═══════════════════════════════════════════════════════════════════════════════
#  UTILIDADES BÁSICAS
# ═══════════════════════════════════════════════════════════════════════════════

def _get_session():
    import requests
    session = requests.Session()
    retries = Retry(total=3, backoff_factor=1.5,
                    status_forcelist=[429, 500, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retries)
    session.mount("https://", adapter)
    session.mount("http://",  adapter)
    session.headers.update(DEFAULT_HEADERS)
    return session

def log(msg, nivel="INFO"):
    iconos = {"INFO": "ℹ️ ", "OK": "✅", "WARN": "⚠️ ", "ERR": "❌"}
    print(f"{iconos.get(nivel,'  ')} [{datetime.datetime.now().strftime('%H:%M:%S')}] {msg}",
        flush=True)

def paso(n, total, desc):
    print(f"\n{'─'*60}\n  PASO {n}/{total}: {desc}\n{'─'*60}", flush=True)

def resolver_credenciales_a_archivo():
    raw = GEMINI_CREDENTIALS_PATH
    if raw is None:
        raise ValueError("GEMINI_CREDENTIALS_PATH no definida.")
    stripped = raw.strip()
    if stripped.startswith("{"):
        creds_dict = json.loads(stripped)
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".json",
                                          delete=False, encoding="utf-8")
        json.dump(creds_dict, tmp)
        tmp.close()
        log("Credenciales resueltas desde variable de entorno.")
        return tmp.name
    log(f"Credenciales resueltas desde archivo: {stripped}")
    return stripped

def _parsear_json_de_respuesta(raw):
    if not raw:
        return None
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        for ch_open, ch_close in [("{", "}"), ("[", "]")]:
            i = raw.find(ch_open)
            if i < 0:
                continue
            depth = 0
            for j in range(i, len(raw)):
                if raw[j] == ch_open:   depth += 1
                elif raw[j] == ch_close:
                    depth -= 1
                    if depth == 0:
                        try:    return json.loads(raw[i:j+1])
                        except: break
    return None

# ═══════════════════════════════════════════════════════════════════════════════
#  BÚSQUEDA CON DUCKDUCKGO (versión mejorada)
# ═══════════════════════════════════════════════════════════════════════════════

from ddgs import DDGS

def duckduckgo_search_simple(query, num_results=10, region="ec-es", time_filter=None):
    """
    Búsqueda mejorada con filtrado de resultados irrelevantes
    """
    try:
        resultados = []
        with DDGS() as ddgs:
            for r in ddgs.text(
                query, 
                region=region, 
                max_results=num_results,
                safesearch='off',
                timelimit=time_filter
            ):
                url = r.get("href", "")
                title = r.get("title", "").lower()
                body = r.get("body", "").lower()
                
                # Palabras que indican contenido NO relevante para productos
                palabras_irrelevantes = [
                    'forum', 'pregunta', 'respuesta', 'comentario', 'qué es',
                    'cómo funciona', 'derecho', 'abogado', 'wikipedia', 
                    'youtube', 'facebook', 'twitter', 'instagram'
                ]
                
                # Palabras que indican página de TIENDA (buenas)
                palabras_tienda = [
                    'tienda', 'shop', 'store', 'comprar', 'price', 'precio',
                    'catalogo', 'product', 'oferta', 'descuento'
                ]
                
                es_irrelevante = False
                for palabra in palabras_irrelevantes:
                    if palabra in title or (palabra in body and len(body) < 300):
                        es_irrelevante = True
                        break
                
                # Si no es irrelevante O parece tienda, incluir
                es_tienda = any(palabra in title or palabra in body for palabra in palabras_tienda)
                
                if not es_irrelevante or es_tienda:
                    resultados.append({
                        "url": url,
                        "title": r.get("title", ""),
                        "body": r.get("body", "")
                    })
                
                if len(resultados) >= num_results:
                    break
                    
        if resultados:
            log(f"    DuckDuckGo: {len(resultados)} resultado(s) relevante(s)")
        return resultados
        
    except Exception as e:
        log(f"    Error en búsqueda DuckDuckGo: {e}", "WARN")
        return []

def buscar_urls_con_sites(query, dominios, num_por_sitio=2):
    """
    Busca en DuckDuckGo restringiendo a dominios específicos con mejor filtrado
    """
    todas_urls = []
    
    # Mejorar la query para productos
    query_mejorada = query
    # Agregar palabras clave de producto si no están
    if not any(p in query.lower() for p in ['comprar', 'precio', 'venta']):
        query_mejorada = f"{query} comprar"
    
    for dominio in dominios:
        query_con_site = f"{query_mejorada} site:{dominio}"
        resultados = duckduckgo_search_simple(query_con_site, num_results=num_por_sitio, region="ec-es")
        
        for r in resultados:
            # Verificar que la URL realmente pertenezca al dominio buscado
            if dominio in r["url"].lower():
                todas_urls.append({
                    "url": r["url"],
                    "titulo": r["title"],
                    "snippet": r["body"],
                    "dominio": dominio
                })
        time.sleep(0.5)  # Pausa para no saturar
    
    # Eliminar duplicados
    vistas = set()
    urls_unicas = []
    for item in todas_urls:
        if item["url"] not in vistas:
            vistas.add(item["url"])
            urls_unicas.append(item)
    
    return urls_unicas


def buscar_producto_google(nombre, marca, modelo, max_por_dominio=2):
    """
    Busca el producto usando DuckDuckGo con mejor filtrado contextual
    """
    candidatos = []
    
    # Construir query más específica
    partes_query = []
    if nombre:
        partes_query.append(nombre)
    if marca and marca.lower() not in nombre.lower():
        partes_query.append(marca)
    if modelo and modelo.lower() not in nombre.lower():
        partes_query.append(modelo)
    
    query_base = " ".join(partes_query).strip()
    
    # Agregar palabras clave de producto
    query_base = f"{query_base} comprar precio"
    
    if not query_base:
        log("  Query de búsqueda vacía.", "WARN")
        return []
    
    log(f"  Buscando en DuckDuckGo: '{query_base}'")
    
    # Primero: búsqueda en dominios específicos
    candidatos_raw = buscar_urls_con_sites(query_base, DOMINIOS_TODOS, max_por_dominio)
    
    for item in candidatos_raw:
        dom = item["dominio"]
        candidatos.append({
            "url":           item["url"],
            "titulo":        item.get("titulo", ""),
            "snippet":       item.get("snippet", ""),
            "es_extranjero": dom in DOMINIOS_EXTRANJEROS_AMERICA,
            "dominio":       dom,
        })
    
    # Segundo: búsqueda general en español
    if not candidatos:
        log("  Sin resultados en proveedores. Buscando en toda la web…", "WARN")
        resultados = duckduckgo_search_simple(query_base, num_results=15, region="ec-es")
        
        for r in resultados:
            url = r["url"]
            dom = _dominio(url)
            # Verificar que sea una tienda/tienda en línea (no foros)
            es_tienda = any(tienda in url.lower() for tienda in [
                'tienda', 'shop', 'store', 'producto', 'comprar', 'product', 'buy',
                'mercado', 'oferta', 'precio', 'catalogo'
            ])
            
            if es_tienda or dom in DOMINIOS_TODOS:
                candidatos.append({
                    "url":           url,
                    "titulo":        r.get("title", ""),
                    "snippet":       r.get("body", ""),
                    "es_extranjero": dom not in DOMINIOS_NACIONALES,
                    "dominio":       dom,
                })
    
    # Priorizar URLs que contengan el producto en el título
    palabras_clave = [p.lower() for p in partes_query if len(p) > 3]
    for c in candidatos:
        titulo_lower = c["titulo"].lower()
        # Puntuar relevancia
        c["relevancia"] = sum(1 for p in palabras_clave if p in titulo_lower)
    
    # Ordenar por relevancia y luego por nacionales primero
    candidatos.sort(key=lambda x: (-x["relevancia"], x["es_extranjero"]))
    
    # Deduplicar y limitar
    vistos = set()
    out = []
    for c in candidatos:
        if c["url"] not in vistos:
            vistos.add(c["url"])
            out.append(c)
            if len(out) >= 15:
                break
    
    log(f"  URLs candidatas encontradas: {len(out)}", "OK" if out else "WARN")
    
    if not out:
        # Sugerencia para búsqueda manual
        log(f"  💡 Sugerencia: Busca manualmente '{query_base}' en Google", "INFO")
        log(f"  Puede que el producto no esté disponible en los proveedores configurados.", "WARN")
        
        # Último recurso: búsqueda sin filtros
        log("  Intentando búsqueda amplia…", "INFO")
        resultados = duckduckgo_search_simple(partes_query[0], num_results=10, region="ec-es")
        for r in resultados[:5]:
            log(f"    Resultado alternativo: {r['title'][:60]}", "INFO")
    
    return out


def buscar_imagen_google(nombre, marca, modelo):
    """
    Busca imagen del producto usando DuckDuckGo Images
    """
    query = " ".join(x for x in [nombre, marca, modelo] if x).strip()
    
    try:
        with DDGS() as ddgs:
            # Búsqueda de imágenes con región en español
            resultados = list(ddgs.images(
                query, 
                region="ec-es", 
                max_results=5,
                safesearch='off'
            ))
            
            for r in resultados:
                img_url = r.get("image", "")
                # Verificar que la imagen sea de un sitio confiable
                if img_url and img_url.startswith("http"):
                    # Evitar imágenes de foros o sitios no comerciales
                    if any(sitio in img_url.lower() for sitio in ['amazon', 'mercadolibre', 'tienda', 'shop']):
                        if validar_url_imagen(img_url):
                            log(f"  Imagen encontrada: {img_url[:70]}", "OK")
                            return img_url
            
            # Segunda pasada: cualquier imagen válida
            for r in resultados:
                img_url = r.get("image", "")
                if img_url and img_url.startswith("http"):
                    if validar_url_imagen(img_url):
                        log(f"  Imagen encontrada (alternativa): {img_url[:70]}", "OK")
                        return img_url
                        
    except Exception as e:
        log(f"  Error buscando imagen: {e}", "WARN")
    
    return None

# ═══════════════════════════════════════════════════════════════════════════════
#  VALIDACIÓN DE URLs E IMÁGENES
# ═══════════════════════════════════════════════════════════════════════════════

_URL_CACHE = {}

def validar_url(url, timeout=10):
    """Verifica que una URL responda HTTP 200-399."""
    import requests
    if not url or not isinstance(url, str) or not url.startswith("http"):
        return False
    if url in _URL_CACHE:
        return _URL_CACHE[url]
    try:
        r = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout,
                         allow_redirects=True, stream=True)
        ok = r.status_code < 400
        r.close()
        _URL_CACHE[url] = ok
        return ok
    except Exception:
        _URL_CACHE[url] = False
        return False


def validar_url_imagen(url, timeout=10):
    """Verifica que una URL devuelve una imagen real (Content-Type + tamaño mínimo)."""
    import requests
    if not url or not isinstance(url, str) or not url.startswith("http"):
        return False
    if "localhost" in url:
        return False
    try:
        r = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout,
                         allow_redirects=True, stream=True)
        if r.status_code >= 400:
            return False
        ct = (r.headers.get("Content-Type") or "").lower()
        if not ct.startswith("image/"):
            return False
        chunk = next(r.iter_content(chunk_size=2048), b"")
        r.close()
        return len(chunk) >= 512
    except Exception:
        return False


# ═══════════════════════════════════════════════════════════════════════════════
#  EXTRACCIÓN DE DATOS DE PÁGINAS
# ═══════════════════════════════════════════════════════════════════════════════

def fetch_html(url, timeout=15, max_bytes=400_000):
    """Descarga HTML parcial de una URL."""
    import requests
    try:
        r = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout,
                         allow_redirects=True, stream=True)
        if r.status_code >= 400:
            return None
        chunks, leido = [], 0
        for chunk in r.iter_content(chunk_size=8192, decode_unicode=False):
            if not chunk: break
            chunks.append(chunk)
            leido += len(chunk)
            if leido >= max_bytes: break
        r.close()
        return b"".join(chunks).decode("utf-8", errors="ignore")
    except Exception:
        return None


def extraer_imagen_de_html(html_text, url_base):
    """Extrae la mejor URL de imagen del HTML (og:image, JSON-LD, twitter:image)."""
    if not html_text:
        return None

    patrones_meta = [
        r'<meta[^>]+property=["\']og:image:secure_url["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+property=["\']og:image:secure_url["\']',
        r'<meta[^>]+property=["\']og:image["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+property=["\']og:image["\']',
        r'<meta[^>]+name=["\']twitter:image["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+name=["\']twitter:image["\']',
    ]
    for pat in patrones_meta:
        m = re.search(pat, html_text, re.IGNORECASE)
        if m:
            url = urllib.parse.urljoin(url_base, html.unescape(m.group(1)))
            if url.startswith("http") and "localhost" not in url:
                return url

    # JSON-LD
    for m in re.finditer(
        r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>',
        html_text, re.IGNORECASE | re.DOTALL
    ):
        try:
            data = json.loads(m.group(1).strip())
            for cand in (data if isinstance(data, list) else [data]):
                if not isinstance(cand, dict): continue
                img = cand.get("image")
                if isinstance(img, str) and img.startswith("http"): return img
                if isinstance(img, list) and img:
                    p = img[0]
                    if isinstance(p, str)  and p.startswith("http"): return p
                    if isinstance(p, dict) and p.get("url","").startswith("http"): return p["url"]
        except Exception:
            continue
    return None


def extraer_precio_de_html(html_text):
    """Extrae precio aproximado en USD del HTML. Devuelve float o None."""
    if not html_text:
        return None
    candidatos = []

    # JSON-LD offers.price
    for m in re.finditer(
        r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>',
        html_text, re.IGNORECASE | re.DOTALL
    ):
        try:
            data = json.loads(m.group(1).strip())
            stack = data if isinstance(data, list) else [data]
            while stack:
                cur = stack.pop()
                if isinstance(cur, dict):
                    offer = cur.get("offers")
                    if isinstance(offer, dict):  stack.append(offer)
                    elif isinstance(offer, list): stack.extend(offer)
                    p = cur.get("price")
                    if p:
                        try: candidatos.append(float(str(p).replace(",", "")))
                        except: pass
                elif isinstance(cur, list):
                    stack.extend(cur)
        except: continue

    # Patrones adicionales: $ 1,234.56 o $1.234,56 o 1 234,56 $
    patrones_extra = [
        r'(?:USD|US\$|\$)\s*([0-9]{1,3}(?:[.,][0-9]{3})*(?:[.,][0-9]{1,2}))',  # $1,234.56
        r'([0-9]{1,3}(?:[.,][0-9]{3})*(?:[.,][0-9]{1,2}))\s*(?:USD|US\$|\$)',  # 1,234.56 $
        r'(?:Precio|PRECIO|precio)\s*:?\s*\$?\s*([0-9]+(?:[.,][0-9]+))',        # Precio: $1234.56
    ]
    for pat in patrones_extra:
        for m in re.finditer(pat, html_text, re.IGNORECASE):
            try:
                val = m.group(1).replace(",", ".").replace(" ", "")
                # manejar formato con punto de miles
                if re.match(r'^\d{1,3}(\.\d{3})*,\d{1,2}$', m.group(1)):  # formato 1.234,56
                    val = m.group(1).replace(".", "").replace(",", ".")
                candidatos.append(float(val))
            except:
                pass

    candidatos = [c for c in candidatos if 1.0 <= c <= 50000.0]
    if not candidatos:
        return None
    candidatos.sort()
    return candidatos[len(candidatos)//2]

def extraer_precio_con_ia(html_text, url_producto, modelo_ai):
    if not html_text or not modelo_ai:
        return None
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_text, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        texto = soup.get_text(separator=' ', strip=True)[:3000]
    except:
        texto = html_text[:3000]

    prompt = f"""Eres un asistente que extrae el precio de un producto desde el texto de una página web.
    URL: {url_producto}

    Texto:
    {texto}

    Devuelve ÚNICAMENTE un JSON con la estructura:
    {{"precio": número, "moneda": "USD"}}

    Reglas:
    - El precio debe ser el valor final que paga el comprador (impuestos incluidos si aparecen).
    - Si ves varios precios, quédate con el que más se repite o el que aparezca junto a "precio", "total", "a pagar".
    - Si no hay precio, devuelve {{"precio": null, "moneda": null}}.
    """
    response = modelo_ai.generate_content(prompt)
    data = _parsear_json_de_respuesta(response.text)
    if data and isinstance(data.get("precio"), (int, float)):
        return float(data["precio"])
    return None

def extraer_info_proveedor(url_producto, nombre_producto, modelo_ai):
    resultado_seguro = {
        "nombre": nombre_producto,
        "caracteristicas": [],
        "especificaciones": [],
        "descripcion": ""
    }

    html_text = fetch_html(url_producto, timeout=15, max_bytes=300_000)
    if not html_text:
        return resultado_seguro

    from bs4 import BeautifulSoup
    try:
        soup = BeautifulSoup(html_text, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        texto = soup.get_text(separator=' ', strip=True)[:4000]
    except:
        texto = html_text[:4000]

    prompt = f"""
Eres un analizador de páginas de producto. Extrae la información técnica del siguiente texto
proveniente de la página del producto "{nombre_producto}":

{texto}

Devuelve ÚNICAMENTE un JSON válido:
{{
  "nombre": "nombre exacto del producto en la tienda",
  "caracteristicas": ["característica 1", "característica 2", ...],
  "especificaciones": ["especificación 1", "especificación 2", ...],
  "descripcion": "descripción comercial de 50 palabras"
}}
Si no encuentras información, devuelve arrays vacíos.
"""
    try:
        response = modelo_ai.generate_content(prompt)
        data = _parsear_json_de_respuesta(response.text)
        if data:
            for k, v in data.items():
                if v:  # solo reemplaza si el valor no está vacío
                    resultado_seguro[k] = v
    except Exception as e:
        log(f"  Error al extraer info del proveedor: {e}", "WARN")

    return resultado_seguro

# ═══════════════════════════════════════════════════════════════════════════════
#  HELPERS DE ARTÍCULO
# ═══════════════════════════════════════════════════════════════════════════════

def _hay_marca_modelo_especificos(info_articulo):
    """Determina si el artículo tiene marca Y modelo concretos."""
    marca  = (info_articulo.get("marca")  or "").strip().lower()
    modelo = (info_articulo.get("modelo") or "").strip().lower()
    indeterminados = {
        "", "no especificada", "no especificado", "no especifica",
        "n/a", "na", "ninguna", "ninguno", "sin marca", "sin modelo",
        "no aplica", "varios", "cualquiera", "—", "-",
    }
    return marca not in indeterminados and modelo not in indeterminados

def _construir_query_caracteristicas(info_articulo, max_terms=6):
    """Arma query de búsqueda desde nombre + características (modo propuesta)."""
    nombre = (info_articulo.get("nombre_articulo") or "").strip()
    items  = (info_articulo.get("caracteristicas", []) or []) + \
             (info_articulo.get("especificaciones_tecnicas", []) or [])
    keywords = []
    vistos = set(w.lower() for w in nombre.split())
    for item in items[:8]:
        if not isinstance(item, str): continue
        val = item.split(":", 1)[-1].strip() if ":" in item else item
        for w in re.split(r"[\s,;]+", val):
            w_low = w.strip().lower()
            if len(w_low) >= 3 and w_low not in vistos:
                keywords.append(w.strip())
                vistos.add(w_low)
                if len(keywords) >= max_terms: break
        if len(keywords) >= max_terms: break
    return f"{nombre} {' '.join(keywords)}".strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 1 – BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════════

def obtener_data_table_1():
    import mysql.connector
    log("Conectando a MySQL…")
    conn   = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT codigo_necesidad, entidad_contratante, entidad_contratante_url,
               direccion_entrega, contacto
        FROM   infimas
        WHERE  LOWER(etapa) = 'en generacion'
    """)
    rows = cursor.fetchall()
    cursor.close()
    conn.close()
    log(f"Registros encontrados en BD: {len(rows)}", "OK")
    for r in rows:
        log(f"  → {r['codigo_necesidad']} | {r['entidad_contratante']}")
    return rows


# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 2 – GOOGLE CLOUD STORAGE
# ═══════════════════════════════════════════════════════════════════════════════

def obtener_cliente_gcs():
    from google.oauth2 import service_account
    from google.cloud  import storage
    creds_path = resolver_credenciales_a_archivo()
    creds = service_account.Credentials.from_service_account_file(
        creds_path, scopes=["https://www.googleapis.com/auth/cloud-platform"]
    )
    return storage.Client(credentials=creds, project=creds.project_id)

def listar_docs_necesidad(gcs_client, codigo_necesidad):
    bucket  = gcs_client.bucket(BUCKET_NAME)
    prefijo = f"{BUCKET_FOLDER}/{codigo_necesidad}/"
    blobs   = list(gcs_client.list_blobs(bucket, prefix=prefijo))
    docs    = [b for b in blobs if b.name.lower().endswith((".doc", ".docx", ".pdf"))]
    log(f"  Docs encontrados para {codigo_necesidad}: {len(docs)}")
    return docs

def descargar_blob_a_tmp(blob):
    suffix = Path(blob.name).suffix
    tmp    = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    blob.download_to_filename(tmp.name)
    return tmp.name

def subir_archivo_a_bucket(gcs_client, local_path, carpeta_destino):
    bucket    = gcs_client.bucket(BUCKET_NAME)
    nombre    = Path(local_path).name
    blob_name = f"{carpeta_destino}/{nombre}"
    blob      = bucket.blob(blob_name)
    blob.upload_from_filename(local_path)
    log(f"  Subido: {blob_name}", "OK")
    return blob_name


# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 3 – VERTEX AI / GEMINI
# ═══════════════════════════════════════════════════════════════════════════════

def inicializar_vertex_ai():
    import vertexai
    from vertexai.generative_models import GenerativeModel
    from google.oauth2 import service_account

    creds_path = resolver_credenciales_a_archivo()
    with open(creds_path, "r", encoding="utf-8") as f:
        creds_data = json.load(f)
    creds = service_account.Credentials.from_service_account_file(creds_path)
    vertexai.init(project=creds_data["project_id"], credentials=creds,
                  location="us-central1")
    modelo = GenerativeModel(AI_MODEL)
    log(f"VertexAI inicializado con modelo: {AI_MODEL}", "OK")
    return modelo

def docx_a_pdf(docx_path):
    """Convierte .docx a PDF con LibreOffice (requerido porque Gemini no acepta .docx)."""
    import subprocess
    SOFFICE = "soffice"
    out_dir  = tempfile.mkdtemp()
    try:
        res = subprocess.run(
            [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=120,
        )
        if res.returncode != 0:
            log(f"    LibreOffice error: {res.stderr[:200]}", "WARN")
            return None
        pdf_path = Path(out_dir) / (Path(docx_path).stem + ".pdf")
        if not pdf_path.exists():
            return None
        size_mb = pdf_path.stat().st_size / (1024*1024)
        if size_mb > 18:
            log(f"    PDF demasiado grande ({size_mb:.1f} MB), omitido.", "WARN")
            shutil.rmtree(out_dir, ignore_errors=True)
            return None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        shutil.copy2(str(pdf_path), tmp.name)
        shutil.rmtree(out_dir, ignore_errors=True)
        return tmp.name
    except FileNotFoundError:
        log("    LibreOffice no encontrado. Instálalo: https://www.libreoffice.org", "ERR")
        return None
    except subprocess.TimeoutExpired:
        log("    Conversión a PDF superó 120s.", "WARN")
        return None
    except Exception as e:
        log(f"    Error al convertir a PDF: {e}", "WARN")
        return None

def blob_a_part(blob_path_local):
    from vertexai.generative_models import Part
    suffix = Path(blob_path_local).suffix.lower()
    if suffix != ".pdf":
        log(f"    Formato no soportado: {suffix}. Omitido.", "WARN")
        return None
    size_mb = Path(blob_path_local).stat().st_size / (1024*1024)
    if size_mb > 18:
        log(f"    PDF muy grande ({size_mb:.1f} MB). Omitido.", "WARN")
        return None
    with open(blob_path_local, "rb") as f:
        data = f.read()
    return Part.from_data(data=data, mime_type="application/pdf")

def llamar_gemini_con_reintentos(modelo, contents, max_intentos=3, espera_inicial=8):
    ultimo_error = None
    for intento in range(1, max_intentos + 1):
        try:
            return modelo.generate_content(contents)
        except Exception as e:
            ultimo_error = e
            msg = str(e)
            transitorio = any(t in msg.lower() for t in
                               ["503","500","429","deadline","unavailable","rate","exhausted"])
            if intento < max_intentos and transitorio:
                espera = espera_inicial * (2 ** (intento - 1))
                log(f"  Gemini error transitorio (intento {intento}/{max_intentos}), "
                    f"reintentando en {espera}s.", "WARN")
                time.sleep(espera)
                continue
            raise
    raise ultimo_error

def extraer_texto_de_office(file_path):
    """
    Extrae texto de archivos Office (docx, xlsx) detectando el tipo real por contenido.
    Devuelve el texto concatenado o cadena vacía si falla.
    """
    try:
        # Leer los primeros bytes para identificar el tipo ZIP (tanto docx como xlsx son ZIP)
        with open(file_path, 'rb') as f:
            header = f.read(4)
        if header[:2] != b'PK':
            # No es un archivo Office comprimido
            log(f"    {Path(file_path).name} no es un archivo Office válido.", "WARN")
            return ""

        # Intentar como Word primero
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if texto:
                return texto
        except Exception:
            pass

        # Si falla, intentar como Excel
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            texts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_text = " ".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        texts.append(row_text)
            texto = "\n".join(texts)
            if texto:
                return texto
        except Exception:
            pass

        log(f"    No se pudo extraer texto del archivo {Path(file_path).name}", "WARN")
        return ""
    except Exception as e:
        log(f"    Error extrayendo texto de {file_path}: {e}", "WARN")
        return ""

def analizar_documentos_con_gemini(modelo, archivos_locales, codigo_necesidad):
    partes = []
    texto_extra_docx = []   # aquí acumularemos texto de los .docx

    for path in archivos_locales:
        suf = Path(path).suffix.lower()
        if suf in (".docx", ".doc", ".xlsx"):
            # Extraer texto sin LibreOffice
            texto = extraer_texto_de_office(path)
            if texto:
                texto_extra_docx.append(f"--- Contenido del archivo {Path(path).name} ---\n{texto}")
                log(f"    Texto extraído de DOCX: {Path(path).name} ({len(texto)} caracteres)")
            else:
                log(f"    No se pudo extraer texto de {Path(path).name}, omitido.", "WARN")
        else:
            # PDF y otros formatos: se envían como Part
            try:
                part = blob_a_part(path)
                if part is None:
                    log(f"    Documento omitido: {Path(path).name}", "WARN")
                    continue
                partes.append(part)
                log(f"    Documento cargado: {Path(path).name}")
            except Exception as e:
                log(f"    No se pudo cargar {path}: {e}", "WARN")

    if not partes and not texto_extra_docx:
        log("No hay documentos válidos para analizar.", "ERR")
        return None

    # Construir el prompt combinando el texto de los DOCX
    prompt_base = f"""
Eres un analista de contratación pública especializado en fichas técnicas.
Analiza TODOS los documentos adjuntos del código de necesidad: {codigo_necesidad}.

IMPORTANTE: Identifica CADA producto/servicio solicitado en los documentos.
Para cada uno, extrae sus características completas.
"""
    # Si hay texto extra de docx, lo añadimos al prompt
    if texto_extra_docx:
        prompt_base += "\n\n===== TEXTOS EXTRAÍDOS DE ARCHIVOS DOCX =====\n"
        prompt_base += "\n\n".join(texto_extra_docx)
        prompt_base += "\n===== FIN TEXTOS DOCX =====\n"

    prompt_base += """
Devuelve ÚNICAMENTE un JSON válido (sin markdown) con esta estructura:
{
  "articulos": [
    {
      "nombre_articulo": "Nombre completo del artículo",
      "marca": "Marca exacta (cadena vacía si no se especifica)",
      "modelo": "Modelo exacto (cadena vacía si no se especifica)",
      "cantidad": 1,
      "caracteristicas": ["Característica 1", "... (al menos 7)"],
      "especificaciones_tecnicas": ["Especificación: valor", "... (al menos 10)"],
      "especificaciones_electricas": ["Especificación eléctrica: valor", "... ([] si no aplica)"],
      "incluye": ["Accesorio incluido 1"],
      "resumen": "Descripción del producto de 50 a 80 palabras.",
      "funcion_principal": "Función principal que debe cumplir este producto"
    }
  ]
}

REGLAS IMPORTANTES:
- Si hay MÚLTIPLES productos diferentes, crea un objeto por cada uno en el array 'articulos'
- Para cada producto, identifica su FUNCIÓN PRINCIPAL (para qué sirve)
- Si no se especifica marca/modelo, deja esos campos vacíos
- Si hay marca/modelo pero no se encuentra, anótalo igualmente
- Sé extremadamente detallado en características y especificaciones
"""
    # Enviar a Gemini: los PDF como partes + el prompt (que ya contiene el texto de los docx)
    if partes:
        log(f"  Enviando {len(partes)} PDF(s) y {len(texto_extra_docx)} texto(s) DOCX a Gemini para análisis…")
        contents = partes + [prompt_base]
    else:
        log(f"  Enviando solo texto de DOCX(s) ({len(texto_extra_docx)}) a Gemini para análisis…")
        contents = [prompt_base]

    response = llamar_gemini_con_reintentos(modelo, contents)
    data = _parsear_json_de_respuesta(response.text)
    if data is None:
        log(f"  Error al parsear JSON de Gemini. Raw: {response.text[:300]}", "ERR")
        return None
    log("  Análisis Gemini completado.", "OK")
    return data

# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 4 – BÚSQUEDA DE PRODUCTOS (Google Custom Search API)
# ═══════════════════════════════════════════════════════════════════════════════
def buscar_producto_en_proveedores(modelo_ai, info_articulo):
    """
    MODIFICADO: Ahora procesa múltiples artículos y busca alternativas
    """
    # Extraer array de artículos
    if isinstance(info_articulo, list):
        articulos = info_articulo
    elif isinstance(info_articulo, dict) and "articulos" in info_articulo:
        articulos = info_articulo["articulos"]
    else:
        articulos = [info_articulo]
    
    todos_resultados = []
    productos_encontrados = []
    
    for idx, articulo in enumerate(articulos):
        log(f"\n  Buscando producto {idx+1}/{len(articulos)}: {articulo.get('nombre_articulo', '')[:60]}")
        
        # Búsqueda normal primero
        resultado = _buscar_producto_unico(articulo,modelo_ai)
        
        # Si no encontró nada o poca relevancia, buscar alternativa
        if not resultado.get('mejor_opcion') or resultado['mejor_opcion'].get('relevancia', 0) < 50:
            log(f"  ⚠️ Producto no encontrado con exactitud. Buscando alternativa funcional...", "WARN")
            
            # Buscar alternativa basada en función
            alternativa, nuevos_resultados = buscar_producto_alternativo_con_ia(
                modelo_ai, articulo, resultado.get('alternativas', [])
            )
            
            if alternativa:
                # Actualizar artículo con la alternativa
                articulo['nombre_alternativo'] = alternativa.get('nombre_alternativo', '')
                articulo['marca_alternativa'] = alternativa.get('marca_sugerida', '')
                articulo['modelo_alternativo'] = alternativa.get('modelo_sugerido', '')
                articulo['justificacion_alternativa'] = alternativa.get('justificacion', '')
                
                # Buscar la alternativa
                resultado_alt = _buscar_producto_unico({
                    'nombre_articulo': alternativa.get('nombre_alternativo', ''),
                    'marca': alternativa.get('marca_sugerida', ''),
                    'modelo': alternativa.get('modelo_sugerido', ''),
                    'caracteristicas': articulo.get('caracteristicas', []),
                    'especificaciones_tecnicas': articulo.get('especificaciones_tecnicas', [])
                }, modelo_ai)
                
                if resultado_alt.get('mejor_opcion'):
                    resultado = resultado_alt
                    log(f"  ✅ Alternativa encontrada: {alternativa.get('nombre_alternativo', '')}", "OK")
        
        # Guardar resultado
        resultado['articulo_original'] = articulo
        todos_resultados.append(resultado)
        
        if resultado.get('mejor_opcion'):
            productos_encontrados.append(resultado)
    
    # Si hay múltiples productos, consolidar
    if len(todos_resultados) > 1:
        return {
            "productos": todos_resultados,
            "articulos_originales": articulos,
            "url_imagen_principal": todos_resultados[0].get('url_imagen_producto', '') if todos_resultados else '',
            "resumen_busqueda": f"Se buscaron {len(articulos)} productos, encontrados {len(productos_encontrados)}"
        }
    
    return todos_resultados[0] if todos_resultados else {}

def _buscar_producto_unico(info_articulo,modelo_ai=None):
    nombre = info_articulo.get("nombre_articulo", "") or ""
    marca = info_articulo.get("marca", "") or ""
    modelo_prod = info_articulo.get("modelo", "") or ""
    es_propuesta = not _hay_marca_modelo_especificos(info_articulo)

    search_engine = ProductSearchEngine(region="ec-es", timeout=15)

    # Búsqueda única sin filtrar dominios
    query_base = " ".join(x for x in [nombre, marca, modelo_prod] if x) if not es_propuesta else _construir_query_caracteristicas(info_articulo)
    log(f"  Modo {'EXACTO' if not es_propuesta else 'PROPUESTA'} — buscando: '{query_base}'")
    resultados_scored = search_engine.buscar_con_scoring(nombre, marca, modelo_prod, num_resultados=30)

    if not resultados_scored:
        log("  No se encontraron productos.", "WARN")
        return {"mejor_opcion": {}, "alternativas": [], "url_imagen_producto": "", "info_articulo_actualizada": None}

    # Separar por origen
    nacionales = []
    extranjeros_america = []
    otros = []
    for res in resultados_scored:
        dom = _dominio(res['url'])
        if dom in DOMINIOS_NACIONALES:
            nacionales.append(res)
        elif dom in DOMINIOS_EXTRANJEROS_AMERICA:
            extranjeros_america.append(res)
        else:
            otros.append(res)

    # Priorizar nacionales, luego americanos, luego otros
    candidatos = nacionales + extranjeros_america + otros

    # Buscar el primero con relevancia >= 50
    mejor = next((c for c in candidatos if c['relevancia'] >= 50), None)
    if not mejor:
        log("  No se encontraron productos con suficiente relevancia (≥50).", "WARN")
        return {"mejor_opcion": {}, "alternativas": [], "url_imagen_producto": "", "info_articulo_actualizada": None}

    # Mostrar los 3 mejores
    for i, res in enumerate(candidatos[:3]):
        log(f"    #{i+1}: {res['relevancia']:.0f}% - {res['url'][:70]}")

    # Obtener imagen
    url_imagen = mejor.get('imagen_url', '')
    if not url_imagen or not validar_url_imagen(url_imagen):
        log(f"  Buscando imagen específica para '{nombre}'...", "INFO")
        query_img = f"{nombre} {marca} {modelo_prod}".strip()
        url_imagen = buscar_imagen_especifica_producto(query_img, nombre)

    dominio = _dominio(mejor['url'])
    es_extranjero = dominio not in DOMINIOS_NACIONALES

    mejor_opcion = {
        "proveedor": dominio,
        "url_producto": mejor['url'],
        "precio_unitario_usd": mejor.get('precio') or 0,
        "precio_total_usd": mejor.get('precio') or 0,
        "es_extranjero": es_extranjero,
        "costos_adicionales_usd": 0,
        "detalle_costos": "",
        "nombre_en_tienda": mejor.get('titulo', nombre)[:100],
        "disponible": True,
        "relevancia": mejor['relevancia']
    }

    # Alternativas (relevancia >= 50)
    alternativas = []
    for res in candidatos[1:6]:
        if res['relevancia'] >= 50:
            dom_alt = _dominio(res['url'])
            alternativas.append({
                "proveedor": dom_alt,
                "url_producto": res['url'],
                "precio_total_usd": res.get('precio') or 0,
                "nombre_en_tienda": res.get('titulo', nombre)[:100],
                "relevancia": res['relevancia']
            })

    precio_encontrado = mejor_opcion.get('precio_total_usd', 0) or 0
    if (precio_encontrado == 0 or precio_encontrado < 1.0) and modelo_ai:
        html = fetch_html(mejor['url'], timeout=15)
        if html:
            precio_ia = extraer_precio_con_ia(html, mejor['url'], modelo_ai)
            if precio_ia:
                mejor_opcion['precio_unitario_usd'] = precio_ia
                mejor_opcion['precio_total_usd'] = precio_ia
                log(f"  Precio extraído con IA: ${precio_ia:.2f}", "OK")
    

    log(f"  ✅ Mejor opción: {dominio} → ${mejor_opcion['precio_total_usd']} ({mejor['relevancia']:.0f}% relevancia)")
    log(f"  URL: {mejor['url'][:80]}")

    return {
        "mejor_opcion": mejor_opcion,
        "alternativas": alternativas,
        "url_imagen_producto": url_imagen,
        "info_articulo_actualizada": None,
    }

def buscar_imagen_especifica_producto(query, nombre_producto):
    """
    Busca imagen específica del producto con verificación de relevancia
    """
    try:
        with DDGS() as ddgs:
            resultados = list(ddgs.images(query, region="ec-es", max_results=10, safesearch='off'))
            
            # Extraer palabras clave del producto
            keywords = set(re.findall(r'\b\w{4,}\b', nombre_producto.lower()))
            
            for r in resultados:
                img_url = r.get("image", "")
                title = r.get("title", "").lower()
                
                if not img_url or not img_url.startswith("http"):
                    continue
                
                # Verificar relevancia de la imagen
                es_relevante = False
                for kw in list(keywords)[:5]:
                    if kw in title:
                        es_relevante = True
                        break
                
                if es_relevante and validar_url_imagen(img_url):
                    log(f"  Imagen específica encontrada: {img_url[:70]}", "OK")
                    return img_url
            
            # Si no hay imagen relevante, tomar la primera válida
            for r in resultados:
                img_url = r.get("image", "")
                if img_url and validar_url_imagen(img_url):
                    log(f"  Imagen encontrada (relevancia no verificada): {img_url[:70]}", "OK")
                    return img_url
                        
    except Exception as e:
        log(f"  Error buscando imagen: {e}", "WARN")
    
    return None

def buscar_imagen_google_producto(query, palabras_clave=None):
    """
    Busca imagen específica del producto
    """
    try:
        with DDGS() as ddgs:
            resultados = list(ddgs.images(query, region="ec-es", max_results=10, safesearch='off'))
            
            for r in resultados:
                img_url = r.get("image", "")
                title = r.get("title", "").lower()
                
                if not img_url or not img_url.startswith("http"):
                    continue
                
                # Verificar que la imagen corresponda al producto
                if palabras_clave:
                    # Buscar palabras clave en el título de la imagen
                    if any(palabra in title for palabra in list(palabras_clave)[:3]):
                        if validar_url_imagen(img_url):
                            log(f"  Imagen específica encontrada: {img_url[:70]}", "OK")
                            return img_url
                
                # Si no hay palabras clave o no coinciden, tomar la primera válida
                if validar_url_imagen(img_url):
                    log(f"  Imagen encontrada: {img_url[:70]}", "OK")
                    return img_url
                        
    except Exception as e:
        log(f"  Error buscando imagen: {e}", "WARN")
    
    return None

def descargar_imagen_producto(url_imagen):
    """Descarga imagen con headers Referer para evitar bloqueos de CDN."""
    import requests, urllib.parse
    if not url_imagen or "localhost" in url_imagen:
        return None
    try:
        parsed  = urllib.parse.urlparse(url_imagen)
        referer = f"{parsed.scheme}://{parsed.netloc}/"
    except Exception:
        referer = "https://www.google.com/"

    headers = {**DEFAULT_HEADERS, "Referer": referer,
               "Accept": "image/webp,image/apng,image/*,*/*;q=0.8"}
    try:
        r = requests.get(url_imagen, headers=headers, timeout=20)
        r.raise_for_status()
        ct = r.headers.get("content-type", "image/jpeg").lower()
        if "image" not in ct:
            return None
        ext = ".png" if "png" in ct else (".webp" if "webp" in ct else ".jpg")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
        tmp.write(r.content)
        tmp.close()
        # Convertir webp a png (python-docx no soporta webp)
        if ext == ".webp":
            try:
                from PIL import Image as PILImage
                im = PILImage.open(tmp.name).convert("RGB")
                tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                tmp2.close()
                im.save(tmp2.name, "PNG")
                os.unlink(tmp.name)
                return tmp2.name
            except Exception as e:
                log(f"  No se pudo convertir webp: {e}", "WARN")
                os.unlink(tmp.name)
                return None
        log(f"  Imagen descargada: {len(r.content)//1024} KB", "OK")
        return tmp.name
    except Exception as e:
        log(f"  No se pudo descargar imagen: {e}", "WARN")
        return None

def buscar_producto_alternativo_con_ia(modelo_ai, articulo_original, resultados_busqueda):
    """
    Cuando no se encuentra el producto exacto, usa IA para proponer alternativa
    basada en la función principal del producto
    """
    if not resultados_busqueda or len(resultados_busqueda) == 0:
        # No se encontró NADA, pedir a IA que proponga
        prompt = f"""
Eres un experto en productos tecnológicos/industriales.
Se necesita un producto con estas características:

Producto solicitado: {articulo_original.get('nombre_articulo', '')}
Marca especificada: {articulo_original.get('marca', 'No especificada')}
Modelo especificado: {articulo_original.get('modelo', 'No especificado')}
Función principal: {articulo_original.get('funcion_principal', articulo_original.get('nombre_articulo', ''))}
Características requeridas: {json.dumps(articulo_original.get('caracteristicas', []), ensure_ascii=False)}
Especificaciones técnicas: {json.dumps(articulo_original.get('especificaciones_tecnicas', []), ensure_ascii=False)}

TAREA: Propone un producto ALTERNATIVO que cumpla la misma FUNCIÓN, 
aunque no tenga exactamente las mismas características.
Devuelve JSON:
{{
  "nombre_alternativo": "Nombre del producto alternativo",
  "marca_sugerida": "Marca recomendada",
  "modelo_sugerido": "Modelo recomendado", 
  "justificacion": "Por qué este producto cumple la función requerida",
  "query_busqueda": "Query optimizada para buscar este producto"
}}
"""
        response = modelo_ai.generate_content(prompt)
        data = _parsear_json_de_respuesta(response.text)
        
        if data:
            log(f"  IA sugiere alternativa: {data.get('nombre_alternativo', '')}", "OK")
            # Buscar esta alternativa
            query = data.get('query_busqueda', '')
            nuevos_resultados = duckduckgo_search_simple(query, num_results=10, region="ec-es")
            return data, nuevos_resultados
    
    return None, resultados_busqueda

# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 5 – GENERACIÓN DE FICHA TÉCNICA (.docx)
# ═══════════════════════════════════════════════════════════════════════════════

def generar_ficha_tecnica(resultados_busqueda, codigo_necesidad, directorio_salida, modelo_ai):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    productos = resultados_busqueda.get('productos', [resultados_busqueda])
    if not productos:
        return None, None

    nombre_archivo = re.sub(r"[^a-zA-Z0-9_\-]", "_", f"{codigo_necesidad}_ficha_tecnica")
    out_path = Path(directorio_salida) / f"{nombre_archivo}.docx"

    doc = Document(str(TEMPLATE_DOCX))
    cuerpo = doc.element.body
    for child in list(cuerpo):
        if not child.tag.endswith("}sectPr"):
            cuerpo.remove(child)

    # ---------- HELPERS REALES (implementación completa) ----------
    def agregar_parrafo(text, bold=False, size_pt=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        espacio_antes=0, espacio_despues=0, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(int(espacio_antes * 20)))
        sp.set(qn("w:after"),  str(int(espacio_despues * 20)))
        pPr.append(sp)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Century Gothic"
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = RGBColor(*color)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Century Gothic")
        rFonts.set(qn("w:hAnsi"), "Century Gothic")
        rPr.insert(0, rFonts)
        return p

    def agregar_item_lista(texto, size_pt=10):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"),    "360")
        ind.set(qn("w:hanging"), "360")
        pPr.append(ind)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"),  "60")
        pPr.append(sp)
        for txt in ["•  ", texto]:
            run = p.add_run(txt)
            run.font.name = "Century Gothic"
            run.font.size = Pt(size_pt)
        return p
    # ----------------------------------------------------------------

    # Título principal
    agregar_parrafo("FICHA TÉCNICA CONSOLIDADA", bold=True, size_pt=14,
                    align=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=12)
    #agregar_parrafo(f"Código: {codigo_necesidad}", size_pt=10,align=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=18)

    for idx, prod in enumerate(productos):
        articulo_original = prod.get('articulo_original', {})
        mejor_opcion = prod.get('mejor_opcion', {})
        es_alternativa = 'nombre_alternativo' in articulo_original

        # Intentar obtener info real del proveedor
        url_producto = mejor_opcion.get('url_producto', '')
        info_proveedor = None
        if url_producto:
            log(f"  Obteniendo características del proveedor para {articulo_original.get('nombre_articulo','')[:40]}")
            info_proveedor = extraer_info_proveedor(url_producto, mejor_opcion.get('nombre_en_tienda', ''), modelo_ai)

        # Datos a mostrar
        nombre_mostrar = info_proveedor.get('nombre') if info_proveedor else mejor_opcion.get('nombre_en_tienda', articulo_original.get('nombre_articulo', ''))
        precio_mostrar = mejor_opcion.get('precio_total_usd', 0) or 0
        caract_mostrar = info_proveedor.get('caracteristicas') if info_proveedor else articulo_original.get('caracteristicas', [])
        specs_mostrar = info_proveedor.get('especificaciones') if info_proveedor else articulo_original.get('especificaciones_tecnicas', [])
        descripcion_mostrar = info_proveedor.get('descripcion') if info_proveedor else articulo_original.get('resumen', '')

        # Si es página nueva, no hace falta separador, pero después del primero añadimos salto de página
        if idx > 0:
            doc.add_page_break()

        # Nombre del producto
        agregar_parrafo(f"PRODUCTO {idx+1}: {nombre_mostrar}", bold=True, size_pt=12, espacio_despues=6)

        if es_alternativa:
            agregar_parrafo("⚠️ PRODUCTO ALTERNATIVO - No coincide exactamente con lo solicitado",
                            size_pt=9, color=(255, 0, 0), espacio_despues=6)
            agregar_parrafo(f"Producto original solicitado: {articulo_original.get('nombre_articulo', '')}",
                            size_pt=9, espacio_despues=4)
            just = articulo_original.get('justificacion_alternativa')
            if just:
                agregar_parrafo(f"Justificación: {just}", size_pt=9, espacio_despues=6)

        # Datos del proveedor
        #agregar_parrafo(f"Proveedor: {mejor_opcion.get('proveedor', 'N/A')}", size_pt=10)
        #agregar_parrafo(f"Precio: ${precio_mostrar:,.2f}", size_pt=10)
        #if url_producto:
            #agregar_parrafo(f"Link de compra: {url_producto}", size_pt=9, color=(0, 0, 255))

        # Descripción
        if descripcion_mostrar:
            agregar_parrafo(descripcion_mostrar, size_pt=10, espacio_despues=6)

        # Imagen del producto
        url_imagen = prod.get('url_imagen_producto', '')
        if url_imagen:
            path_img_local = descargar_imagen_producto(url_imagen)
            if path_img_local and Path(path_img_local).exists():
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(path_img_local, width=Inches(3.5))
                except:
                    pass
            else:
                agregar_parrafo("[Imagen del producto no disponible]", size_pt=9,
                                align=WD_ALIGN_PARAGRAPH.CENTER, color=(128,128,128))

        # Características
        if caract_mostrar:
            agregar_parrafo("Características:", bold=True, size_pt=11, espacio_antes=8)
            for c in caract_mostrar[:12]:
                agregar_item_lista(c, size_pt=10)


        # Especificaciones técnicas
        if specs_mostrar:
            agregar_parrafo("Especificaciones técnicas:", bold=True, size_pt=11, espacio_antes=6)
            for s in specs_mostrar[:10]:
                agregar_item_lista(s, size_pt=10)

        # Fallback extremo: si no hay características ni especificaciones, usar resumen o función principal
        if not caract_mostrar and not specs_mostrar:
            fallback = articulo_original.get('resumen') or articulo_original.get('funcion_principal') or ''
            if fallback:
                caract_mostrar = [fallback]

        if not caract_mostrar:
            caract_mostrar = ["Información no disponible. Consulte el enlace de compra para más detalles."]

        agregar_parrafo("", size_pt=8)

    doc.save(str(out_path))
    log(f"  Ficha técnica guardada: {out_path.name}", "OK")
    return str(out_path), None

# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 7 – GENERACIÓN DE PROFORMA (.xlsm) - VERSIÓN MEJORADA
# ═══════════════════════════════════════════════════════════════════════════════

def generar_proforma(registro_bd, resultados_busqueda, directorio_salida, img_path=None):
    """
    MODIFICADO: Genera proforma para múltiples productos con links de compra.
    Ahora inserta filas adicionales en ambas hojas si hay más de 6 productos.
    """
    import openpyxl
    from openpyxl.drawing.image import Image as XlImage
    from openpyxl.utils import get_column_letter
    from openpyxl.styles import Font, PatternFill, Alignment
    from copy import copy

    codigo      = registro_bd["codigo_necesidad"]
    entidad     = str(registro_bd.get("entidad_contratante",     "") or "").upper()
    entidad_url = str(registro_bd.get("entidad_contratante_url", "") or "")
    direccion   = str(registro_bd.get("direccion_entrega",       "") or "")
    contacto    = str(registro_bd.get("contacto",                "") or "")
    ruc_ish     = codigo[4:17] if len(codigo) >= 17 else codigo

    # Extraer productos
    if isinstance(resultados_busqueda, dict):
        if 'productos' in resultados_busqueda:
            productos = resultados_busqueda['productos']
            alternativas_globales = []
            for prod in productos:
                alternativas_globales.extend(prod.get('alternativas', []))
        else:
            productos = [resultados_busqueda]
            alternativas_globales = resultados_busqueda.get('alternativas', [])
    else:
        productos = []
        alternativas_globales = []

    fecha_hoy   = datetime.date.today().strftime("%d/%m/%Y")
    nombre_arch = re.sub(r"[^a-zA-Z0-9_\-]", "_", codigo)
    out_path    = Path(directorio_salida) / f"{nombre_arch}.xlsm"

    shutil.copy2(str(TEMPLATE_XLSX), str(out_path))
    wb = openpyxl.load_workbook(str(out_path), keep_vba=True)

    # Identificar hojas
    hoja_cot = next((wb[s] for s in wb.sheetnames if "cotizaci" in s.lower()), wb.worksheets[0])
    hoja_cos = next((wb[s] for s in wb.sheetnames if "costo" in s.lower()), wb.worksheets[1] if len(wb.worksheets) > 1 else None)

    def escribir_celda(ws, ref, valor):
        try:
            for mr in ws.merged_cells.ranges:
                if ref in mr:
                    ws[f"{get_column_letter(mr.min_col)}{mr.min_row}"].value = valor
                    return
            ws[ref].value = valor
        except Exception as e:
            log(f"  Error escribiendo en {ref}: {e}", "WARN")

    # ═══════════════════════════════════════════
    # HOJA COTIZACIÓN
    # ═══════════════════════════════════════════
    log("  Llenando hoja de Cotización...")
    escribir_celda(hoja_cot, "B8",  entidad)
    escribir_celda(hoja_cot, "B9",  ruc_ish)
    escribir_celda(hoja_cot, "B10", direccion)
    escribir_celda(hoja_cot, "B11", contacto)
    escribir_celda(hoja_cot, "D12", codigo)
    escribir_celda(hoja_cot, "I8",  fecha_hoy)

    # --- Inserción dinámica de filas en Cotización ---
    FILA_INICIO_COT = 16
    MAX_FILAS_PLANTILLA = 6  # filas 16..21

    if len(productos) > MAX_FILAS_PLANTILLA:
        filas_extra = len(productos) - MAX_FILAS_PLANTILLA
        for i in range(filas_extra):
            hoja_cot.insert_rows(FILA_INICIO_COT + MAX_FILAS_PLANTILLA + i)
        for i in range(filas_extra):
            fila_src = FILA_INICIO_COT + MAX_FILAS_PLANTILLA - 1  # 21
            fila_dst = FILA_INICIO_COT + MAX_FILAS_PLANTILLA + i
            for col in ['A','B','C','D','E','F','G','H','I','J']:
                celda_src = hoja_cot[f"{col}{fila_src}"]
                celda_dst = hoja_cot[f"{col}{fila_dst}"]
                if celda_src.has_style:
                    celda_dst.font = copy(celda_src.font)
                    celda_dst.border = copy(celda_src.border)
                    celda_dst.fill = copy(celda_src.fill)
                    celda_dst.number_format = celda_src.number_format
                    celda_dst.alignment = copy(celda_src.alignment)
                    celda_dst.protection = copy(celda_src.protection)
            hoja_cot.merge_cells(f"C{fila_dst}:E{fila_dst}")   # ← Agregar esto
        
    # Llenar productos
    for i, prod in enumerate(productos):
        fila = FILA_INICIO_COT + i
        articulo = prod.get('articulo_original', {})
        nombre_producto = articulo.get('nombre_articulo', '')
        marca_producto = articulo.get('marca', '')
        modelo_producto = articulo.get('modelo', '')
        if 'nombre_alternativo' in articulo:
            nombre_mostrar = f"{articulo['nombre_alternativo']} (ALT: {nombre_producto})"
            marca_mostrar = articulo.get('marca_alternativa', marca_producto)
            modelo_mostrar = articulo.get('modelo_alternativo', modelo_producto)
        else:
            nombre_mostrar = nombre_producto
            marca_mostrar = marca_producto
            modelo_mostrar = modelo_producto
        texto_producto = f"{nombre_mostrar}\nMarca: {marca_mostrar} | Modelo: {modelo_mostrar}"
        escribir_celda(hoja_cot, f"C{fila}", texto_producto)
        escribir_celda(hoja_cot, f"G{fila}", articulo.get("cantidad", 1))
        log(f"    Producto {i+1}: {nombre_mostrar[:50]}")

    # Insertar imagen del primer producto si existe
    if img_path and Path(img_path).exists():
        try:
            img_xl = XlImage(img_path)
            img_xl.width  = 120
            img_xl.height = 90
            img_xl.anchor = f"G{FILA_INICIO_COT}"
            hoja_cot.add_image(img_xl)
            log("  Imagen insertada en hoja Cotización", "OK")
        except Exception as e:
            log(f"  No se pudo insertar imagen: {e}", "WARN")

    # ═══════════════════════════════════════════
    # HOJA COSTOS
    # ═══════════════════════════════════════════
    if hoja_cos is not None:
        log("  Llenando hoja de Costos...")
        escribir_celda(hoja_cos, "C7", entidad_url)

        FILA_INICIO_COS = 15
        MAX_FILAS_COS = 6  # filas 15..20

        if len(productos) > MAX_FILAS_COS:
            filas_extra = len(productos) - MAX_FILAS_COS
            for i in range(filas_extra):
                hoja_cos.insert_rows(FILA_INICIO_COS + MAX_FILAS_COS + i)
            for i in range(filas_extra):
                fila_src = FILA_INICIO_COS + MAX_FILAS_COS - 1  # 20
                fila_dst = FILA_INICIO_COS + MAX_FILAS_COS + i
                for col in ['A','B','C','D','E','F','G','H','I','J','K','L']:
                    celda_src = hoja_cos[f"{col}{fila_src}"]
                    celda_dst = hoja_cos[f"{col}{fila_dst}"]
                    if celda_src.has_style:
                        celda_dst.font = copy(celda_src.font)
                        celda_dst.border = copy(celda_src.border)
                        celda_dst.fill = copy(celda_src.fill)
                        celda_dst.number_format = celda_src.number_format
                        celda_dst.alignment = copy(celda_src.alignment)
                        celda_dst.protection = copy(celda_src.protection)
                hoja_cos.merge_cells(f"C{fila_dst}:D{fila_dst}")   # ← Agregar esto
                
        # Para cada producto
        for i, prod in enumerate(productos):
            fila_cos = FILA_INICIO_COS + i
            mejor_opcion = prod.get('mejor_opcion', {})
            articulo = prod.get('articulo_original', {})

            # Link del producto
            url_producto = mejor_opcion.get('url_producto', '')
            if url_producto:
                escribir_celda(hoja_cos, f"C{fila_cos}", url_producto)
                log(f"    Link producto {i+1}: {url_producto[:80]}", "OK")
                try:
                    cell = hoja_cos[f"C{fila_cos}"]
                    cell.hyperlink = url_producto
                    cell.font = Font(color="0563C1", underline="single")
                except:
                    pass
            else:
                log(f"    ⚠️ Sin link para producto {i+1}", "WARN")
                escribir_celda(hoja_cos, f"C{fila_cos}", "No disponible")

            # Precio unitario
            precio_real = mejor_opcion.get('precio_unitario_usd', 0) or mejor_opcion.get('precio_total_usd', 0)
            if precio_real:
                escribir_celda(hoja_cos, f"E{fila_cos}", precio_real)
                log(f"    Precio {i+1}: ${precio_real:.2f}")
            else:
                escribir_celda(hoja_cos, f"E{fila_cos}", 0)
                log(f"    ⚠️ Sin precio para producto {i+1}", "WARN")

            # Cantidad
            cantidad_c = articulo.get("cantidad", 1) or 1
            escribir_celda(hoja_cos, f"F{fila_cos}", cantidad_c)

            # Costo final unitario
            es_ext = mejor_opcion.get("es_extranjero", False)
            costos_ext = mejor_opcion.get("costos_adicionales_usd", 0)
            if es_ext and costos_ext:
                costo_final = round(precio_real + costos_ext / cantidad_c, 2)
            else:
                costo_final = precio_real
            if costo_final:
                escribir_celda(hoja_cos, f"G{fila_cos}", costo_final)
                log(f"    Costo final {i+1}: ${costo_final:.2f}")

            # Marcar alternativas en amarillo
            if 'nombre_alternativo' in articulo:
                try:
                    for col in ['C', 'E', 'F', 'G']:
                        cell = hoja_cos[f"{col}{fila_cos}"]
                        cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                except:
                    pass

        # Links alternativos
        log("  Agregando links alternativos...")
        urls_alt = []
        for alt in alternativas_globales[:5]:
            url = alt.get('url_producto', '')
            if url and url not in urls_alt:
                urls_alt.append(url)
        if not urls_alt:
            for prod in productos:
                for alt in prod.get('alternativas', [])[:3]:
                    url = alt.get('url_producto', '')
                    if url and url not in urls_alt:
                        urls_alt.append(url)

        filas_alt = [16, 17, 18]  # igual que antes, pero ahora podrían moverse si insertamos filas en costos?
        # Como los links alternativos se ponen en J16:J18, y esos están después de los productos,
        # si insertamos filas, esas filas de alternativas se habrán desplazado correctamente porque
        # insertamos filas *antes* de la sección de alternativas (si la plantilla tiene esa sección).
        # Si la plantilla tiene las celdas de alternativas más abajo, debemos ajustar la referencia.
        # Asumiremos que están justo debajo de los productos, así que con el insert_rows se desplazan bien.
        for i, url_a in enumerate(urls_alt[:3]):
            if url_a:
                fila = filas_alt[i]
                escribir_celda(hoja_cos, f"J{fila}", url_a)
                try:
                    cell = hoja_cos[f"J{fila}"]
                    cell.hyperlink = url_a
                    cell.font = Font(color="0563C1", underline="single")
                except:
                    pass
                log(f"    Link alternativo {i+1}: {url_a[:80]}", "OK")
        if urls_alt:
            escribir_celda(hoja_cos, "J15", "OTRAS OPCIONES")
            log(f"  Total links alternativos: {len(urls_alt[:3])}", "OK")

    # Guardar
    try:
        wb.save(str(out_path))
        log(f"  Proforma guardada: {out_path.name}", "OK")
    except Exception as e:
        log(f"  Error guardando proforma: {e}", "ERR")
        out_path_simple = Path(directorio_salida) / f"{nombre_arch}_simple.xlsx"
        wb.save(str(out_path_simple))
        log(f"  Proforma simple guardada: {out_path_simple.name}", "WARN")
        return str(out_path_simple)

    return str(out_path)


# ═══════════════════════════════════════════════════════════════════════════════
#  PASO 9 – VERIFICACIÓN Y ACTUALIZACIÓN DE BD
# ═══════════════════════════════════════════════════════════════════════════════

def verificar_blobs_en_bucket(gcs_client, blob_docx, blob_xlsx):
    bucket = gcs_client.bucket(BUCKET_NAME)
    def existe(name):
        b = bucket.blob(name)
        b.reload()
        return b.exists() and (b.size or 0) > 0
    ok_docx = ok_xlsx = False
    try: ok_docx = existe(blob_docx)
    except Exception as e: log(f"  No se pudo verificar ficha: {e}", "WARN")
    try: ok_xlsx = existe(blob_xlsx)
    except Exception as e: log(f"  No se pudo verificar proforma: {e}", "WARN")
    return ok_docx, ok_xlsx

def actualizar_etapa_bd(codigo_necesidad):
    import mysql.connector
    conn   = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE infimas SET etapa='finalizada' "
        "WHERE codigo_necesidad=%s AND LOWER(etapa)='en generacion'",
        (codigo_necesidad,),
    )
    filas = cursor.rowcount
    conn.commit(); cursor.close(); conn.close()
    if filas: log(f"  BD actualizada: {codigo_necesidad} → 'finalizada'", "OK")
    else:      log(f"  BD: no se actualizó ningún registro para {codigo_necesidad}.", "WARN")


# ═══════════════════════════════════════════════════════════════════════════════
#  FUNCIÓN PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("\n" + "═"*60)
    print("  Preform_generator — Automatización de Documentos")
    print("  VERSIÓN MEJORADA: Múltiples productos + Alternativas")
    print("═"*60 + "\n")

    log("  Usando DuckDuckGo para búsqueda de productos.", "INFO")
    log(f"  Proveedores configurados: {len(DOMINIOS_TODOS)} dominios totales.", "INFO")
    log(f"  Modelo AI: {AI_MODEL}", "INFO")

    # ═══════════════════════════════════════════
    # PASO 1: Obtener registros de BD
    # ═══════════════════════════════════════════
    paso(1, 9, "Obteniendo registros de la BD MySQL")
    data_table_1 = obtener_data_table_1()
    if not data_table_1:
        log("No hay registros 'en generacion'. Fin.", "WARN")
        return

    # ═══════════════════════════════════════════
    # PASO 2: Inicializar clientes
    # ═══════════════════════════════════════════
    paso(2, 9, "Inicializando clientes GCS y VertexAI")
    try:
        gcs_client = obtener_cliente_gcs()
        log("Cliente GCS inicializado.", "OK")
    except Exception as e:
        log(f"Error inicializando GCS: {e}", "ERR")
        return

    try:
        modelo_ai = inicializar_vertex_ai()
        log("VertexAI inicializado.", "OK")
    except Exception as e:
        log(f"Error inicializando VertexAI: {e}", "ERR")
        return

    # Directorio temporal de trabajo
    dir_salida = Path(tempfile.mkdtemp(prefix="nexus_out_"))
    log(f"Directorio de trabajo temporal: {dir_salida}")

    # Registro de errores y estadísticas
    errores = []
    procesados_exitosos = []
    total_productos_encontrados = 0
    total_productos_alternativos = 0

    # ═══════════════════════════════════════════
    # PROCESAR CADA REGISTRO
    # ═══════════════════════════════════════════
    for idx, registro in enumerate(data_table_1, 1):
        codigo = registro["codigo_necesidad"]
        print(f"\n{'━'*60}")
        print(f"  [{idx}/{len(data_table_1)}] Procesando: {codigo}")
        print(f"  Entidad: {registro.get('entidad_contratante', 'N/A')[:60]}")
        print(f"{'━'*60}")

        dir_necesidad = dir_salida / codigo
        dir_necesidad.mkdir(exist_ok=True)
        path_img = None
        archivos_locales = []

        try:
            # ═══════════════════════════════════
            # PASO 3: Descargar y analizar documentos
            # ═══════════════════════════════════
            paso(3, 9, f"Descargando y analizando documentos — {codigo}")
            
            blobs = listar_docs_necesidad(gcs_client, codigo)
            if not blobs:
                log(f"Sin documentos en bucket para {codigo}. Saltando.", "WARN")
                errores.append(f"{codigo}: sin documentos.")
                continue

            log(f"  Documentos encontrados: {len(blobs)}")
            for blob in blobs:
                local_path = descargar_blob_a_tmp(blob)
                archivos_locales.append(local_path)
                log(f"  ✓ Descargado: {blob.name.split('/')[-1]}")

            # Analizar con Gemini
            log("  Enviando documentos a Gemini para análisis...")
            info_articulo = analizar_documentos_con_gemini(
                modelo_ai, archivos_locales, codigo)
            
            if not info_articulo:
                log(f"Análisis fallido para {codigo}.", "ERR")
                errores.append(f"{codigo}: análisis Gemini fallido.")
                continue

            # Extraer artículos del análisis
            if isinstance(info_articulo, dict):
                if "articulos" in info_articulo:
                    articulos = info_articulo["articulos"]
                    if isinstance(articulos, list) and len(articulos) > 0:
                        log(f"  📦 Encontrados {len(articulos)} productos en los documentos", "OK")
                        for i, art in enumerate(articulos, 1):
                            log(f"    {i}. {art.get('nombre_articulo', 'N/A')[:50]}")
                    else:
                        log("  El array 'articulos' está vacío", "WARN")
                        errores.append(f"{codigo}: sin artículos detectados.")
                        continue
                else:
                    articulos = [info_articulo]
                    log(f"  📦 1 producto encontrado: {info_articulo.get('nombre_articulo', 'N/A')[:50]}")
            elif isinstance(info_articulo, list):
                articulos = info_articulo
                log(f"  📦 Encontrados {len(articulos)} productos (formato array)", "OK")
            else:
                log("  Formato de análisis no reconocido", "ERR")
                errores.append(f"{codigo}: formato Gemini no reconocido.")
                continue

            # ═══════════════════════════════════
            # PASO 4: Buscar productos en proveedores
            # ═══════════════════════════════════
            paso(4, 9, f"Buscando productos en proveedores — {codigo}")
            
            log(f"  Buscando {len(articulos)} producto(s)...")
            resultado_busqueda = buscar_producto_en_proveedores(modelo_ai, articulos)
            
            if not resultado_busqueda:
                log(f"  ❌ No se encontró ningún producto", "ERR")
                errores.append(f"{codigo}: sin resultados de búsqueda.")
                continue

            # Verificar resultados
            if 'productos' in resultado_busqueda:
                productos_encontrados = resultado_busqueda['productos']
                log(f"  Resultados: {len(productos_encontrados)} producto(s) encontrado(s)")
                
                # Contar productos alternativos
                for prod in productos_encontrados:
                    if prod.get('articulo_original', {}).get('nombre_alternativo'):
                        total_productos_alternativos += 1
                        log(f"    ⚠️ Producto alternativo: {prod['articulo_original'].get('nombre_alternativo', '')[:50]}", "WARN")
                    else:
                        total_productos_encontrados += 1
            else:
                log("  Resultado único encontrado", "OK")
                if resultado_busqueda.get('mejor_opcion'):
                    total_productos_encontrados += 1

            # ═══════════════════════════════════
            # PASO 5: Generar ficha técnica consolidada
            # ═══════════════════════════════════
            paso(5, 9, f"Generando ficha técnica .docx — {codigo}")
            
            log("  Creando ficha técnica con todos los productos...")
            path_docx, path_img = generar_ficha_tecnica(
                resultado_busqueda, codigo, str(dir_necesidad), modelo_ai
            )
            
            if not path_docx:
                log("  Error generando ficha técnica", "ERR")
                errores.append(f"{codigo}: error en ficha técnica.")
                continue

            # ═══════════════════════════════════
            # PASO 6: Subir ficha técnica al bucket
            # ═══════════════════════════════════
            paso(6, 9, f"Subiendo ficha técnica al bucket — {codigo}")
            
            blob_docx = subir_archivo_a_bucket(gcs_client, path_docx, "Fichas Técnicas")
            
            if not blob_docx:
                log("  Error subiendo ficha técnica", "ERR")
                errores.append(f"{codigo}: error subiendo ficha.")
                continue

            log("  Esperando 60 segundos antes de generar la proforma...")
            time.sleep(60)

            # ═══════════════════════════════════
            # PASO 7: Generar proforma con links
            # ═══════════════════════════════════
            paso(7, 9, f"Generando proforma .xlsm — {codigo}")
            
            log("  Creando proforma con links de productos...")
            path_xlsx = generar_proforma(
                registro, resultado_busqueda,
                str(dir_necesidad), img_path=path_img)
            
            if not path_xlsx:
                log("  Error generando proforma", "ERR")
                errores.append(f"{codigo}: error en proforma.")
                continue

            # ═══════════════════════════════════
            # PASO 8: Subir proforma al bucket
            # ═══════════════════════════════════
            paso(8, 9, f"Subiendo proforma al bucket — {codigo}")
            
            blob_xlsx = subir_archivo_a_bucket(gcs_client, path_xlsx, "Proformas")
            
            if not blob_xlsx:
                log("  Error subiendo proforma", "ERR")
                errores.append(f"{codigo}: error subiendo proforma.")
                continue

            # ═══════════════════════════════════
            # PASO 9: Verificar y actualizar BD
            # ═══════════════════════════════════
            paso(9, 9, f"Verificando archivos y actualizando BD — {codigo}")
            
            ok_docx, ok_xlsx = verificar_blobs_en_bucket(gcs_client, blob_docx, blob_xlsx)

            if ok_docx and ok_xlsx:
                log("  ✓ Ficha técnica en bucket verificada", "OK")
                log("  ✓ Proforma en bucket verificada", "OK")
                actualizar_etapa_bd(codigo)
                procesados_exitosos.append(codigo)
            else:
                partes = []
                if not ok_docx: 
                    partes.append("ficha no verificada")
                    log("  ✗ Ficha técnica no verificada en bucket", "ERR")
                if not ok_xlsx: 
                    partes.append("proforma no verificada")
                    log("  ✗ Proforma no verificada en bucket", "ERR")
                motivo = "; ".join(partes)
                log(f"  BD NO actualizada: {motivo}", "WARN")
                errores.append(f"{codigo}: {motivo}")

            log(f"  ✔ {codigo} procesado exitosamente", "OK")

        except Exception as e:
            log(f"  ❌ Error procesando {codigo}: {e}", "ERR")
            traceback.print_exc()
            errores.append(f"{codigo}: {e}")
        
        finally:
            # Limpiar archivos temporales
            for f in archivos_locales:
                try: 
                    os.unlink(f)
                except: 
                    pass
            if path_img:
                try: 
                    os.unlink(path_img)
                except: 
                    pass

    # ═══════════════════════════════════════════
    # RESUMEN FINAL
    # ═══════════════════════════════════════════
    print(f"\n{'═'*60}")
    print(f"  RESUMEN FINAL")
    print(f"{'═'*60}")
    print(f"  Total registros procesados: {len(data_table_1)}")
    print(f"  Procesados exitosamente:    {len(procesados_exitosos)}")
    print(f"  Productos exactos:          {total_productos_encontrados}")
    print(f"  Productos alternativos:     {total_productos_alternativos}")
    
    if procesados_exitosos:
        log(f"Exitosos: {len(procesados_exitosos)}/{len(data_table_1)}", "OK")
        for cod in procesados_exitosos:
            print(f"    ✓ {cod}")
    
    if errores:
        log(f"Errores: {len(errores)}/{len(data_table_1)}", "ERR")
        for e in errores:
            print(f"    ✗ {e}")
    
    # Limpiar directorio temporal
    try: 
        shutil.rmtree(str(dir_salida))
        log("Directorio temporal eliminado.", "INFO")
    except Exception as e:
        log(f"No se pudo eliminar directorio temporal: {e}", "WARN")

    print(f"\n{'═'*60}")
    print(f"  Preform_generator finalizado.")
    print(f"  Hora: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'═'*60}\n")

    # Retornar resumen para posible uso externo
    return {
        "total": len(data_table_1),
        "exitosos": len(procesados_exitosos),
        "errores": len(errores),
        "productos_exactos": total_productos_encontrados,
        "productos_alternativos": total_productos_alternativos,
        "lista_exitosos": procesados_exitosos,
        "lista_errores": errores
    }


if __name__ == "__main__":
    main()