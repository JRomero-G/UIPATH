#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
═══════════════════════════════════════════════════════════════════════════════
 5_Preforms_generator.py   ·   Motor de IA:  Claude Sonnet 4.6  (Vertex AI)
───────────────────────────────────────────────────────────────────────────────
 Proyecto NEXUS · IMPOCRUZ EC S.A.S.

 Genera, para cada ínfima cuantía en etapa "en generacion":
   1) Una FICHA TÉCNICA (.docx) por código de necesidad, respetando la plantilla
      "FICHA TECNICA MICROFONO Y MEMORIA.docx" (fuente, encabezado y pie).
   2) Una PROFORMA (.xlsm) a partir de "FORMATO DE PROFORMA RECREADO VACÍO.xlsm",
      conservando macros, hojas, formato y el botón de impresión PDF.

 El análisis de documentos, la búsqueda de productos/proveedores/precios/imágenes
 y la redacción de contenidos los realiza Claude Sonnet 4.6 en Vertex AI, con la
 herramienta de búsqueda web nativa (web_search_20250305).  El SCRIPT verifica de
 forma determinista que todo enlace e imagen propuestos por la IA sean reales y
 accesibles, con un ciclo de reparación si algo falla.

 Ejecutar (tres niveles por debajo de la raíz, donde vive el paquete Config):
     python 5_Preforms_generator.py
═══════════════════════════════════════════════════════════════════════════════
"""

# ── REQUISITOS (pip) ──────────────────────────────────────────────────────────
#   anthropic[vertex]  ·  google-cloud-storage  ·  mysql-connector-python
#   python-docx  ·  openpyxl  ·  requests  ·  pdfplumber  ·  pillow
#   (opcional, respaldo de imágenes)  ddgs
#   (opcional, para archivos .doc)    LibreOffice (soffice) en el sistema
#
# ── NOTAS DE VERTEX AI ────────────────────────────────────────────────────────
#   · La búsqueda web nativa de Claude debe estar HABILITADA en el proyecto:
#     org policy  constraints/vertexai.allowedPartnerModelFeatures → permitir web
#     search  (obligatorio a partir de julio 2026 para claude-sonnet-4-6).
#   · No debe haber un perímetro VPC-SC que bloquee web_search_tool.
# ──────────────────────────────────────────────────────────────────────────────

import os, sys, json, io, re, time, shutil, tempfile, zipfile, datetime, traceback
import urllib.parse, html
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

# ── Raíz del proyecto (el script vive 3 niveles por debajo, junto a Config) ────
RAIZ_PROYECTO = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(RAIZ_PROYECTO))
from Config import Global   # noqa: E402

# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIGURACIÓN GLOBAL
# ═══════════════════════════════════════════════════════════════════════════════

MYSQL_CONFIG = {
    "host":     Global.DB_HOST,
    "user":     Global.DB_USER,
    "password": Global.DB_PASSWORD,
    "database": Global.DATABASE,
    "connection_timeout": 20,
}

BUCKET_NAME        = Global.BUCKET_NAME
CARPETA_DOCS       = "Documentos de Contratación"   # entrada  (docs de contratación)
CARPETA_FICHAS     = "Fichas Técnicas"              # salida   (.docx)
CARPETA_PROFORMAS  = "Proformas"                    # salida   (.xlsm)

# ── Motor de IA · Claude Sonnet 4.6 en Vertex AI ──────────────────────────────
# En Vertex el identificador es SIN sufijo de fecha (formato "dateless" fijado).
MODELO_IA          = "claude-sonnet-4-6"
VERTEX_REGION      = "global"      # endpoint global (mejor disponibilidad/latencia)
# La búsqueda web nativa de Claude en Vertex se habilita con este encabezado beta.
BETA_WEB_SEARCH    = "web-search-2025-03-05"
MAX_TOKENS_IA      = 16000         # holgado: la ficha lleva mucho texto estructurado
MAX_BUSQUEDAS_WEB  = 18            # cota de búsquedas por solicitud (coste acotado)
MAX_REINTENTOS_IA  = 3            # reintentos ante errores transitorios de la API
MAX_CICLOS_REPARAR = 2            # ciclos de reparación de enlaces/imágenes muertos

# ── Reglas de negocio ─────────────────────────────────────────────────────────
LIMITE_ARTICULOS   = 10            # > 10 artículos distintos ⇒ no se genera proforma
ENVIO_MIN, ENVIO_MAX   = 86.0, 155.0   # USD, entregas fuera de Guayaquil
INSTAL_MIN, INSTAL_MAX = 60.0, 80.0    # USD por artículo, si requiere instalación
# Un proveedor es "compartido" cuando surte a este número de artículos o más.
UMBRAL_PROVEEDOR_COMPARTIDO = 2

# Al haber < 10 artículos, las filas de artículo sobrantes se limpian (spec §7).
# La plantilla trae en la fila 15 (hoja Cotización) un ítem de ejemplo residual
# (Ítem=1, CPC=5327000115) que NO corresponde a ningún artículo: se limpia para
# que la proforma no muestre una línea fantasma.  Poner en False para respetarla.
LIMPIAR_FILA_EJEMPLO_COTIZACION = True

# ── Plantillas (junto al script en producción) ────────────────────────────────
SCRIPT_DIR    = Path(__file__).resolve().parent
TEMPLATE_DOCX = SCRIPT_DIR / "FICHA TECNICA MICROFONO Y MEMORIA.docx"
TEMPLATE_XLSX = SCRIPT_DIR / "FORMATO DE PROFORMA RECREADO VACÍO.xlsm"

# ── HTTP (verificación de enlaces / descarga de imágenes) ─────────────────────
USER_AGENT = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
             "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
HTTP_HEADERS = {
    "User-Agent":      USER_AGENT,
    "Accept-Language": "es-ES,es;q=0.9,en;q=0.7",
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,"
                       "image/avif,image/webp,*/*;q=0.8",
}
HTTP_TIMEOUT = 20

# ── Listas de proveedores de confianza (se inyectan al prompt de la IA) ────────
PROVEEDORES_NACIONALES = [
    "https://mibodega.ec", "https://bodeguitadelahorro.com/", "https://comercialvaca.ec/",
    "https://kissu.com.ec/", "https://altatenalmacen.com.ec/", "https://www.tventas.com/",
    "https://store.intcomex.com/es-XPE/Home", "https://lavictoria.ec/", "https://www.marcimex.com/",
    "https://almacenesespana.ec/", "https://www.novicompu.com/", "https://technoprimec.com/",
    "https://point.com.ec/", "https://www.computron.com.ec/", "https://nomadaware.com.ec/",
    "https://www.idcmayoristas.com/", "https://tecnit.com.ec/", "https://tecnomegastore.ec/",
    "https://www.artefacta.com/", "https://mundotek.com.ec/", "https://eljuri.store/",
    "https://www.kywi.com.ec/", "https://tecnocostoec.com/", "https://www.almacenesjapon.com/",
    "https://electromegaecuador.com/", "https://www.compraecuador.com/", "https://granhogar.com.ec/",
    "https://miamihome-ec.com/",
]
PROVEEDORES_EXTRANJEROS = [
    "https://www.amazon.com/", "https://www.mercadolibre.com.mx/", "https://www.mercadolibre.com.co/",
    "https://www.mercadolibre.com.ar/", "https://www.mercadolibre.cl/", "https://www.mercadolibre.com.pe/",
    "https://www.mercadolibre.com.hn/", "https://www.ebay.com/", "https://www.walmart.com/",
    "https://www.bestbuy.com/", "https://www.homedepot.com/", "https://www.costco.com/",
]

def _dominio(url: str) -> str:
    try:
        return urllib.parse.urlparse(url).netloc.lower().replace("www.", "")
    except Exception:
        return ""

DOMINIOS_NACIONALES  = [d for d in (_dominio(u) for u in PROVEEDORES_NACIONALES)  if d]
DOMINIOS_EXTRANJEROS = [d for d in (_dominio(u) for u in PROVEEDORES_EXTRANJEROS) if d]

# Respaldo opcional de búsqueda de imágenes (no crítico si no está instalado).
try:
    from ddgs import DDGS
except Exception:
    try:
        from duckduckgo_search import DDGS   # nombre antiguo del paquete
    except Exception:
        DDGS = None

_URL_CACHE: Dict[str, bool] = {}   # memo de verificación de URLs (evita repetir HEAD/GET)


# ═══════════════════════════════════════════════════════════════════════════════
#  UTILIDADES DE TERMINAL  ·  pasos, banners, cronómetros y resúmenes
# ═══════════════════════════════════════════════════════════════════════════════

class C:
    """Códigos ANSI de color (se desactivan si la salida no es una TTY)."""
    _tty = sys.stdout.isatty()
    RESET = "\033[0m"  if _tty else ""
    BOLD  = "\033[1m"  if _tty else ""
    DIM   = "\033[2m"  if _tty else ""
    AZUL  = "\033[94m" if _tty else ""
    CIAN  = "\033[96m" if _tty else ""
    VERDE = "\033[92m" if _tty else ""
    AMBAR = "\033[93m" if _tty else ""
    ROJO  = "\033[91m" if _tty else ""
    GRIS  = "\033[90m" if _tty else ""

def _ahora() -> str:
    return datetime.datetime.now().strftime("%H:%M:%S")

def banner(titulo: str) -> None:
    """Banner de ETAPA principal del script."""
    linea = "═" * 79
    print(f"\n{C.CIAN}{C.BOLD}{linea}{C.RESET}")
    print(f"{C.CIAN}{C.BOLD}  {titulo}{C.RESET}")
    print(f"{C.CIAN}{C.BOLD}{linea}{C.RESET}")

def sub(titulo: str) -> None:
    """Sub-banner (por código de necesidad)."""
    print(f"\n{C.AZUL}{C.BOLD}┌─ {titulo}{C.RESET}")

def paso(msg: str) -> None:
    print(f"  {C.GRIS}{_ahora()}{C.RESET} {C.AZUL}▸{C.RESET} {msg}")

def ok(msg: str) -> None:
    print(f"  {C.GRIS}{_ahora()}{C.RESET} {C.VERDE}✔{C.RESET} {msg}")

def warn(msg: str) -> None:
    print(f"  {C.GRIS}{_ahora()}{C.RESET} {C.AMBAR}⚠{C.RESET} {C.AMBAR}{msg}{C.RESET}")

def err(msg: str) -> None:
    print(f"  {C.GRIS}{_ahora()}{C.RESET} {C.ROJO}✘{C.RESET} {C.ROJO}{msg}{C.RESET}")

def dato(clave: str, valor: Any) -> None:
    print(f"      {C.DIM}{clave}:{C.RESET} {valor}")

def fmt_dur(seg: float) -> str:
    """Formatea una duración en s → 'Xm Ys' ó 'Zs'."""
    seg = max(0.0, float(seg))
    if seg < 60:
        return f"{seg:.1f}s"
    m, s = divmod(int(round(seg)), 60)
    if m < 60:
        return f"{m}m {s}s"
    h, m = divmod(m, 60)
    return f"{h}h {m}m {s}s"


class Cronometro:
    """Cronómetro simple con arranque perezoso."""
    def __init__(self):
        self.t0 = time.perf_counter()
    def reset(self):
        self.t0 = time.perf_counter()
    def leer(self) -> float:
        return time.perf_counter() - self.t0


class Resumen:
    """
    Acumulador de métricas por etapa y global.  Al final imprime un resumen
    tabulado con el tiempo por etapa y el tiempo total de ejecución del script.
    """
    def __init__(self):
        self.etapas: List[Tuple[str, float, str]] = []   # (nombre, seg, detalle)
        self.contadores: Dict[str, int] = {}
        self.cron_global = Cronometro()

    def etapa(self, nombre: str, seg: float, detalle: str = "") -> None:
        self.etapas.append((nombre, seg, detalle))

    def sumar(self, clave: str, n: int = 1) -> None:
        self.contadores[clave] = self.contadores.get(clave, 0) + n

    def imprimir_final(self) -> None:
        total = self.cron_global.leer()
        banner("RESUMEN FINAL DE EJECUCIÓN")
        # Contadores de negocio
        c = self.contadores
        print(f"  {C.BOLD}Ínfimas procesadas:{C.RESET}            "
              f"{c.get('procesadas', 0)}")
        print(f"  {C.BOLD}Fichas técnicas generadas:{C.RESET}     "
              f"{c.get('fichas', 0)}")
        print(f"  {C.BOLD}Proformas generadas:{C.RESET}           "
              f"{c.get('proformas', 0)}")
        print(f"  {C.BOLD}Finalizadas por límite (>10):{C.RESET}  "
              f"{c.get('limite_excedido', 0)}")
        print(f"  {C.BOLD}Con errores/omitidas:{C.RESET}          "
              f"{c.get('errores', 0)}")
        if c.get('enlaces_ok') or c.get('enlaces_muertos'):
            print(f"  {C.BOLD}Enlaces verificados OK / muertos:{C.RESET} "
                  f"{c.get('enlaces_ok',0)} / {c.get('enlaces_muertos',0)}")
        if c.get('busquedas_web'):
            print(f"  {C.BOLD}Búsquedas web (IA):{C.RESET}            "
                  f"{c.get('busquedas_web',0)}")
        # Tiempos por etapa
        print(f"\n  {C.BOLD}Tiempo por etapa:{C.RESET}")
        for nombre, seg, detalle in self.etapas:
            extra = f"  {C.DIM}({detalle}){C.RESET}" if detalle else ""
            print(f"    {C.GRIS}·{C.RESET} {nombre:<42} "
                  f"{C.CIAN}{fmt_dur(seg):>9}{C.RESET}{extra}")
        print(f"\n  {C.BOLD}{C.VERDE}TIEMPO TOTAL DE EJECUCIÓN: "
              f"{fmt_dur(total)}{C.RESET}")
        print(f"{C.CIAN}{C.BOLD}{'═'*79}{C.RESET}\n")

RESUMEN = Resumen()


# ═══════════════════════════════════════════════════════════════════════════════
#  GENERADOR DE FICHA TÉCNICA  (.docx)  ·  respeta la plantilla proporcionada
# ───────────────────────────────────────────────────────────────────────────────
#  Estrategia: se ABRE la plantilla como base (así se conservan encabezado con
#  logo, pie con imágenes, márgenes, orientación y definiciones de numeración/
#  viñetas) y se REEMPLAZA únicamente el cuerpo, reconstruyéndolo con la fuente
#  original (Century Gothic).  El título va centrado en negrita 12 pt; la
#  descripción justificada; las secciones usan viñetas de la plantilla.
# ═══════════════════════════════════════════════════════════════════════════════

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FUENTE_FICHA   = "Century Gothic"
NUMID_VINETA   = 3           # numId de viñeta existente en la plantilla (numbering.xml)
ANCHO_IMG_FICHA_PULG = 4.3   # ancho máximo de la imagen de producto en la ficha


def _set_fuente(run, negrita=False, tam=None, color=None):
    """Fija la fuente Century Gothic (incl. East Asian) sobre un run."""
    run.font.name = FUENTE_FICHA
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FUENTE_FICHA)
    run.font.bold = negrita
    if tam is not None:
        run.font.size = Pt(tam)
    if color is not None:
        run.font.color.rgb = color


def _limpiar_cuerpo(doc: Document) -> None:
    """Elimina todo el contenido del cuerpo salvo el <w:sectPr> final
    (que preserva la sección, encabezado, pie, márgenes y orientación)."""
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def _p_nuevo(doc: Document):
    """Inserta un nuevo párrafo ANTES del sectPr final (para no romper la sección)."""
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    p = OxmlElement('w:p')
    if sectPr is not None:
        sectPr.addprevious(p)
    else:
        body.append(p)
    return Paragraph(p, doc._body)


def _aplicar_vineta(par):
    """Convierte un párrafo en viñeta reutilizando la numeración de la plantilla."""
    par.style = "List Paragraph"
    pPr = par._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl');  ilvl.set(qn('w:val'), "0")
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), str(NUMID_VINETA))
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)


def _titulo_producto(doc, texto):
    par = _p_nuevo(doc)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_fuente(par.add_run(texto), negrita=True, tam=12)
    return par


def _imagen_centrada(doc, ruta_img, ancho_pulg=ANCHO_IMG_FICHA_PULG):
    par = _p_nuevo(doc)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    try:
        run.add_picture(str(ruta_img), width=Emu(int(ancho_pulg * 914400)))
    except Exception as e:
        # Si la imagen no pudo insertarse, se deja constancia textual (nunca se aborta).
        _set_fuente(par.add_run(f"[imagen no disponible: {e}]"), tam=8,
                    color=RGBColor(0x99, 0x99, 0x99))
    return par


def _parrafo_texto(doc, texto, negrita=False, tam=None, justificado=False,
                   color=None):
    par = _p_nuevo(doc)
    if justificado:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_fuente(par.add_run(texto), negrita=negrita, tam=tam, color=color)
    return par


def _seccion_vinetas(doc, encabezado, items):
    """Encabezado en negrita + lista de viñetas (omite la sección si no hay items)."""
    if not items:
        return
    _parrafo_texto(doc, encabezado, negrita=True)
    for it in items:
        it = (it or "").strip()
        if not it:
            continue
        par = _p_nuevo(doc)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _aplicar_vineta(par)
        _set_fuente(par.add_run(it))


def _salto_pagina(doc):
    par = _p_nuevo(doc)
    run = par.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    run._element.append(br)


def construir_ficha_tecnica(articulos: List[dict], ruta_salida: str) -> str:
    """
    Construye la ficha técnica .docx con uno o varios artículos.
    `articulos`: lista de dicts con claves:
        nombre, imagen (ruta local o None), descripcion,
        caracteristicas[], especificaciones_tecnicas[],
        especificaciones_electricas[], incluye[], garantia_texto,
        anexos[] (rutas locales de imágenes adicionales)
    Devuelve la ruta del archivo generado.
    """
    doc = Document(str(TEMPLATE_DOCX))
    _limpiar_cuerpo(doc)

    for idx, art in enumerate(articulos):
        if idx > 0:
            _salto_pagina(doc)
        _titulo_producto(doc, art.get("nombre", "Artículo"))
        if art.get("imagen"):
            _imagen_centrada(doc, art["imagen"])
        _parrafo_texto(doc, "")   # espaciado bajo la imagen
        if art.get("descripcion"):
            _parrafo_texto(doc, art["descripcion"], justificado=True)
        _seccion_vinetas(doc, "Características:",
                         art.get("caracteristicas", []))
        _seccion_vinetas(doc, "Especificaciones técnicas:",
                         art.get("especificaciones_tecnicas", []))
        _seccion_vinetas(doc, "Especificaciones eléctricas:",
                         art.get("especificaciones_electricas", []))
        _seccion_vinetas(doc, "Incluye:", art.get("incluye", []))
        if art.get("garantia_texto"):
            _parrafo_texto(doc, art["garantia_texto"])
        # Anexos (imágenes adicionales del mismo producto)
        anexos = [a for a in art.get("anexos", []) if a]
        if anexos:
            _parrafo_texto(doc, "Anexos", negrita=True)
            for ap in anexos:
                _imagen_centrada(doc, ap, ancho_pulg=3.5)

    doc.save(ruta_salida)
    return ruta_salida


def construir_ficha_limite_excedido(ruta_salida: str, n_articulos: int) -> str:
    """
    Ficha especial cuando se superan los 10 artículos distintos: un único
    mensaje en rojo, Arial Bold 14 (spec §3.a).  Conserva encabezado y pie.
    """
    doc = Document(str(TEMPLATE_DOCX))
    _limpiar_cuerpo(doc)
    par = _p_nuevo(doc)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(
        "La cantidad de artículos supera el límite permitido de 10 artículos de compra")
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), "Arial")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    doc.save(ruta_salida)
    return ruta_salida


# ═══════════════════════════════════════════════════════════════════════════════
#  GENERADOR DE PROFORMA  (.xlsm)  ·  a partir de la plantilla, celda por celda
# ───────────────────────────────────────────────────────────────────────────────
#  Mapa de celdas (spec §7).  Existe un DESFASE de 1 fila entre hojas:
#     artículo i (i = 1..n, n ≤ 10):
#        Cotización → fila (15 + i)   ·   Costos → fila (14 + i)
#     (p.ej. artículo 1: Cotización C16/G16  ↔  Costos C15/E15/A15/G15)
#     El CPC (tomado de la BD) va en Cotización col. B (B16:B25), igual en todas
#     las líneas de artículo.
#
#  Lógica de costos (spec §7, hoja Costos):
#     E = valor unitario real del producto (sin extras).
#     A = costo extra de artículos de un MISMO proveedor, DISTRIBUIDO entre ellos.
#     G = costo extra EFECTIVO por artículo (es el que alimenta el total, ya que
#         la plantilla calcula  H = (E + G) * cantidad  y  Cotización!H = (E+G)*1.5).
#         · proveedor compartido (≥2 art.):  G = (envío+aduana del proveedor)/k + instalación
#         · proveedor único           (1 art.): G = (envío+aduana) + instalación   (directo)
#     Para proveedores compartidos, A = G (desglose "mismo proveedor"); para únicos A = 0.
#     Se SOBREESCRIBE la fórmula original de G con el valor calculado (spec: "pondrá el valor").
# ═══════════════════════════════════════════════════════════════════════════════

import openpyxl
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter

# Filas base de la tabla de artículos
FILA0_COTIZACION = 15   # artículo i → fila 15 + i  (16..25)
FILA0_COSTOS     = 14   # artículo i → fila 14 + i  (15..24)
FILA_EJEMPLO_COT = 15   # fila con ítem de ejemplo residual en Cotización
ALTURA_FILA_ART  = 220  # altura recomendada para filas de artículo (spec §7)
IMG_MAX_W_PX     = 360
IMG_MAX_H_PX     = 280

# Celdas de alternativas en hoja Costos (3 por artículo, decreciente)
CELDAS_ALTERNATIVAS = {1: ["J15", "J16", "J17"],
                       2: ["J19", "J20", "J21"],
                       3: ["J23", "J24", "J25"]}


def numero_proforma(id_infima) -> str:
    """'00100100###' → 8 dígitos fijos + id_infima a la derecha (11 en total).
    Si id_infima ≤ 3 dígitos se completa con ceros; si es mayor, invade los
    ceros finales del bloque fijo.  Ej.: 21 → '00100100021'."""
    base = "00100100000"           # 11 chars: 00100100 + 000
    s = str(id_infima).strip()
    if len(s) >= len(base):
        return s[-len(base):]
    return base[:len(base) - len(s)] + s


def _extraer_ancla_boton(template_path: str) -> Optional[str]:
    """Devuelve el <xdr:twoCellAnchor> del botón 'Imprimir PDF' (macro
    Imprimir_PDF) de la plantilla, con los namespaces xdr/a declarados en
    línea para que sea autónomo al reinyectarlo."""
    with zipfile.ZipFile(template_path) as z:
        if "xl/drawings/drawing1.xml" not in z.namelist():
            return None
        draw = z.read("xl/drawings/drawing1.xml").decode("utf-8")
    for a in re.findall(r"<xdr:twoCellAnchor.*?</xdr:twoCellAnchor>", draw, re.S):
        if "Imprimir_PDF" in a or "Rectángulo" in a:
            return a.replace(
                "<xdr:twoCellAnchor>",
                '<xdr:twoCellAnchor '
                'xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing" '
                'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">',
                1)
    return None


def _restaurar_boton_pdf(xlsm_path: str, template_path: str) -> bool:
    """openpyxl descarta las formas con macro al guardar; se reinyecta el ancla
    del botón 'Imprimir PDF' dentro del drawing1.xml generado (junto a las
    imágenes de producto ya añadidas).  El vbaProject.bin se conserva con
    keep_vba=True.  Devuelve True si el botón quedó presente."""
    boton = _extraer_ancla_boton(template_path)
    if not boton:
        return False
    tmp = xlsm_path + ".tmp"
    with zipfile.ZipFile(xlsm_path) as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "xl/drawings/drawing1.xml":
                    txt = data.decode("utf-8")
                    if "Imprimir_PDF" not in txt and "</wsDr>" in txt:
                        txt = txt.replace("</wsDr>", boton + "</wsDr>", 1)
                    data = txt.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, xlsm_path)
    with zipfile.ZipFile(xlsm_path) as z:
        return "Imprimir_PDF" in z.read("xl/drawings/drawing1.xml").decode("utf-8")


def _redimensionar(img: "XLImage", max_w=IMG_MAX_W_PX, max_h=IMG_MAX_H_PX) -> None:
    """Escala la imagen preservando proporción para que quepa en la celda combinada."""
    try:
        w, h = img.width, img.height
        if not w or not h:
            return
        factor = min(max_w / w, max_h / h, 1.0)
        img.width  = int(w * factor)
        img.height = int(h * factor)
    except Exception:
        pass


def _solo_contacto(texto: str) -> str:
    """De 'contacto' conserva únicamente nombre + (correo ó teléfono)."""
    if not texto:
        return ""
    correo = re.search(r"[\w\.\-\+]+@[\w\.\-]+\.\w+", texto)
    tel    = re.search(r"(?<!\d)(\+?\d[\d\s\-]{6,}\d)", texto)
    # Nombre: primer fragmento antes de separadores comunes o del propio dato
    nombre = re.split(r"[·|/\n]| - |,|;", texto)[0].strip()
    if correo:
        nombre = nombre.replace(correo.group(0), "").strip(" -·,;")
    partes = [p for p in [nombre, (correo.group(0) if correo else
              (tel.group(0).strip() if tel else ""))] if p]
    return " · ".join(dict.fromkeys(partes)) if partes else texto.strip()


def _distribuir_costos(articulos: List[dict], en_guayaquil: bool) -> List[dict]:
    """
    Calcula, por artículo, E (valor unitario), A (extra mismo proveedor
    distribuido) y G (extra efectivo por artículo).  Modifica y devuelve la
    lista de artículos con las claves '_E', '_A', '_G' añadidas.
    Cada artículo debe traer: valor_unitario, proveedor (dominio), es_extranjero,
    costo_envio, costo_aduana, requiere_instalacion, costo_instalacion.
    """
    # Agrupar por proveedor
    grupos: Dict[str, List[dict]] = {}
    for art in articulos:
        prov = (art.get("proveedor") or _dominio(art.get("url", "")) or "?").lower()
        art["_prov"] = prov
        grupos.setdefault(prov, []).append(art)

    for prov, arts in grupos.items():
        k = len(arts)
        # Extra de envío por proveedor (0 si la entrega es en Guayaquil).
        envio  = 0.0 if en_guayaquil else float(arts[0].get("costo_envio", 0.0) or 0.0)
        envio  = min(max(envio, 0.0), ENVIO_MAX) if envio else 0.0
        # Aduana: sólo si algún artículo del proveedor es extranjero.
        aduana = max((float(a.get("costo_aduana", 0.0) or 0.0) for a in arts), default=0.0)
        extra_envio_prov = envio + aduana
        compartido = (k >= UMBRAL_PROVEEDOR_COMPARTIDO)
        for a in arts:
            instal = float(a.get("costo_instalacion", 0.0) or 0.0) \
                     if a.get("requiere_instalacion") else 0.0
            instal = min(instal, INSTAL_MAX) if instal else 0.0
            if compartido:
                efectivo = extra_envio_prov / k + instal
                a["_A"] = round(efectivo, 2)      # desglose "mismo proveedor"
                a["_G"] = round(efectivo, 2)      # efectivo (alimenta el total)
            else:
                efectivo = extra_envio_prov + instal
                a["_A"] = 0.0
                a["_G"] = round(efectivo, 2)      # directo (proveedor único)
            a["_E"] = round(float(a.get("valor_unitario", 0.0) or 0.0), 2)
    return articulos


def construir_proforma(datos: dict, articulos: List[dict],
                       imagenes: Dict[int, str], ruta_salida: str) -> str:
    """
    Genera la proforma .xlsm a partir de la plantilla.
    `datos`: dict con entidad_contratante, entidad_contratante_url, ruc13,
             direccion_entrega, contacto, codigo_necesidad, id_infima,
             tiempo_entrega_texto, garantia_texto, en_guayaquil.
    `articulos`: lista (≤10) ya enriquecida por _distribuir_costos, con claves:
             nombre, cantidad, url, alternativas[], _E, _A, _G.
    `imagenes`: {indice_articulo(0-based): ruta_imagen_local}.
    """
    wb = openpyxl.load_workbook(str(TEMPLATE_XLSX), keep_vba=True, data_only=False)
    ws_cot = wb["Cotización"]
    ws_cos = wb["Costos"]
    n = len(articulos)

    # ── Hoja Cotización · cabecera ────────────────────────────────────────────
    ws_cot["B8"]  = (datos.get("entidad_contratante") or "").upper()
    ws_cot["B9"]  = datos.get("ruc13", "")
    ws_cot["B10"] = datos.get("direccion_entrega", "")
    ws_cot["B11"] = _solo_contacto(datos.get("contacto", ""))
    ws_cot["D12"] = datos.get("codigo_necesidad", "")
    ws_cot["I3"]  = numero_proforma(datos.get("id_infima", ""))
    ws_cot["I8"]  = datetime.date.today()
    ws_cot["I8"].number_format = "DD/MM/YYYY"
    if datos.get("tiempo_entrega_texto"):
        ws_cot["I10"] = datos["tiempo_entrega_texto"]
    if datos.get("garantia_texto"):
        ws_cot["I12"] = datos["garantia_texto"]

    # Fila de ejemplo residual (ítem 1 / CPC de muestra) → se limpia.
    if LIMPIAR_FILA_EJEMPLO_COTIZACION:
        for col in ("A", "B", "F", "I"):
            ws_cot[f"{col}{FILA_EJEMPLO_COT}"] = None

    # ── Hoja Costos · cabecera ────────────────────────────────────────────────
    ws_cos["C7"] = datos.get("entidad_contratante_url", "")

    # CPC de la ínfima (tomado de la BD): mismo valor para todos los artículos,
    # se escribe en la columna B de cada línea (B16:B25) de la hoja Cotización.
    _cpc = datos.get("cpc")
    cpc = _cpc if (_cpc is not None and str(_cpc).strip() != "") else None

    # ── Artículos (ambas hojas) ───────────────────────────────────────────────
    for i, art in enumerate(articulos, start=1):
        rc = FILA0_COTIZACION + i     # fila en Cotización
        rk = FILA0_COSTOS     + i     # fila en Costos
        # Cotización: nombre/marca/modelo + cantidad
        ws_cot[f"C{rc}"] = art.get("nombre", "")
        if cpc is not None:
            ws_cot[f"B{rc}"] = cpc     # CPC (mismo para todos los artículos)
        ws_cot[f"G{rc}"] = art.get("cantidad", 0)
        ws_cot.row_dimensions[rc].height = ALTURA_FILA_ART
        # Imagen del producto anclada en la celda combinada C{rc}:E{rc}
        ruta_img = imagenes.get(i - 1)
        if ruta_img and os.path.exists(ruta_img):
            try:
                xi = XLImage(ruta_img)
                _redimensionar(xi)
                xi.anchor = f"C{rc}"
                ws_cot.add_image(xi)
            except Exception as e:
                warn(f"No se pudo insertar imagen del artículo {i} en la proforma: {e}")
        # Costos: URL, valor unitario (E), extra mismo proveedor (A), extra efectivo (G)
        ws_cos[f"C{rk}"] = art.get("url", "")
        ws_cos[f"E{rk}"] = art.get("_E", 0.0)
        ws_cos[f"A{rk}"] = art.get("_A", 0.0)
        ws_cos[f"G{rk}"] = art.get("_G", 0.0)     # sobreescribe la fórmula original
        # Alternativas (sólo para los 3 primeros artículos, orden decreciente)
        if i in CELDAS_ALTERNATIVAS:
            alts = (art.get("alternativas") or [])[:3]
            for celda, url in zip(CELDAS_ALTERNATIVAS[i], alts):
                if url:
                    ws_cos[celda] = url

    # ── Limpiar filas de artículo sobrantes en Cotización (spec §7) ───────────
    for i in range(n + 1, LIMITE_ARTICULOS + 1):
        rc = FILA0_COTIZACION + i
        for col in ("A", "B", "C", "F", "G", "H", "I"):
            ws_cot[f"{col}{rc}"] = None
        # (En Costos las filas sobrantes se anulan solas: F=Cotización!G(vacío)→0.)

    # ── Guardar y restaurar el botón de impresión PDF ─────────────────────────
    wb.save(ruta_salida)
    if not _restaurar_boton_pdf(ruta_salida, str(TEMPLATE_XLSX)):
        warn("No se pudo restaurar el botón 'Imprimir PDF' (revisar plantilla).")
    return ruta_salida


# ═══════════════════════════════════════════════════════════════════════════════
#  CREDENCIALES  ·  resolución con orden de prioridad (compatibilidad NEXUS)
# ═══════════════════════════════════════════════════════════════════════════════
#  Se resuelve el JSON de la cuenta de servicio y se expone vía
#  GOOGLE_APPLICATION_CREDENTIALS, de modo que tanto google-cloud-storage como
#  AnthropicVertex (que usa ADC) tomen la misma credencial.
#  El typo 'RENDER_CRENDENTIALS_JSON' es intencional y load-bearing en NEXUS.
# ═══════════════════════════════════════════════════════════════════════════════

def _resolver_credenciales() -> Tuple[str, str]:
    """Devuelve (ruta_json_credenciales, project_id).  Escribe el JSON a un
    archivo temporal y fija GOOGLE_APPLICATION_CREDENTIALS."""
    raw = None
    for nombre in ("GEMINI_CREDENTIALS", "RENDER_CRENDENTIALS_JSON", "CREDENTIALS_GEMINI"):
        val = getattr(Global, nombre, None)
        if val:
            raw = val
            paso(f"Credenciales tomadas de Global.{nombre}")
            break
    if not raw:
        raise RuntimeError("No se encontró ninguna credencial "
                           "(GEMINI_CREDENTIALS / RENDER_CRENDENTIALS_JSON / CREDENTIALS_GEMINI).")

    # `raw` puede ser un dict, un JSON en texto o una ruta a un archivo .json
    if isinstance(raw, dict):
        cred = raw
    elif isinstance(raw, str) and os.path.exists(raw):
        with open(raw, encoding="utf-8") as f:
            cred = json.load(f)
    else:
        cred = json.loads(raw)

    fd, ruta = tempfile.mkstemp(suffix=".json", prefix="nexus_cred_")
    with os.fdopen(fd, "w", encoding="utf-8") as f:
        json.dump(cred, f)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = ruta
    project_id = cred.get("project_id") or getattr(Global, "PROJECT_ID", "") or ""
    return ruta, project_id


# ═══════════════════════════════════════════════════════════════════════════════
#  BASE DE DATOS  ·  MySQL (Cloud SQL)  ·  tabla gestorex.infimas
# ═══════════════════════════════════════════════════════════════════════════════

import mysql.connector

def conectar_db():
    return mysql.connector.connect(**MYSQL_CONFIG)


def obtener_data_table_1() -> List[dict]:
    """
    data_table_1 = filas de `infimas` con etapa = 'en generacion'.
    Nombres de columnas literales (incluida la tilde de 'direccion_entrega').
    """
    sql = ("SELECT codigo_necesidad, entidad_contratante, entidad_contratante_url, "
           "`direccion_entrega`, id_infima, contacto, CPC "
           "FROM infimas WHERE etapa = %s")
    conn = conectar_db()
    try:
        cur = conn.cursor(dictionary=True)
        cur.execute(sql, ("en generacion",))
        filas = cur.fetchall()
        cur.close()
    finally:
        conn.close()
    # Normalizar la clave con tilde a un nombre cómodo
    for f in filas:
        f["direccion_entrega"] = f.get("direccion_entrega", "")
    return filas


def marcar_finalizada(id_infima) -> None:
    """Cambia etapa de 'en generacion' a 'finalizada' para la ínfima dada."""
    conn = conectar_db()
    try:
        cur = conn.cursor()
        cur.execute("UPDATE infimas SET etapa = %s WHERE id_infima = %s",
                    ("finalizada", id_infima))
        conn.commit()
        cur.close()
    finally:
        conn.close()


# ═══════════════════════════════════════════════════════════════════════════════
#  GOOGLE CLOUD STORAGE  ·  descarga de documentos y subida de resultados
# ═══════════════════════════════════════════════════════════════════════════════

from google.cloud import storage

_GCS_CLIENT = None

def _bucket():
    global _GCS_CLIENT
    if _GCS_CLIENT is None:
        _GCS_CLIENT = storage.Client()
    return _GCS_CLIENT.bucket(BUCKET_NAME)


EXT_DOCS = (".pdf", ".doc", ".docx")

def descargar_documentos(codigo_necesidad: str, destino: str) -> List[str]:
    """Descarga a `destino` los archivos doc/docx/pdf de la carpeta del bucket
    'Documentos de Contratación/<codigo_necesidad>/'.  Devuelve rutas locales."""
    os.makedirs(destino, exist_ok=True)
    prefijo = f"{CARPETA_DOCS}/{codigo_necesidad}/"
    rutas = []
    for blob in _bucket().list_blobs(prefix=prefijo):
        nombre = blob.name[len(prefijo):]
        if not nombre or nombre.endswith("/"):
            continue
        if not nombre.lower().endswith(EXT_DOCS):
            continue
        local = os.path.join(destino, os.path.basename(nombre))
        blob.download_to_filename(local)
        rutas.append(local)
    return rutas


def subir_a_bucket(ruta_local: str, carpeta_bucket: str,
                   content_type: Optional[str] = None) -> str:
    """Sube un archivo a 'carpeta_bucket/<nombre>' y devuelve la ruta en el bucket."""
    nombre = os.path.basename(ruta_local)
    destino = f"{carpeta_bucket}/{nombre}"
    blob = _bucket().blob(destino)
    if content_type:
        blob.upload_from_filename(ruta_local, content_type=content_type)
    else:
        blob.upload_from_filename(ruta_local)
    return destino


# ═══════════════════════════════════════════════════════════════════════════════
#  EXTRACCIÓN DE DOCUMENTOS  ·  pdf / docx / doc → texto  (+ bloques PDF para IA)
# ═══════════════════════════════════════════════════════════════════════════════

import requests
import base64 as _b64

# Sesión HTTP con reintentos
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

def _sesion_http() -> requests.Session:
    s = requests.Session()
    s.headers.update(HTTP_HEADERS)
    retry = Retry(total=2, backoff_factor=0.4,
                  status_forcelist=[429, 500, 502, 503, 504],
                  allowed_methods=["HEAD", "GET"])
    ad = HTTPAdapter(max_retries=retry)
    s.mount("http://", ad); s.mount("https://", ad)
    return s

SESION = _sesion_http()

MAX_PDF_MB_IA   = 12     # tope de PDF a adjuntar como bloque para la IA
MAX_PDF_ADJUNTOS = 4     # nº máx. de PDFs adjuntos por código de necesidad


def _texto_de_pdf(ruta: str) -> str:
    try:
        import pdfplumber
        partes = []
        with pdfplumber.open(ruta) as pdf:
            for pag in pdf.pages:
                partes.append(pag.extract_text() or "")
        return "\n".join(partes).strip()
    except Exception:
        try:
            from pypdf import PdfReader
            return "\n".join((p.extract_text() or "") for p in PdfReader(ruta).pages).strip()
        except Exception as e:
            warn(f"No se pudo extraer texto del PDF {os.path.basename(ruta)}: {e}")
            return ""


def _texto_de_docx(ruta: str) -> str:
    try:
        d = Document(ruta)
        partes = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for fila in t.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
                if celdas:
                    partes.append(" | ".join(celdas))
        return "\n".join(partes).strip()
    except Exception as e:
        warn(f"No se pudo extraer texto del DOCX {os.path.basename(ruta)}: {e}")
        return ""


def _doc_a_docx(ruta: str, destino: str) -> Optional[str]:
    """Convierte .doc → .docx con LibreOffice si está disponible."""
    for soffice in ("soffice", "libreoffice"):
        if shutil.which(soffice):
            try:
                import subprocess
                subprocess.run([soffice, "--headless", "--convert-to", "docx",
                                "--outdir", destino, ruta],
                               check=True, timeout=120,
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                cand = os.path.join(destino, Path(ruta).stem + ".docx")
                return cand if os.path.exists(cand) else None
            except Exception:
                return None
    return None


def extraer_documentos(rutas: List[str], carpeta_tmp: str) -> Tuple[str, List[dict]]:
    """
    Devuelve (texto_consolidado, bloques_pdf) donde bloques_pdf es una lista de
    bloques de contenido tipo 'document' (base64) para adjuntar a la IA.
    """
    textos, bloques = [], []
    pdfs_adjuntos = 0
    for ruta in rutas:
        low = ruta.lower()
        nombre = os.path.basename(ruta)
        if low.endswith(".pdf"):
            texto = _texto_de_pdf(ruta)
            # Adjuntar el PDF como bloque para máxima fidelidad (tablas, escaneos)
            tam_mb = os.path.getsize(ruta) / (1024 * 1024)
            if pdfs_adjuntos < MAX_PDF_ADJUNTOS and tam_mb <= MAX_PDF_MB_IA:
                try:
                    with open(ruta, "rb") as f:
                        b64 = _b64.standard_b64encode(f.read()).decode("ascii")
                    bloques.append({"type": "document",
                                    "source": {"type": "base64",
                                               "media_type": "application/pdf",
                                               "data": b64}})
                    pdfs_adjuntos += 1
                except Exception:
                    pass
        elif low.endswith(".docx"):
            texto = _texto_de_docx(ruta)
        elif low.endswith(".doc"):
            conv = _doc_a_docx(ruta, carpeta_tmp)
            texto = _texto_de_docx(conv) if conv else ""
            if not texto:
                warn(f"No se pudo procesar {nombre} (.doc). Instale LibreOffice para soportarlo.")
        else:
            continue
        if texto:
            textos.append(f"### DOCUMENTO: {nombre}\n{texto}")
    return "\n\n".join(textos).strip(), bloques


# ═══════════════════════════════════════════════════════════════════════════════
#  VERIFICACIÓN DE ENLACES E IMÁGENES  (el script confirma lo que propone la IA)
# ═══════════════════════════════════════════════════════════════════════════════

def verificar_url(url: str) -> bool:
    """True si la URL responde 200 y no es una página de error evidente."""
    if not url or not url.lower().startswith(("http://", "https://")):
        return False
    if url in _URL_CACHE:
        return _URL_CACHE[url]
    ok_ = False
    try:
        r = SESION.get(url, timeout=HTTP_TIMEOUT, allow_redirects=True, stream=True)
        code = r.status_code
        # Algunas tiendas responden 405/403 a bots pero la URL existe; se acepta 200.
        if code == 200:
            ct = r.headers.get("Content-Type", "").lower()
            ok_ = ("text/html" in ct or "application" in ct or ct == "")
        r.close()
    except Exception:
        ok_ = False
    _URL_CACHE[url] = ok_
    RESUMEN.sumar("enlaces_ok" if ok_ else "enlaces_muertos")
    return ok_


def descargar_imagen(url: str, destino: str, min_px: int = 120) -> Optional[str]:
    """Descarga una imagen y confirma que es válida y de tamaño razonable.
    Devuelve la ruta local o None."""
    if not url or not url.lower().startswith(("http://", "https://")):
        return None
    try:
        r = SESION.get(url, timeout=HTTP_TIMEOUT, allow_redirects=True)
        if r.status_code != 200 or "image" not in r.headers.get("Content-Type", "").lower():
            # Reintento sin filtrar content-type (algunos CDNs lo omiten)
            if r.status_code != 200:
                return None
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(r.content))
        img.verify()
        img = PILImage.open(io.BytesIO(r.content))
        if min(img.size) < min_px:
            return None
        ext = {"JPEG": ".jpg", "PNG": ".png", "WEBP": ".webp",
               "GIF": ".gif"}.get(img.format, ".img")
        if img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB"); ext = ".jpg"
        ruta = os.path.join(destino, f"img_{abs(hash(url)) % 10**8}{ext}")
        img.save(ruta)
        return ruta
    except Exception:
        return None


# ═══════════════════════════════════════════════════════════════════════════════
#  MOTOR DE IA  ·  Claude Sonnet 4.6 en Vertex AI  (búsqueda web nativa)
# ═══════════════════════════════════════════════════════════════════════════════

from anthropic import AnthropicVertex

_IA_CLIENT = None
_PROJECT_ID = ""

def _cliente_ia():
    global _IA_CLIENT
    if _IA_CLIENT is None:
        _IA_CLIENT = AnthropicVertex(region=VERTEX_REGION, project_id=_PROJECT_ID)
    return _IA_CLIENT


HERRAMIENTA_WEB = {"type": "web_search_20250305", "name": "web_search",
                   "max_uses": MAX_BUSQUEDAS_WEB}

SYSTEM_PROMPT = f"""Eres un analista experto en contratación pública del Ecuador (SERCOP) y \
en búsqueda de productos en tiendas en línea. Trabajas para IMPOCRUZ EC S.A.S. Tu tarea es \
analizar documentos de contratación (ínfimas cuantías) y producir, en ESPAÑOL, la información \
necesaria para generar una ficha técnica y una proforma.

REGLAS DE ANÁLISIS
1. Identifica los ARTÍCULOS DE COMPRA DISTINTOS solicitados. Para cada uno determina nombre, \
marca y modelo EXACTOS tal como se piden en los documentos. Si el documento no especifica \
nombre/marca/modelo, propón un producto real que cumpla las características solicitadas.
2. Si hay MÁS DE {LIMITE_ARTICULOS} artículos distintos, responde únicamente con \
{{"supera_limite": true, "num_articulos": <n>}} y nada más.
3. Los artículos deben ser NUEVOS. La información de la web y las características de los \
documentos deben coincidir en más del 70 % (idealmente 100 %).
4. Busca cada artículo, en ESTE ORDEN DE PRIORIDAD, usando la herramienta web:
   a) Proveedores nacionales de confianza: {', '.join(DOMINIOS_NACIONALES)}.
   b) Sólo si no se encuentra, proveedores extranjeros de confianza: {', '.join(DOMINIOS_EXTRANJEROS)}.
   c) Sólo si se agotan, propón un proveedor de Ecuador fuera de la lista.
   d) Sólo como último recurso, un proveedor extranjero fuera de la lista.
   Busca hasta 4 opciones por artículo. Conserva los enlaces TAL CUAL, actuales y reales.
5. Elige la MEJOR OPCIÓN: que cumpla exactamente lo pedido y con el MENOR precio unitario. \
Si la mejor es extranjera, considera envío a Guayaquil, aduana y logística de modo que su \
costo final siga siendo el menor; indica esos costos por separado.
6. Determina la CANTIDAD de cada artículo, el TIEMPO DE ENTREGA (en días) y la GARANTÍA \
técnica (en meses o años) desde los documentos. El tiempo de entrega y la garantía son los \
mismos para todos los artículos del código de necesidad.
7. Evalúa si la ENTREGA es en Guayaquil (revisa la direccion de entrega). Si NO lo es, hay \
costo de envío entre {ENVIO_MIN:.0f} y {ENVIO_MAX:.0f} USD según distancia/accesibilidad. \
Si un artículo requiere instalación, añade mano de obra entre {INSTAL_MIN:.0f} y \
{INSTAL_MAX:.0f} USD por artículo según complejidad.
8. Para cada artículo redacta: una DESCRIPCIÓN de 50 a 80 palabras (tomada de la web del \
proveedor o del fabricante), al menos 7 CARACTERÍSTICAS generales, al menos 10 \
ESPECIFICACIONES TÉCNICAS, especificaciones ELÉCTRICAS si aplica, y lo que INCLUYE (extras). \
Provee una imagen del producto (imagen_url) desde la página del proveedor o fabricante, y \
opcionalmente imágenes de anexos.

FORMATO DE SALIDA
Responde EXCLUSIVAMENTE con un objeto JSON válido (sin markdown, sin texto adicional) con la \
estructura:
{{
  "supera_limite": false,
  "num_articulos": <int>,
  "tiempo_entrega_texto": "<p.ej. '15 DÍAS'>",
  "garantia_texto": "<p.ej. '12 MESES'>",
  "en_guayaquil": <bool>,
  "articulos": [
    {{
      "nombre": "<nombre marca modelo>",
      "descripcion": "<50-80 palabras>",
      "caracteristicas": ["..."],
      "especificaciones_tecnicas": ["..."],
      "especificaciones_electricas": ["..."],
      "incluye": ["..."],
      "cantidad": <int>,
      "requiere_instalacion": <bool>,
      "costo_instalacion": <float>,
      "mejor_opcion": {{
        "url": "<url real del producto>",
        "proveedor": "<dominio>",
        "es_extranjero": <bool>,
        "valor_unitario": <float>,
        "costo_envio": <float>,
        "costo_aduana": <float>,
        "imagen_url": "<url de imagen del producto>",
        "imagenes_anexos": ["<url>", "..."]
      }},
      "alternativas": ["<url2>", "<url3>", "<url4>"]
    }}
  ]
}}
Todas las URLs deben ser reales y accesibles. Prioriza precios bajos y coincidencia exacta."""


def _extraer_json(texto: str) -> Optional[dict]:
    """Extrae el primer objeto JSON del texto (tolerante a ``` y preámbulos)."""
    if not texto:
        return None
    t = texto.strip()
    t = re.sub(r"^```(?:json)?|```$", "", t, flags=re.M).strip()
    ini = t.find("{")
    if ini == -1:
        return None
    # Buscar el cierre equilibrando llaves
    prof, fin = 0, -1
    for i in range(ini, len(t)):
        if t[i] == "{": prof += 1
        elif t[i] == "}":
            prof -= 1
            if prof == 0:
                fin = i; break
    if fin == -1:
        return None
    try:
        return json.loads(t[ini:fin + 1])
    except Exception:
        try:
            return json.loads(t[ini:fin + 1].replace(",\n}", "\n}").replace(",}", "}"))
        except Exception:
            return None


def _texto_respuesta(resp) -> str:
    return "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")


def _contar_busquedas(resp) -> int:
    try:
        stu = getattr(resp.usage, "server_tool_use", None)
        return getattr(stu, "web_search_requests", 0) or 0
    except Exception:
        return 0


def _llamar_ia(messages, con_web=True):
    """Invoca a Claude en Vertex con reintentos ante errores transitorios."""
    kwargs = dict(model=MODELO_IA, max_tokens=MAX_TOKENS_IA,
                  system=SYSTEM_PROMPT, messages=messages)
    if con_web:
        kwargs["tools"] = [HERRAMIENTA_WEB]
        kwargs["extra_headers"] = {"anthropic-beta": BETA_WEB_SEARCH}
    ult = None
    for intento in range(1, MAX_REINTENTOS_IA + 1):
        try:
            return _cliente_ia().messages.create(**kwargs)
        except Exception as e:
            ult = e
            warn(f"Error IA (intento {intento}/{MAX_REINTENTOS_IA}): {e}")
            time.sleep(2 * intento)
    raise ult


def analizar_necesidad(datos: dict, texto_docs: str, bloques_pdf: List[dict]) -> Optional[dict]:
    """Envía documentos + contexto a Claude Sonnet 4.6 y devuelve el JSON analizado."""
    contexto = (
        f"CÓDIGO DE NECESIDAD: {datos.get('codigo_necesidad','')}\n"
        f"ENTIDAD CONTRATANTE: {datos.get('entidad_contratante','')}\n"
        f"DIRECCIÓN DE ENTREGA: {datos.get('direccion_entrega','')}\n\n"
        f"DOCUMENTOS DE CONTRATACIÓN (texto extraído):\n{texto_docs or '(sin texto extraíble)'}"
    )
    contenido = [{"type": "text", "text": contexto}] + bloques_pdf
    resp = _llamar_ia([{"role": "user", "content": contenido}], con_web=True)
    RESUMEN.sumar("busquedas_web", _contar_busquedas(resp))
    data = _extraer_json(_texto_respuesta(resp))

    # Ciclo de reparación: verificar enlaces de la mejor opción y pedir reemplazos.
    if data and not data.get("supera_limite"):
        messages = [{"role": "user", "content": contenido}]
        for ciclo in range(MAX_CICLOS_REPARAR):
            muertos = _enlaces_muertos(data)
            if not muertos:
                break
            paso(f"Reparando {len(muertos)} enlace(s) inválido(s) (ciclo {ciclo+1})…")
            messages.append({"role": "assistant", "content": resp.content})
            messages.append({"role": "user", "content":
                "Las siguientes URLs NO son accesibles o no existen: "
                + "; ".join(muertos) +
                ". Búscalas de nuevo y devuelve el MISMO JSON completo y corregido, "
                "con URLs reales y verificables (mejor_opcion.url, imagen_url y alternativas)."})
            resp = _llamar_ia(messages, con_web=True)
            RESUMEN.sumar("busquedas_web", _contar_busquedas(resp))
            nueva = _extraer_json(_texto_respuesta(resp))
            if nueva:
                data = nueva
    return data


def _enlaces_muertos(data: dict) -> List[str]:
    """Lista de URLs de mejor_opcion (producto + imagen) que no verifican."""
    muertos = []
    for art in data.get("articulos", []):
        mo = art.get("mejor_opcion", {}) or {}
        url = mo.get("url", "")
        if url and not verificar_url(url):
            muertos.append(url)
    return muertos


def imagen_corresponde_al_producto(ruta_img: str, nombre: str) -> bool:
    """Verificación por VISIÓN: pregunta a Claude si la imagen corresponde al producto."""
    try:
        with open(ruta_img, "rb") as f:
            b64 = _b64.standard_b64encode(f.read()).decode("ascii")
        media = "image/png" if ruta_img.lower().endswith(".png") else "image/jpeg"
        resp = _cliente_ia().messages.create(
            model=MODELO_IA, max_tokens=10, system="Responde únicamente SI o NO.",
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64",
                 "media_type": media, "data": b64}},
                {"type": "text", "text": f"¿Esta imagen corresponde al producto "
                 f"'{nombre}'? Responde SI o NO."}]}])
        return _texto_respuesta(resp).strip().upper().startswith("SI")
    except Exception:
        return True   # ante fallo de verificación, no se descarta la imagen


# ═══════════════════════════════════════════════════════════════════════════════
#  RESPALDO DE IMÁGENES  (DDGS)  ·  sólo si la imagen de la IA falla
# ═══════════════════════════════════════════════════════════════════════════════

def _buscar_imagen_fallback(consulta: str, destino: str) -> Optional[str]:
    if DDGS is None:
        return None
    try:
        with DDGS() as d:
            for r in d.images(consulta, max_results=6):
                url = r.get("image") or r.get("thumbnail")
                ruta = descargar_imagen(url, destino) if url else None
                if ruta:
                    return ruta
    except Exception:
        pass
    return None


# ═══════════════════════════════════════════════════════════════════════════════
#  PREPARACIÓN DE ARTÍCULOS  ·  del JSON de la IA a las estructuras de salida
# ═══════════════════════════════════════════════════════════════════════════════

def _preparar_articulos(data: dict, carpeta_tmp: str) -> Tuple[List[dict], List[dict], Dict[int, str]]:
    """
    Devuelve (articulos_ficha, articulos_proforma, imagenes_proforma).
    Descarga y verifica imágenes de producto y anexos por el camino.
    """
    art_ficha, art_prof, imgs_prof = [], [], {}
    gt = data.get("garantia_texto", "")
    for idx, art in enumerate(data.get("articulos", [])):
        mo = art.get("mejor_opcion", {}) or {}
        nombre = art.get("nombre", f"Artículo {idx+1}")

        # ── Imagen principal (descarga + verificación por visión) ─────────────
        img_local = None
        if mo.get("imagen_url"):
            cand = descargar_imagen(mo["imagen_url"], carpeta_tmp)
            if cand and imagen_corresponde_al_producto(cand, nombre):
                img_local = cand
        if not img_local:
            img_local = _buscar_imagen_fallback(nombre, carpeta_tmp)
        if img_local:
            imgs_prof[idx] = img_local

        # ── Anexos (imágenes adicionales del mismo producto) ──────────────────
        anexos = []
        for aurl in (mo.get("imagenes_anexos") or [])[:3]:
            a = descargar_imagen(aurl, carpeta_tmp)
            if a:
                anexos.append(a)

        gar_ficha = (f"Garantía de {gt.lower()} por defecto de fábrica"
                     if gt else "Garantía de 12 meses por defecto de fábrica")
        art_ficha.append({
            "nombre": nombre,
            "imagen": img_local,
            "descripcion": art.get("descripcion", ""),
            "caracteristicas": art.get("caracteristicas", []),
            "especificaciones_tecnicas": art.get("especificaciones_tecnicas", []),
            "especificaciones_electricas": art.get("especificaciones_electricas", []),
            "incluye": art.get("incluye", []),
            "garantia_texto": gar_ficha,
            "anexos": anexos,
        })
        art_prof.append({
            "nombre": nombre,
            "cantidad": int(art.get("cantidad", 0) or 0),
            "url": mo.get("url", ""),
            "alternativas": art.get("alternativas", []),
            "valor_unitario": float(mo.get("valor_unitario", 0.0) or 0.0),
            "proveedor": (mo.get("proveedor") or _dominio(mo.get("url", ""))).lower(),
            "es_extranjero": bool(mo.get("es_extranjero", False)),
            "costo_envio": float(mo.get("costo_envio", 0.0) or 0.0),
            "costo_aduana": float(mo.get("costo_aduana", 0.0) or 0.0),
            "requiere_instalacion": bool(art.get("requiere_instalacion", False)),
            "costo_instalacion": float(art.get("costo_instalacion", 0.0) or 0.0),
        })
    return art_ficha, art_prof, imgs_prof


# ═══════════════════════════════════════════════════════════════════════════════
#  ORQUESTADOR  ·  procesamiento de un código de necesidad
# ═══════════════════════════════════════════════════════════════════════════════

def _ruc13(codigo: str) -> str:
    """Caracteres 5 a 17 del código de necesidad (RUC de la entidad)."""
    return (codigo or "")[4:17]


def procesar_necesidad(fila: dict, carpeta_base: str) -> str:
    """
    Procesa una ínfima completa. Devuelve un estado:
    'ok' | 'limite' | 'sin_docs' | 'error'.
    """
    codigo = fila.get("codigo_necesidad", "")
    id_infima = fila.get("id_infima", "")
    sub(f"{codigo}   (id_infima={id_infima})")
    tmp = os.path.join(carpeta_base, re.sub(r"[^\w\-]", "_", str(codigo)))
    os.makedirs(tmp, exist_ok=True)

    # 1) Descargar documentos del bucket
    paso("Descargando documentos de contratación del bucket…")
    rutas = descargar_documentos(codigo, tmp)
    if not rutas:
        warn("No se encontraron documentos en el bucket para este código. Se omite.")
        return "sin_docs"
    dato("Documentos", f"{len(rutas)} archivo(s)")

    # 2) Extraer texto (+ PDFs para la IA)
    paso("Extrayendo contenido de los documentos…")
    texto_docs, bloques_pdf = extraer_documentos(rutas, tmp)

    # 3) Análisis con Claude Sonnet 4.6
    paso("Analizando con Claude Sonnet 4.6 (búsqueda web de productos)…")
    data = analizar_necesidad(fila, texto_docs, bloques_pdf)
    if not data:
        err("La IA no devolvió un JSON válido. Se omite este código.")
        return "error"

    # 4) Regla de límite de artículos (> 10)
    n = int(data.get("num_articulos", len(data.get("articulos", []))) or 0)
    if data.get("supera_limite") or n > LIMITE_ARTICULOS:
        warn(f"Se detectaron {n} artículos distintos (> {LIMITE_ARTICULOS}). "
             "Se genera ficha de aviso y NO se crea proforma.")
        ruta_ficha = os.path.join(tmp, f"{codigo}_Ficha_técnica_{id_infima}.docx")
        construir_ficha_limite_excedido(ruta_ficha, n)
        subir_a_bucket(ruta_ficha, CARPETA_FICHAS,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        marcar_finalizada(id_infima)
        ok("Ficha de aviso subida y ínfima marcada como 'finalizada'.")
        return "limite"

    dato("Artículos", n)
    dato("Tiempo de entrega", data.get("tiempo_entrega_texto", "?"))
    dato("Garantía", data.get("garantia_texto", "?"))

    # 5) Preparar artículos (descarga/verificación de imágenes)
    paso("Descargando y verificando imágenes de producto…")
    art_ficha, art_prof, imgs_prof = _preparar_articulos(data, tmp)

    # 6) Ficha técnica (.docx)
    paso("Generando ficha técnica (.docx)…")
    ruta_ficha = os.path.join(tmp, f"{codigo}_Ficha_técnica_{id_infima}.docx")
    construir_ficha_tecnica(art_ficha, ruta_ficha)
    subir_a_bucket(ruta_ficha, CARPETA_FICHAS,
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    ok(f"Ficha técnica subida a '{CARPETA_FICHAS}/'.")
    RESUMEN.sumar("fichas")

    # 7) Proforma (.xlsm)
    paso("Generando proforma (.xlsm)…")
    art_prof = _distribuir_costos(art_prof, en_guayaquil=bool(data.get("en_guayaquil", False)))
    datos_prof = {
        "entidad_contratante":     fila.get("entidad_contratante", ""),
        "entidad_contratante_url": fila.get("entidad_contratante_url", ""),
        "ruc13":                   _ruc13(codigo),
        "direccion_entrega":       fila.get("direccion_entrega", ""),
        "contacto":                fila.get("contacto", ""),
        "codigo_necesidad":        codigo,
        "id_infima":               id_infima,
        "cpc":                     fila.get("CPC"),
        "tiempo_entrega_texto":    data.get("tiempo_entrega_texto", ""),
        "garantia_texto":          data.get("garantia_texto", ""),
    }
    ruta_prof = os.path.join(tmp, f"{codigo}_Proforma_{id_infima}.xlsm")
    construir_proforma(datos_prof, art_prof, imgs_prof, ruta_prof)
    subir_a_bucket(ruta_prof, CARPETA_PROFORMAS, "application/vnd.ms-excel.sheet.macroEnabled.12")
    ok(f"Proforma subida a '{CARPETA_PROFORMAS}/'.")
    RESUMEN.sumar("proformas")

    # 8) Marcar como finalizada
    marcar_finalizada(id_infima)
    ok("Ínfima marcada como 'finalizada' en la base de datos.")
    return "ok"


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main() -> int:
    global _PROJECT_ID
    banner("5_PREFORMS_GENERATOR · MOTOR CLAUDE SONNET 4.6 (VERTEX AI)")
    print(f"  Inicio: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ·  "
          f"Modelo: {MODELO_IA}  ·  Región: {VERTEX_REGION}")

    # ── Etapa 0 · Credenciales ────────────────────────────────────────────────
    cron = Cronometro()
    banner("ETAPA 0 · Credenciales y clientes")
    try:
        _, _PROJECT_ID = _resolver_credenciales()
        ok(f"Credenciales resueltas (project_id={_PROJECT_ID or 'ADC'}).")
    except Exception as e:
        err(f"No se pudieron resolver las credenciales: {e}")
        return 1
    RESUMEN.etapa("Etapa 0 · Credenciales", cron.leer())

    # ── Etapa 1 · Leer data_table_1 ───────────────────────────────────────────
    cron.reset()
    banner("ETAPA 1 · Lectura de data_table_1 (infimas 'en generacion')")
    try:
        data_table_1 = obtener_data_table_1()
    except Exception as e:
        err(f"Error consultando la base de datos: {e}")
        traceback.print_exc()
        return 1
    ok(f"{len(data_table_1)} ínfima(s) en etapa 'en generacion'.")
    RESUMEN.etapa("Etapa 1 · Lectura de base de datos", cron.leer(),
                  f"{len(data_table_1)} filas")
    if not data_table_1:
        warn("No hay ínfimas por procesar. Fin.")
        RESUMEN.imprimir_final()
        return 0

    # ── Etapa 2 · Procesamiento por código de necesidad ───────────────────────
    cron.reset()
    banner("ETAPA 2 · Generación de fichas y proformas")
    base_tmp = tempfile.mkdtemp(prefix="nexus_preforms_")
    for i, fila in enumerate(data_table_1, start=1):
        print(f"\n  {C.BOLD}[{i}/{len(data_table_1)}]{C.RESET}")
        t_fila = Cronometro()
        try:
            estado = procesar_necesidad(fila, base_tmp)
            RESUMEN.sumar("procesadas")
            if estado == "limite":
                RESUMEN.sumar("limite_excedido")
            elif estado in ("error", "sin_docs"):
                RESUMEN.sumar("errores")
        except Exception as e:
            err(f"Error inesperado procesando {fila.get('codigo_necesidad','?')}: {e}")
            traceback.print_exc()
            RESUMEN.sumar("errores")
        finally:
            print(f"  {C.DIM}└─ tiempo de esta ínfima: {fmt_dur(t_fila.leer())}{C.RESET}")

    RESUMEN.etapa("Etapa 2 · Fichas y proformas", cron.leer(),
                  f"{len(data_table_1)} ínfimas")

    # Limpieza de temporales
    try:
        shutil.rmtree(base_tmp, ignore_errors=True)
    except Exception:
        pass

    # ── Resumen final ─────────────────────────────────────────────────────────
    RESUMEN.imprimir_final()
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\nInterrumpido por el usuario.")
        sys.exit(130)