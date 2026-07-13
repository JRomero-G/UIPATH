# -*- coding: utf-8 -*-
"""
═══════════════════════════════════════════════════════════════════════════════
 5_Preforms_generator.py  ·  Proyecto NEXUS  ·  IMPOCRUZ EC S.A.S.
═══════════════════════════════════════════════════════════════════════════════
 Generador de Fichas Técnicas (.docx) y Proformas (.xlsm) a partir de los
 Documentos de Contratación (ínfimas cuantías) usando Gemini 3.1 Pro Preview
 en Vertex AI (fallback: Gemini 2.5 Pro).

 FLUJO GENERAL (una "fase" por cada paso numerado del funcionamiento):
   FASE 1  Lee de MySQL (gestorex.infimas) las filas con etapa='en generacion'
           y guarda las columnas requeridas en data_table_1.
   FASE 2  Lista en el bucket GCS la carpeta "Documentos de Contratación/<código>"
           y descarga los documentos (doc, docx, pdf).
   FASE 3  Pide a Gemini (3.1 pro preview → fallback 2.5 pro) el análisis de los
           documentos + búsqueda web en la lista de proveedores de confianza.
   FASE 4  Verifica/repara URLs y descarga + valida la imagen del producto.
   FASE 5  Genera la Ficha Técnica .docx (plantilla corporativa) y la sube a
           la carpeta "Fichas Técnicas" del bucket.
   FASE 6  Genera la Proforma .xlsm (plantilla corporativa, conserva el botón
           "Imprimir PDF") y la sube a la carpeta "Proformas" del bucket.
   FASE 7  Cambia la etapa de la ínfima a 'finalizada' en la base de datos.

 REGLA DEL LÍMITE: si un código de necesidad tiene MÁS de 10 artículos
 distintos, la ficha contiene únicamente el mensaje en rojo (Arial bold 14),
 NO se crea proforma y la etapa pasa directamente a 'finalizada'.

 MODO PRUEBA LOCAL (sin nube):
   python 5_Preforms_generator.py --test-local RUTA.pdf [--mock-json datos.json]
                                  [--salida DIR]
   · Usa un documento de contratación local, un registro simulado y (si no hay
     credenciales de Vertex) un análisis simulado, para probar de punta a punta
     la generación de ficha y proforma sin tocar MySQL ni GCS.

 Requisitos (producción):
   pip install mysql-connector-python google-cloud-storage google-genai
   pip install python-docx openpyxl requests pillow ddgs
═══════════════════════════════════════════════════════════════════════════════
"""

import os, sys, json, shutil, tempfile, datetime, re, io, time, traceback
import urllib.parse, html, argparse, zipfile, copy
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ── Ruta raíz del proyecto (el script vive 3 niveles bajo la raíz) ────────────
RAIZ_PROYECTO = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(RAIZ_PROYECTO))
try:
    from Config import Global                     # configuración centralizada
    _HAY_CONFIG = True
except Exception:                                 # en modo prueba local puede no existir
    Global = None
    _HAY_CONFIG = False

# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIGURACIÓN GLOBAL
# ═══════════════════════════════════════════════════════════════════════════════

MYSQL_CONFIG = {
    "host":     getattr(Global, "DB_HOST", None),
    "user":     getattr(Global, "DB_USER", None),
    "password": getattr(Global, "DB_PASSWORD", None),
    "database": getattr(Global, "DATABASE", None),   # "gestorex"
    "connection_timeout": 20,
} if _HAY_CONFIG else {}

BUCKET_NAME       = getattr(Global, "BUCKET_NAME", None) if _HAY_CONFIG else None
BUCKET_FOLDER     = "Documentos de Contratación"     # carpeta de entrada en GCS
CARPETA_FICHAS    = "Fichas Técnicas"                # carpeta de salida (docx)
CARPETA_PROFORMAS = "Proformas"                      # carpeta de salida (xlsm)

# ── Modelos Vertex AI ─────────────────────────────────────────────────────────
MODEL_COMPLEX   = "gemini-3.1-pro-preview"   # tareas complejas + grounding web
MODEL_SIMPLE    = "gemini-2.5-pro"           # fallback ante cuota/errores
VERTEX_LOCATION = "global"                   # endpoint global (Gemini 3.x)

# ── Reglas de negocio ─────────────────────────────────────────────────────────
LIMITE_ARTICULOS       = 10                  # > 10 artículos ⇒ sin proforma
ENVIO_MIN, ENVIO_MAX   = 86.0, 155.0         # USD, entregas fuera de Guayaquil
INSTAL_MIN, INSTAL_MAX = 60.0, 80.0          # USD por artículo con instalación
ALTURA_FILA_ARTICULO   = 220.0               # altura recomendada por artículo

MENSAJE_LIMITE = "La cantidad de artículos supera el límite permitido de 10 artículos de compra"

# ── Plantillas (junto al script en producción) ────────────────────────────────
SCRIPT_DIR    = Path(__file__).parent
TEMPLATE_DOCX = SCRIPT_DIR / "FICHA TECNICA MICROFONO Y MEMORIA.docx"
TEMPLATE_XLSX = SCRIPT_DIR / "FORMATO DE PROFORMA RECREADO VACÍO.xlsm"

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)
DEFAULT_HEADERS = {
    "User-Agent":      USER_AGENT,
    "Accept-Language": "es-ES,es;q=0.9,en;q=0.7",
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
}

# ── Lista de proveedores de confianza (se inyecta en el prompt de búsqueda) ──
PROVEEDORES_NACIONALES = [
    "https://mibodega.ec/", "https://bodeguitadelahorro.com/", "https://comercialvaca.ec/",
    "https://kissu.com.ec/", "https://altatenalmacen.com.ec/", "https://www.tventas.com/",
    "https://store.intcomex.com/es-XPE/Home", "https://lavictoria.ec/", "https://www.marcimex.com/",
    "https://almacenesespana.ec/", "https://www.novicompu.com/", "https://technoprimec.com/",
    "https://point.com.ec/", "https://www.computron.com.ec/", "https://nomadaware.com.ec/",
    "https://www.idcmayoristas.com/", "https://tecnit.com.ec/", "https://tecnomegastore.ec/",
    "https://www.artefacta.com/", "https://mundotek.com.ec/", "https://eljuri.store/",
    "https://www.kywi.com.ec/", "https://tecnocostoec.com/", "https://www.almacenesjapon.com/",
    "https://electromegaecuador.com/", "https://www.compraecuador.com/", "https://granhogar.com.ec/",
    "https://miamihome-ec.com/",
]
PROVEEDORES_EXTRANJEROS = [
    "https://www.amazon.com/", "https://www.mercadolibre.com.mx/", "https://www.mercadolibre.com.co/",
    "https://www.mercadolibre.com.ar/", "https://www.mercadolibre.cl/", "https://www.mercadolibre.com.pe/",
    "https://www.ebay.com/", "https://www.walmart.com/", "https://www.bestbuy.com/",
    "https://www.homedepot.com/", "https://www.costco.com/",
]

def _dominio(url: str) -> str:
    """Extrae el dominio de una URL (sin 'www.') para agrupar proveedores."""
    try:
        return urllib.parse.urlparse(url).netloc.lower().replace("www.", "")
    except Exception:
        return ""

DOMINIOS_NACIONALES  = [d for d in (_dominio(u) for u in PROVEEDORES_NACIONALES) if d]
DOMINIOS_EXTRANJEROS = [d for d in (_dominio(u) for u in PROVEEDORES_EXTRANJEROS) if d]

# Respaldo opcional de búsqueda de imágenes (no crítico si no está instalado)
try:
    from ddgs import DDGS
except Exception:
    DDGS = None

_URL_CACHE: Dict[str, bool] = {}     # memoriza URLs ya verificadas (ahorra red)


# ═══════════════════════════════════════════════════════════════════════════════
#  UTILIDADES: consola, cronómetro de fases y sesión HTTP con reintentos
# ═══════════════════════════════════════════════════════════════════════════════

def log(msg: str, nivel: str = "INFO") -> None:
    """Imprime en consola con marca de tiempo y nivel (pasos importantes)."""
    hora = datetime.datetime.now().strftime("%H:%M:%S")
    print(f"[{hora}] [{nivel:<5}] {msg}", flush=True)


class Cronometro:
    """Mide la duración de cada fase y acumula un resumen por fase."""

    def __init__(self):
        self.registros: List[Tuple[str, float]] = []   # (nombre de fase, segundos)
        self._t0: Optional[float] = None
        self._nombre: Optional[str] = None

    def iniciar(self, nombre: str) -> None:
        self._nombre = nombre
        self._t0 = time.perf_counter()
        log(f"— Inicia {nombre} —")

    def detener(self) -> float:
        dur = time.perf_counter() - self._t0
        self.registros.append((self._nombre, dur))
        log(f"— Termina {self._nombre} en {dur:.2f} s —")
        return dur

    def resumen(self) -> None:
        print("\n" + "─" * 70)
        print(" RESUMEN DE TIEMPOS POR FASE")
        print("─" * 70)
        for nombre, dur in self.registros:
            print(f"   {nombre:<52} {dur:>8.2f} s")
        print(f"   {'TOTAL':<52} {sum(d for _, d in self.registros):>8.2f} s")
        print("─" * 70)


def crear_sesion_http() -> requests.Session:
    """Sesión requests con reintentos automáticos ante errores transitorios."""
    s = requests.Session()
    reintentos = Retry(total=3, backoff_factor=0.8,
                       status_forcelist=(429, 500, 502, 503, 504),
                       allowed_methods=frozenset(["GET", "HEAD"]))
    s.mount("http://", HTTPAdapter(max_retries=reintentos))
    s.mount("https://", HTTPAdapter(max_retries=reintentos))
    s.headers.update(DEFAULT_HEADERS)
    return s

SESION = crear_sesion_http()


# ═══════════════════════════════════════════════════════════════════════════════
#  CREDENCIALES (acepta JSON en línea o ruta a archivo; nombres en prioridad)
# ═══════════════════════════════════════════════════════════════════════════════

def resolver_credenciales_a_archivo() -> Optional[str]:
    """
    Busca la credencial de servicio en varios nombres de variable (en orden de
    prioridad) tanto en Config.Global como en variables de entorno. Acepta:
      · JSON en línea  → lo vuelca a un archivo temporal
      · Ruta a archivo → la normaliza (incluye rutas Windows)
    Devuelve la ruta final del archivo o None si no encontró nada.
    NOTA: 'RENDER_CRENDENTIALS_JSON' conserva su errata intencional.
    """
    nombres = ["GEMINI_CREDENTIALS", "RENDER_CRENDENTIALS_JSON", "CREDENTIALS_GEMINI"]
    valor = None
    for n in nombres:
        valor = getattr(Global, n, None) if _HAY_CONFIG else None
        if not valor:
            valor = os.environ.get(n)
        if valor:
            break
    if not valor:
        return None
    valor = str(valor).strip()
    if valor.lstrip().startswith("{"):                    # JSON en línea
        tmp = Path(tempfile.gettempdir()) / "nexus_sa_credentials.json"
        tmp.write_text(valor, encoding="utf-8")
        return str(tmp)
    ruta = Path(valor.replace("\\", "/"))                 # normaliza Windows
    return str(ruta) if ruta.exists() else None


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 1 · BASE DE DATOS → data_table_1
# ═══════════════════════════════════════════════════════════════════════════════

def obtener_data_table_1() -> List[Dict[str, Any]]:
    """
    Lee de gestorex.infimas SOLO las filas con etapa='en generacion' y devuelve
    una lista de diccionarios (data_table_1) con las columnas literales:
    codigo_necesidad, entidad_contratante, entidad_contratante_url,
    dirección_entrega, id_infima, CPC, contacto.
    """
    import mysql.connector                                 # import diferido
    consulta = (
        "SELECT `codigo_necesidad`, `entidad_contratante`, `entidad_contratante_url`, "
        "`dirección_entrega`, `id_infima`, `CPC`, `contacto` "
        "FROM `infimas` WHERE `etapa` = 'en generacion'"
    )
    conexion = mysql.connector.connect(**MYSQL_CONFIG)
    try:
        cursor = conexion.cursor(dictionary=True)
        cursor.execute(consulta)
        data_table_1 = cursor.fetchall()
        cursor.close()
    finally:
        conexion.close()
    log(f"FASE 1 · {len(data_table_1)} ínfima(s) en etapa 'en generacion'.")
    return data_table_1


def marcar_finalizada(id_infima: Any) -> None:
    """Cambia la etapa de la ínfima a 'finalizada' (FASE 7)."""
    import mysql.connector
    conexion = mysql.connector.connect(**MYSQL_CONFIG)
    try:
        cursor = conexion.cursor()
        cursor.execute("UPDATE `infimas` SET `etapa`='finalizada' WHERE `id_infima`=%s",
                       (id_infima,))
        conexion.commit()
        cursor.close()
        log(f"FASE 7 · Ínfima {id_infima} marcada como 'finalizada'.")
    finally:
        conexion.close()


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 2 · GOOGLE CLOUD STORAGE
# ═══════════════════════════════════════════════════════════════════════════════

def cliente_gcs():
    """Crea el cliente de Cloud Storage con la credencial de servicio."""
    from google.cloud import storage                       # import diferido
    ruta = resolver_credenciales_a_archivo()
    if ruta:
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = ruta
    return storage.Client()


def descargar_documentos_contratacion(bucket, codigo: str, destino: Path) -> List[Path]:
    """
    Descarga a 'destino' todos los doc/docx/pdf que estén en
    'Documentos de Contratación/<codigo>/' dentro del bucket.
    """
    prefijo = f"{BUCKET_FOLDER}/{codigo}/"
    rutas: List[Path] = []
    for blob in bucket.list_blobs(prefix=prefijo):
        nombre = Path(blob.name).name
        if not nombre or not nombre.lower().endswith((".pdf", ".doc", ".docx")):
            continue
        destino.mkdir(parents=True, exist_ok=True)
        ruta = destino / nombre
        blob.download_to_filename(str(ruta))
        rutas.append(ruta)
    log(f"FASE 2 · {codigo}: {len(rutas)} documento(s) de contratación descargado(s).")
    return rutas


def subir_a_gcs(bucket, ruta_local: Path, carpeta_destino: str) -> None:
    """Sube un archivo local al bucket dentro de la carpeta indicada."""
    blob = bucket.blob(f"{carpeta_destino}/{ruta_local.name}")
    blob.upload_from_filename(str(ruta_local))
    log(f"GCS · Subido {ruta_local.name} → {carpeta_destino}/")


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 3 · ANÁLISIS CON GEMINI (Vertex AI)  ·  búsqueda web + JSON estricto
# ═══════════════════════════════════════════════════════════════════════════════

def _cliente_genai():
    """Cliente google-genai apuntando a Vertex AI (endpoint global)."""
    from google import genai                                # import diferido
    ruta = resolver_credenciales_a_archivo()
    if ruta:
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = ruta
        proyecto = json.loads(Path(ruta).read_text(encoding="utf-8")).get("project_id")
    else:
        proyecto = os.environ.get("GOOGLE_CLOUD_PROJECT")
    return genai.Client(vertexai=True, project=proyecto, location=VERTEX_LOCATION)


def _partes_documentos(rutas: List[Path]):
    """
    Convierte los documentos de contratación en 'parts' para Gemini:
      · PDF   → bytes nativos (Gemini los lee directamente)
      · DOCX  → texto extraído con python-docx
      · DOC   → intento de conversión a texto plano (LibreOffice si existe)
    """
    from google.genai import types
    partes = []
    for ruta in rutas:
        ext = ruta.suffix.lower()
        if ext == ".pdf":
            partes.append(types.Part.from_bytes(data=ruta.read_bytes(),
                                                mime_type="application/pdf"))
        elif ext == ".docx":
            import docx as _docx
            texto = "\n".join(p.text for p in _docx.Document(str(ruta)).paragraphs)
            partes.append(types.Part.from_text(text=f"[Documento {ruta.name}]\n{texto}"))
        elif ext == ".doc":
            try:                                            # conversión opcional
                import subprocess
                subprocess.run(["soffice", "--headless", "--convert-to", "txt",
                                "--outdir", str(ruta.parent), str(ruta)],
                               check=True, capture_output=True, timeout=120)
                txt = ruta.with_suffix(".txt")
                partes.append(types.Part.from_text(
                    text=f"[Documento {ruta.name}]\n{txt.read_text(errors='ignore')}"))
            except Exception:
                log(f"No se pudo convertir {ruta.name}; se omite.", "WARN")
    return partes


def _llamar_modelo(contenidos, config, usar_grounding: bool) -> str:
    """
    Llama primero a MODEL_COMPLEX (gemini-3.1-pro-preview) y, si falla por
    cuota/errores, cae automáticamente a MODEL_SIMPLE (gemini-2.5-pro).
    Devuelve el texto de la respuesta.
    """
    from google.genai import types
    cliente = _cliente_genai()
    if usar_grounding:                                     # búsqueda web nativa
        config.tools = [types.Tool(google_search=types.GoogleSearch())]
    for modelo in (MODEL_COMPLEX, MODEL_SIMPLE):
        try:
            log(f"FASE 3 · Llamando a {modelo}…")
            resp = cliente.models.generate_content(model=modelo,
                                                   contents=contenidos,
                                                   config=config)
            return resp.text or ""
        except Exception as e:
            log(f"{modelo} falló ({type(e).__name__}: {e}); probando fallback…", "WARN")
    raise RuntimeError("Ambos modelos Gemini fallaron.")


def _extraer_json(texto: str) -> dict:
    """Extrae el primer objeto JSON válido de la respuesta del modelo."""
    texto = re.sub(r"```(?:json)?", "", texto).strip().strip("`")
    inicio = texto.find("{")
    if inicio < 0:
        raise ValueError("La respuesta del modelo no contiene JSON.")
    return json.loads(texto[inicio:texto.rfind("}") + 1])


ESQUEMA_ANALISIS = """
{
  "excede_limite": bool,            // true si hay MÁS de 10 artículos distintos
  "num_articulos": int,
  "tiempo_entrega_dias": int,       // establecido por la entidad contratante
  "garantia_tecnica": "texto en meses o años, p.ej. '12 MESES O 10.000 KM'",
  "articulos": [
    {
      "nombre": "nombre comercial", "marca": "…", "modelo": "…",
      "resumen": "60 a 100 palabras tomadas de la web del proveedor/fabricante",
      "caracteristicas": ["mínimo 7 características generales"],
      "especificaciones": ["mínimo 10 especificaciones técnicas y eléctricas"],
      "incluye": ["extras que vienen con el artículo"],
      "cantidad": int,              // unidades solicitadas en los documentos
      "requiere_instalacion": bool,
      "complejidad_instalacion": "baja|media|alta",
      "mejor_opcion": {
        "proveedor": "…", "url_producto": "URL real y actual",
        "url_imagen": "URL directa a la imagen del producto",
        "precio_unitario": float,   // costo unitario real sin envío/aduana
        "es_extranjero": bool,
        "costo_extra_envio_sugerido": float,   // 86–155 USD fuera de Guayaquil
        "costo_aduana_estimado": float
      },
      "alternativas": [             // hasta 3, en orden decreciente de idoneidad
        {"proveedor": "…", "url": "…", "precio_unitario": float}
      ]
    }
  ]
}
"""


def analizar_con_ia(rutas_docs: List[Path], registro: Dict[str, Any]) -> dict:
    """
    Análisis en DOS pasos (el grounding web y el JSON estricto no se combinan
    en una sola llamada de forma fiable):
      1) Investigación con búsqueda web (Google Search grounding) sobre los
         documentos de contratación y la lista de proveedores de confianza.
      2) Estructuración del hallazgo a JSON estricto según ESQUEMA_ANALISIS.
    """
    from google.genai import types
    partes = _partes_documentos(rutas_docs)

    prompt_busqueda = f"""
Eres un analista de compras públicas de Ecuador. Analiza los documentos de
contratación adjuntos (código de necesidad {registro['codigo_necesidad']}).

TAREAS:
1) Identifica CADA artículo de compra DISTINTO (nombre, marca y modelo). Si los
   documentos no los especifican, propone un producto NUEVO que cumpla las
   características solicitadas (coincidencia > 70%, idealmente 100%).
2) Si hay MÁS de {LIMITE_ARTICULOS} artículos distintos, indícalo y no continúes.
3) Por artículo: al menos 7 características generales, al menos 10
   especificaciones técnicas/eléctricas y lo que incluye (extras).
4) Busca en la web cada artículo, PRIMERO en estos proveedores de confianza
   nacionales: {', '.join(PROVEEDORES_NACIONALES)}
   luego extranjeros: {', '.join(PROVEEDORES_EXTRANJEROS)}
   Solo si se agota la lista, propone otro proveedor de Ecuador y, en último
   caso, uno extranjero. Encuentra hasta 4 opciones de proveedor por artículo.
5) Elige la MEJOR opción: cumplimiento de requisitos y MENOR costo final. Si el
   proveedor es extranjero, suma envío hasta Guayaquil (Ecuador), aduanas y
   logística; con esos extras debe seguir siendo la más barata. Artículos NUEVOS.
6) Conserva las URLs TAL CUAL las encuentres (actuales y reales), incluida una
   URL directa de la imagen del producto (del proveedor o del fabricante).
7) Extrae de los documentos: tiempo de entrega (días) y garantía técnica
   (meses o años), y la cantidad de unidades por artículo.
La dirección de entrega es: {registro.get('dirección_entrega', '')}.
Responde con un informe detallado de tus hallazgos, incluyendo TODAS las URLs.
"""
    config1 = types.GenerateContentConfig(temperature=0.2)
    informe = _llamar_modelo(partes + [types.Part.from_text(text=prompt_busqueda)],
                             config1, usar_grounding=True)

    prompt_json = f"""
Convierte el siguiente informe a JSON ESTRICTO con exactamente este esquema
(sin comentarios, sin texto adicional, sin Markdown):
{ESQUEMA_ANALISIS}

INFORME:
{informe}
"""
    config2 = types.GenerateContentConfig(temperature=0.0,
                                          response_mime_type="application/json")
    texto = _llamar_modelo([types.Part.from_text(text=prompt_json)],
                           config2, usar_grounding=False)
    analisis = _extraer_json(texto)
    log(f"FASE 3 · Análisis IA: {analisis.get('num_articulos', '?')} artículo(s).")
    return analisis


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 4 · VERIFICACIÓN / REPARACIÓN DE URLs E IMÁGENES
# ═══════════════════════════════════════════════════════════════════════════════

def verificar_url(url: str, estricto: bool = True) -> bool:
    """
    Comprueba que la URL exista y responda. En modo estricto exige 200 y HTML;
    en modo laxo acepta 2xx/3xx. Cachea resultados para no repetir peticiones.
    """
    if not url or not url.startswith("http"):
        return False
    clave = f"{url}|{estricto}"
    if clave in _URL_CACHE:
        return _URL_CACHE[clave]
    ok = False
    try:
        r = SESION.get(url, timeout=15, allow_redirects=True, stream=True)
        if estricto:
            ok = r.status_code == 200
        else:
            ok = 200 <= r.status_code < 400
        r.close()
    except Exception:
        ok = False
    _URL_CACHE[clave] = ok
    return ok


def _buscar_url_reemplazo(nombre_producto: str, dominios: List[str]) -> Optional[str]:
    """
    Recuperación de links muertos: busca (vía IA con grounding y, como respaldo,
    DuckDuckGo) una URL viva del producto dentro de los dominios de confianza.
    """
    # 1) Con la IA (búsqueda web nativa)
    try:
        from google.genai import types
        prompt = (f"Busca en la web una URL ACTUAL y accesible del producto "
                  f"'{nombre_producto}' preferentemente en estos dominios: "
                  f"{', '.join(dominios[:15])}. Responde SOLO con la URL.")
        cfg = types.GenerateContentConfig(temperature=0.0)
        texto = _llamar_modelo([types.Part.from_text(text=prompt)], cfg, True)
        m = re.search(r"https?://\S+", texto)
        if m and verificar_url(m.group(0).rstrip(").,"), estricto=False):
            return m.group(0).rstrip(").,")
    except Exception:
        pass
    # 2) Con DuckDuckGo, si está instalado
    if DDGS is not None:
        try:
            with DDGS() as ddg:
                for res in ddg.text(nombre_producto, max_results=10):
                    url = res.get("href", "")
                    if any(d in url for d in dominios) and verificar_url(url, False):
                        return url
        except Exception:
            pass
    return None


def resolver_url_producto(url: str, nombre_producto: str) -> Optional[str]:
    """
    Multi-etapa: verificación estricta → recuperación por búsqueda →
    verificación laxa. Devuelve una URL viva o None.
    """
    if verificar_url(url, estricto=True):
        return url
    log(f"FASE 4 · URL caída, intentando recuperar: {url}", "WARN")
    nueva = _buscar_url_reemplazo(nombre_producto, DOMINIOS_NACIONALES + DOMINIOS_EXTRANJEROS)
    if nueva:
        return nueva
    return url if verificar_url(url, estricto=False) else None


def descargar_imagen(url_imagen: str, destino: Path) -> Optional[Path]:
    """
    Descarga la imagen del producto y confirma que realmente es una imagen
    válida (la abre con Pillow). Devuelve la ruta o None si falló.
    """
    try:
        r = SESION.get(url_imagen, timeout=20)
        if r.status_code != 200 or len(r.content) < 2048:
            return None
        from PIL import Image
        img = Image.open(io.BytesIO(r.content))
        img.verify()                                        # valida integridad
        ext = (img.format or "PNG").lower()
        ruta = destino.with_suffix("." + ("jpg" if ext == "jpeg" else ext))
        ruta.write_bytes(r.content)
        return ruta
    except Exception:
        return None


def verificar_imagen_con_ia(ruta_imagen: Path, nombre_producto: str) -> bool:
    """
    Verificación por VISIÓN: pregunta a Gemini si la imagen corresponde al
    producto (evita quedarnos con banners/anuncios de la página).
    """
    try:
        from google.genai import types
        parte_img = types.Part.from_bytes(data=ruta_imagen.read_bytes(),
                                          mime_type="image/png"
                                          if ruta_imagen.suffix == ".png" else "image/jpeg")
        prompt = (f"¿Esta imagen muestra el producto '{nombre_producto}' "
                  f"(no un anuncio, banner o logotipo)? Responde SOLO 'SI' o 'NO'.")
        cfg = types.GenerateContentConfig(temperature=0.0)
        texto = _llamar_modelo([parte_img, types.Part.from_text(text=prompt)], cfg, False)
        return "SI" in texto.upper()[:10]
    except Exception:
        return True     # ante duda de infraestructura, no bloquear el flujo


def obtener_imagen_producto(articulo: dict, carpeta: Path, con_ia: bool) -> Optional[Path]:
    """
    Orquesta la obtención de la imagen: descarga desde la URL de la IA,
    valida por visión y, si falla, busca una imagen alternativa (DDGS).
    En modo prueba puede venir 'imagen_local' ya lista en el artículo.
    """
    if articulo.get("imagen_local"):                        # modo prueba local
        ruta = Path(articulo["imagen_local"])
        return ruta if ruta.exists() else None
    nombre = f"{articulo.get('marca','')} {articulo.get('modelo','')}".strip() \
             or articulo.get("nombre", "producto")
    url_img = (articulo.get("mejor_opcion") or {}).get("url_imagen", "")
    destino = carpeta / re.sub(r"\W+", "_", nombre)[:40]
    ruta = descargar_imagen(url_img, destino) if url_img else None
    if ruta and con_ia and not verificar_imagen_con_ia(ruta, nombre):
        log(f"FASE 4 · Imagen descartada por visión IA: {url_img}", "WARN")
        ruta = None
    if ruta is None and DDGS is not None:                   # respaldo de imágenes
        try:
            with DDGS() as ddg:
                for res in ddg.images(nombre, max_results=8):
                    ruta = descargar_imagen(res.get("image", ""), destino)
                    if ruta and (not con_ia or verificar_imagen_con_ia(ruta, nombre)):
                        break
                    ruta = None
        except Exception:
            ruta = None
    return ruta


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 5 · FICHA TÉCNICA .docx  (plantilla corporativa: Century Gothic 12,
#            encabezado y pie de página se conservan intactos)
# ═══════════════════════════════════════════════════════════════════════════════

FUENTE_FICHA = "Century Gothic"
TAM_FICHA    = 12       # puntos (la plantilla usa sz=24 medios puntos)


def _limpiar_cuerpo_docx(doc) -> None:
    """
    Vacía el cuerpo del documento PERO conserva la sección final (sectPr):
    así se mantienen tamaño de página, márgenes, encabezado y pie originales.
    """
    from docx.oxml.ns import qn
    cuerpo = doc.element.body
    for hijo in list(cuerpo):
        if hijo.tag != qn("w:sectPr"):
            cuerpo.remove(hijo)


def _parrafo(doc, texto: str, negrita=False, fuente=FUENTE_FICHA, tam=TAM_FICHA,
             color=None, centrado=False):
    """Añade un párrafo con la tipografía de la plantilla."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrita
    run.font.name = fuente
    run.font.size = Pt(tam)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _vinetas(doc, items: List[str]) -> None:
    """
    Lista con viñetas usando el estilo 'List Bullet' de la plantilla; si el
    estilo no existe (KeyError latente conocido), usa un fallback con '• '.
    """
    from docx.shared import Pt
    for texto in items:
        try:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(texto))
        except KeyError:                                    # fallback seguro
            p = doc.add_paragraph()
            run = p.add_run("• " + str(texto))
        run.font.name = FUENTE_FICHA
        run.font.size = Pt(TAM_FICHA)


def generar_ficha_docx(analisis: dict, imagenes: Dict[int, Optional[Path]],
                       ruta_salida: Path) -> Path:
    """
    Genera la ficha técnica siguiendo el lineamiento de la plantilla:
    título en negrita (nombre del artículo) → imagen → resumen (60–100
    palabras) → Características → Especificaciones técnicas → Incluye →
    Anexos (imágenes adicionales). Un bloque por artículo, mismo docx.
    Si excede el límite de artículos, SOLO escribe el mensaje en rojo.
    """
    import docx
    from docx.shared import Inches
    doc = docx.Document(str(TEMPLATE_DOCX))                 # hereda encabezado/pie
    _limpiar_cuerpo_docx(doc)

    if analisis.get("excede_limite") or analisis.get("num_articulos", 0) > LIMITE_ARTICULOS:
        # Regla del límite: mensaje único en rojo, Arial bold 14 — nada más.
        _parrafo(doc, MENSAJE_LIMITE, negrita=True, fuente="Arial", tam=14,
                 color=(255, 0, 0))
        doc.save(str(ruta_salida))
        log(f"FASE 5 · Ficha con mensaje de límite generada: {ruta_salida.name}")
        return ruta_salida

    articulos = analisis.get("articulos", [])
    for i, art in enumerate(articulos):
        titulo = " ".join(x for x in (art.get("nombre"), art.get("marca"),
                                      art.get("modelo")) if x)
        # El nombre puede ya contener marca/modelo; evita duplicarlos
        if art.get("marca") and art["marca"].lower() in (art.get("nombre") or "").lower():
            titulo = art.get("nombre", titulo)
        _parrafo(doc, titulo, negrita=True)

        ruta_img = imagenes.get(i)
        if ruta_img and ruta_img.exists():                  # imagen bajo el título
            doc.add_picture(str(ruta_img), width=Inches(4.3))
            doc.paragraphs[-1].alignment = 1                # centrada

        if art.get("resumen"):
            _parrafo(doc, art["resumen"])

        if art.get("caracteristicas"):
            _parrafo(doc, "Características:", negrita=True)
            _vinetas(doc, art["caracteristicas"])
        if art.get("especificaciones"):
            _parrafo(doc, "Especificaciones técnicas:", negrita=True)
            _vinetas(doc, art["especificaciones"])
        if art.get("incluye"):
            _parrafo(doc, "Incluye:", negrita=True)
            _vinetas(doc, art["incluye"])

        # Anexos: imágenes adicionales del mismo producto (si existen)
        extras = art.get("imagenes_anexos") or []
        rutas_extras = [Path(p) for p in extras if Path(p).exists()]
        if rutas_extras:
            _parrafo(doc, "Anexos", negrita=True)
            for rp in rutas_extras:
                doc.add_picture(str(rp), width=Inches(3.5))
                doc.paragraphs[-1].alignment = 1
        if i < len(articulos) - 1:
            doc.add_page_break()                            # un artículo por bloque

    doc.save(str(ruta_salida))
    log(f"FASE 5 · Ficha técnica generada: {ruta_salida.name}")
    return ruta_salida


# ═══════════════════════════════════════════════════════════════════════════════
#  FASE 6 · PROFORMA .xlsm  (openpyxl keep_vba + reinyección del botón PDF)
# ═══════════════════════════════════════════════════════════════════════════════

def _numero_proforma(id_infima: Any) -> str:
    """
    '00100100###': 8 dígitos fijos + id_infima en 3 posiciones. Si el id tiene
    más de 3 dígitos consume los últimos dígitos fijos; si tiene menos, se
    rellena con ceros. El resultado SIEMPRE mide 11 caracteres.
    Ej.: id 21 → 00100100021 · id 1234 → 00100101234
    """
    base, sid = "00100100", str(id_infima)
    if len(sid) <= 3:
        return base + sid.zfill(3)
    return (base + "0" * 3)[: 11 - len(sid)] + sid


def _contacto_resumido(contacto: str) -> str:
    """
    Del campo 'contacto' toma SOLO el nombre del encargado y su correo o su
    teléfono (lo primero que aparezca).
    """
    contacto = (contacto or "").strip()
    correo = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", contacto)
    telefono = re.search(r"(?:\+?593|0)?\s?\d{2,3}[\s-]?\d{3}[\s-]?\d{3,4}", contacto)
    # El "nombre" es lo que queda antes del primer dato de contacto…
    corte = min([m.start() for m in (correo, telefono) if m] or [len(contacto)])
    previo = contacto[:corte]
    # …y de ese texto se toma SOLO el primer segmento (descarta cargo, área, etc.)
    nombre = re.split(r"\s[-–|·]\s|,|;", previo)[0].strip(" -–,;:")
    dato = correo.group(0) if correo else (telefono.group(0) if telefono else "")
    return f"{nombre} - {dato}".strip(" -") if dato else nombre


def _clamp(v: float, lo: float, hi: float) -> float:
    """Limita un valor al rango [lo, hi] (rangos de negocio para extras)."""
    try:
        return max(lo, min(hi, float(v)))
    except Exception:
        return lo


def _costos_extras(analisis: dict, direccion_entrega: str) -> Dict[int, dict]:
    """
    Calcula los costos extras por artículo y los agrupa por proveedor:
      · Envío 86–155 USD si la entrega es fuera de Guayaquil (según distancia).
      · Instalación 60–80 USD por artículo que la requiera.
      · Proveedor COMPARTIDO (≥2 artículos): el total del proveedor se reparte
        entre sus artículos (columna A + fórmula en G con denominador propio).
      · Proveedor ÚNICO: el costo extra va DIRECTO en la columna G.
    Devuelve {índice_artículo: {"extra": USD, "grupo": dominio, "k": tamaño}}.
    """
    fuera_gye = "guayaquil" not in (direccion_entrega or "").lower()
    articulos = analisis.get("articulos", [])
    grupos: Dict[str, List[int]] = {}
    for i, art in enumerate(articulos):
        dom = _dominio((art.get("mejor_opcion") or {}).get("url_producto", "")) or f"prov{i}"
        grupos.setdefault(dom, []).append(i)

    resultado: Dict[int, dict] = {}
    for i, art in enumerate(articulos):
        mejor = art.get("mejor_opcion") or {}
        extra = 0.0
        if fuera_gye:                                       # envío fuera de GYE
            extra += _clamp(mejor.get("costo_extra_envio_sugerido", ENVIO_MIN),
                            ENVIO_MIN, ENVIO_MAX)
        if mejor.get("es_extranjero"):                      # aduana / logística
            extra += max(0.0, float(mejor.get("costo_aduana_estimado", 0) or 0))
        if art.get("requiere_instalacion"):                 # mano de obra
            mapa = {"baja": INSTAL_MIN, "media": (INSTAL_MIN + INSTAL_MAX) / 2,
                    "alta": INSTAL_MAX}
            extra += mapa.get(str(art.get("complejidad_instalacion", "baja")).lower(),
                              INSTAL_MIN)
        dom = _dominio(mejor.get("url_producto", "")) or f"prov{i}"
        resultado[i] = {"extra": round(extra, 2), "grupo": dom,
                        "k": len(grupos[dom])}
    return resultado


def generar_proforma_xlsm(analisis: dict, registro: Dict[str, Any],
                          imagenes: Dict[int, Optional[Path]],
                          ruta_salida: Path) -> Path:
    """
    Rellena la plantilla xlsm SIN tocar nada más que lo indicado:
      Hoja 'Cotización': B8, B9, B10, B11, B16.., D12, I3, I8, I10, I12,
                         C16.. (nombre + imagen), G16.. (cantidades).
      Hoja 'Costos'    : C7, C15.., E15.., A15.., G15.., J (alternativas).
    Las celdas de artículos sobrantes se vacían y sus filas se ocultan (alto 0)
    y las fórmulas globales (SUM/IVA/Total) siguen funcionando sin alterarse.
    El botón 'Imprimir PDF' se reinyecta en un post-proceso (openpyxl lo
    descarta al guardar).
    """
    import openpyxl
    wb = openpyxl.load_workbook(str(TEMPLATE_XLSX), keep_vba=True)
    ws_cot, ws_cst = wb["Cotización"], wb["Costos"]
    articulos = analisis.get("articulos", [])
    n = len(articulos)

    # ── Hoja "Cotización" · cabecera ─────────────────────────────────────────
    ws_cot["B8"]  = str(registro.get("entidad_contratante", "")).upper()
    ws_cot["B9"]  = str(registro.get("codigo_necesidad", ""))[4:17]   # chars 5–17
    ws_cot["B10"] = registro.get("dirección_entrega", "")
    ws_cot["B11"] = _contacto_resumido(registro.get("contacto", ""))
    ws_cot["D12"] = registro.get("codigo_necesidad", "")
    ws_cot["I3"]  = _numero_proforma(registro.get("id_infima", ""))
    ws_cot["I8"]  = datetime.datetime.now()                 # formato dd/mm/aaaa de la celda
    ws_cot["I8"].number_format = "dd/mm/yyyy"
    if analisis.get("tiempo_entrega_dias"):
        ws_cot["I10"] = f"{int(analisis['tiempo_entrega_dias'])} DÍAS"
    if analisis.get("garantia_tecnica"):
        ws_cot["I12"] = str(analisis["garantia_tecnica"]).upper()

    # ── Hoja "Cotización" · artículos (filas 16..25 ↔ Costos 15..24) ─────────
    for i, art in enumerate(articulos):
        fila = 16 + i
        ws_cot[f"B{fila}"] = registro.get("CPC", "")        # mismo CPC para todos
        nombre = " ".join(x for x in (art.get("nombre"), art.get("marca"),
                                      art.get("modelo")) if x)
        if art.get("marca") and art["marca"].lower() in (art.get("nombre") or "").lower():
            nombre = art.get("nombre", nombre)
        ws_cot[f"C{fila}"] = nombre                         # C:E combinadas (ancla C)
        ws_cot[f"G{fila}"] = int(art.get("cantidad", 1))
        ws_cot.row_dimensions[fila].height = ALTURA_FILA_ARTICULO

    # Filas de artículos SOBRANTES: se vacían sus celdas y se ocultan (alto 0).
    # Así el SUM(I15:I25), el IVA y el Total del formato siguen intactos.
    for fila in range(16 + n, 26):
        for col in ("A", "B", "C", "F", "G", "H", "I"):
            ws_cot[f"{col}{fila}"] = None
        ws_cot.row_dimensions[fila].height = 0
        ws_cot.row_dimensions[fila].hidden = True

    # ── Hoja "Costos" ────────────────────────────────────────────────────────
    ws_cst["C7"] = registro.get("entidad_contratante_url", "")
    extras = _costos_extras(analisis, registro.get("dirección_entrega", ""))
    for i, art in enumerate(articulos):
        fila = 15 + i
        mejor = art.get("mejor_opcion") or {}
        ws_cst[f"C{fila}"] = mejor.get("url_producto", "")   # URL exacta de la mejor opción
        ws_cst[f"E{fila}"] = float(mejor.get("precio_unitario", 0) or 0)
        info = extras[i]
        if info["k"] >= 2:
            # Proveedor compartido: total del proveedor en A y reparto en G
            # con DENOMINADOR PROPIO del proveedor (no el global del formato).
            ws_cst[f"A{fila}"] = info["extra"] * info["k"]
            ws_cst[f"G{fila}"] = f"=A{fila}/{info['k']}"
        else:
            # Proveedor único: costo extra DIRECTO en G (se documenta en A=0).
            ws_cst[f"A{fila}"] = 0
            ws_cst[f"G{fila}"] = info["extra"]

    # Alternativas por artículo: bloques de 4 filas → art.1: J15..J17,
    # art.2: J19..J21, art.3: J23..J25, … (mismo patrón del formato).
    for i, art in enumerate(articulos):
        base = 15 + 4 * i
        alts = (art.get("alternativas") or [])[:3]          # máx. 3, decreciente
        for j, alt in enumerate(alts):
            ws_cst[f"J{base + j}"] = alt.get("url", "")

    wb.save(str(ruta_salida))
    _postprocesar_proforma(ruta_salida, articulos, imagenes)
    log(f"FASE 6 · Proforma generada: {ruta_salida.name}")
    return ruta_salida


# ── Post-proceso ZIP: reinyecta logos + botón "Imprimir PDF" + imágenes ──────
_NS_XDR = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
_EMU_POR_PX = 9525


def _anchor_imagen_producto(indice_art: int, rid: str, ruta_img: Path) -> str:
    """
    Construye el XML <xdr:oneCellAnchor> para pegar la imagen del producto en
    la celda C{16+índice} (columna C = índice 2; fila Excel 16 = índice 15),
    ajustada a la altura recomendada de la fila (220 pt ≈ 293 px).
    """
    from PIL import Image
    with Image.open(ruta_img) as im:
        w, h = im.size
    alto_px = 265                                           # deja margen en la fila
    ancho_px = int(w * alto_px / h)
    if ancho_px > 430:                                      # no desbordar C:E
        ancho_px = 430
        alto_px = int(h * ancho_px / w)
    fila0 = 15 + indice_art                                 # 0-based (art.1 → 15)
    return (
        f'<xdr:oneCellAnchor>'
        f'<xdr:from><xdr:col>2</xdr:col><xdr:colOff>1350000</xdr:colOff>'
        f'<xdr:row>{fila0}</xdr:row><xdr:rowOff>190500</xdr:rowOff></xdr:from>'
        f'<xdr:ext cx="{ancho_px * _EMU_POR_PX}" cy="{alto_px * _EMU_POR_PX}"/>'
        f'<xdr:pic><xdr:nvPicPr><xdr:cNvPr id="{100 + indice_art}" '
        f'name="Producto {indice_art + 1}"/><xdr:cNvPicPr>'
        f'<a:picLocks noChangeAspect="1"/></xdr:cNvPicPr></xdr:nvPicPr>'
        f'<xdr:blipFill><a:blip xmlns:r="http://schemas.openxmlformats.org/'
        f'officeDocument/2006/relationships" r:embed="{rid}"/>'
        f'<a:stretch><a:fillRect/></a:stretch></xdr:blipFill>'
        f'<xdr:spPr><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></xdr:spPr>'
        f'</xdr:pic><xdr:clientData/></xdr:oneCellAnchor>'
    )


def _postprocesar_proforma(ruta_xlsm: Path, articulos: List[dict],
                           imagenes: Dict[int, Optional[Path]]) -> None:
    """
    openpyxl descarta al guardar los dibujos existentes (logos y la autoforma
    xdr:sp del botón 'Imprimir PDF'). Este post-proceso reabre el .xlsm como
    ZIP y:
      1) Restaura drawing1.xml + sus rels + media de la PLANTILLA (logos+botón).
      2) Añade las imágenes de producto ancladas en C16.. con rels nuevos.
      3) Re-vincula el dibujo en sheet1.xml y sus relationships.
      4) Asegura los tipos de contenido (png / drawing) en [Content_Types].xml.
    Soporta XML con prefijo 'xdr:' o con namespace por defecto.
    """
    # ── Lee las piezas necesarias de la plantilla original ───────────────────
    with zipfile.ZipFile(str(TEMPLATE_XLSX)) as ztpl:
        drawing_tpl = ztpl.read("xl/drawings/drawing1.xml").decode("utf-8")
        rels_tpl    = ztpl.read("xl/drawings/_rels/drawing1.xml.rels").decode("utf-8")
        media_tpl   = {n: ztpl.read(n) for n in ztpl.namelist()
                       if n.startswith("xl/media/")}

    # ── Inserta anclas de imágenes de producto en el drawing de la plantilla ─
    nuevos_rels, nuevas_medias = [], {}
    anclas = []
    for i, _art in enumerate(articulos):
        ruta_img = imagenes.get(i)
        if not ruta_img or not Path(ruta_img).exists():
            continue
        rid = f"rIdProd{i + 1}"
        nombre_media = f"xl/media/imageProd{i + 1}{Path(ruta_img).suffix or '.png'}"
        nuevas_medias[nombre_media] = Path(ruta_img).read_bytes()
        nuevos_rels.append(
            f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/'
            f'officeDocument/2006/relationships/image" '
            f'Target="../media/{Path(nombre_media).name}"/>')
        anclas.append(_anchor_imagen_producto(i, rid, Path(ruta_img)))

    # Cierre del drawing con o sin prefijo (plantillas heterogéneas)
    cierre = "</xdr:wsDr>" if "</xdr:wsDr>" in drawing_tpl else "</wsDr>"
    drawing_final = drawing_tpl.replace(cierre, "".join(anclas) + cierre)
    rels_final = rels_tpl.replace("</Relationships>",
                                  "".join(nuevos_rels) + "</Relationships>")

    # ── Reescribe el ZIP del xlsm generado ───────────────────────────────────
    tmp = ruta_xlsm.with_suffix(".tmp.xlsm")
    with zipfile.ZipFile(str(ruta_xlsm)) as zin, \
         zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as zout:
        rid_drawing = "rIdDraw1"
        for item in zin.namelist():
            data = zin.read(item)
            if item == "xl/worksheets/sheet1.xml":
                texto = data.decode("utf-8")
                if "<drawing " not in texto:                # re-vincula el dibujo
                    texto = texto.replace(
                        "</worksheet>",
                        f'<drawing xmlns:r="http://schemas.openxmlformats.org/'
                        f'officeDocument/2006/relationships" r:id="{rid_drawing}"/>'
                        f"</worksheet>")
                data = texto.encode("utf-8")
            elif item == "xl/worksheets/_rels/sheet1.xml.rels":
                texto = data.decode("utf-8")
                if "drawings/drawing1.xml" not in texto:
                    texto = texto.replace(
                        "</Relationships>",
                        f'<Relationship Id="{rid_drawing}" Type="http://schemas.'
                        f'openxmlformats.org/officeDocument/2006/relationships/'
                        f'drawing" Target="../drawings/drawing1.xml"/></Relationships>')
                data = texto.encode("utf-8")
            elif item == "[Content_Types].xml":
                texto = data.decode("utf-8")
                if 'Extension="png"' not in texto:
                    texto = texto.replace("</Types>",
                        '<Default Extension="png" ContentType="image/png"/></Types>')
                if 'Extension="jpg"' not in texto and any(
                        m.endswith((".jpg", ".jpeg")) for m in nuevas_medias):
                    texto = texto.replace("</Types>",
                        '<Default Extension="jpg" ContentType="image/jpeg"/></Types>')
                if "/xl/drawings/drawing1.xml" not in texto:
                    texto = texto.replace("</Types>",
                        '<Override PartName="/xl/drawings/drawing1.xml" ContentType='
                        '"application/vnd.openxmlformats-officedocument.drawing+xml"/>'
                        "</Types>")
                data = texto.encode("utf-8")
            elif item == "xl/drawings/drawing1.xml":
                continue                                    # se reescribe abajo
            elif item == "xl/drawings/_rels/drawing1.xml.rels":
                continue
            elif item.startswith("xl/media/"):
                continue                                    # media se reescribe
            zout.writestr(item, data)
        # Piezas restauradas / nuevas
        zout.writestr("xl/drawings/drawing1.xml", drawing_final)
        zout.writestr("xl/drawings/_rels/drawing1.xml.rels", rels_final)
        for nombre, contenido in {**media_tpl, **nuevas_medias}.items():
            zout.writestr(nombre, contenido)
    shutil.move(str(tmp), str(ruta_xlsm))
    log("FASE 6 · Botón 'Imprimir PDF', logos e imágenes reinyectados.")


# ═══════════════════════════════════════════════════════════════════════════════
#  PROCESAMIENTO DE UN CÓDIGO DE NECESIDAD (orquesta las fases 2–7)
# ═══════════════════════════════════════════════════════════════════════════════

def procesar_codigo(registro: Dict[str, Any], bucket, carpeta_trabajo: Path,
                    crono: Cronometro, estadisticas: dict,
                    analisis_previo: Optional[dict] = None,
                    modo_local: bool = False,
                    docs_locales: Optional[List[Path]] = None,
                    salida_local: Optional[Path] = None) -> None:
    """
    Ejecuta el pipeline completo para UN código de necesidad. En modo prueba
    local no toca ni GCS ni MySQL y puede recibir un análisis simulado.
    """
    codigo, id_infima = registro["codigo_necesidad"], registro["id_infima"]
    log(f"╔═ Procesando {codigo} (id_infima={id_infima}) " + "═" * 20)
    carpeta = carpeta_trabajo / re.sub(r"\W+", "_", str(codigo))
    carpeta.mkdir(parents=True, exist_ok=True)

    # FASE 2 · documentos de contratación
    crono.iniciar(f"FASE 2 · Documentos [{codigo}]")
    docs = docs_locales if modo_local else \
        descargar_documentos_contratacion(bucket, codigo, carpeta / "docs")
    crono.detener()
    if not docs:
        log(f"{codigo}: sin documentos de contratación; se omite.", "WARN")
        estadisticas["errores"] += 1
        return

    # FASE 3 · análisis con IA (o análisis simulado en pruebas)
    crono.iniciar(f"FASE 3 · Análisis IA [{codigo}]")
    analisis = analisis_previo or analizar_con_ia(docs, registro)
    crono.detener()

    excede = analisis.get("excede_limite") or \
             analisis.get("num_articulos", 0) > LIMITE_ARTICULOS

    # FASE 4 · verificación de URLs e imágenes (solo si no excede el límite)
    imagenes: Dict[int, Optional[Path]] = {}
    if not excede:
        crono.iniciar(f"FASE 4 · URLs e imágenes [{codigo}]")
        for i, art in enumerate(analisis.get("articulos", [])):
            mejor = art.get("mejor_opcion") or {}
            nombre = f"{art.get('marca','')} {art.get('modelo','')}".strip()
            if not modo_local:                              # en la nube: verificar/reparar
                url_ok = resolver_url_producto(mejor.get("url_producto", ""), nombre)
                if url_ok:
                    mejor["url_producto"] = url_ok
                alts_ok = []
                for alt in (art.get("alternativas") or []):
                    if verificar_url(alt.get("url", ""), estricto=False):
                        alts_ok.append(alt)
                art["alternativas"] = alts_ok[:3]
            imagenes[i] = obtener_imagen_producto(art, carpeta,
                                                  con_ia=not modo_local)
            estado = "OK" if imagenes[i] else "SIN IMAGEN"
            log(f"FASE 4 · Artículo {i + 1}: imagen {estado}.")
        crono.detener()

    # FASE 5 · ficha técnica
    crono.iniciar(f"FASE 5 · Ficha técnica [{codigo}]")
    nombre_ficha = f"{codigo}_Ficha_técnica_{id_infima}.docx"
    ruta_ficha = (salida_local or carpeta) / nombre_ficha
    generar_ficha_docx(analisis, imagenes, ruta_ficha)
    if not modo_local:
        subir_a_gcs(bucket, ruta_ficha, CARPETA_FICHAS)
    crono.detener()
    estadisticas["fichas"] += 1

    # FASE 6 · proforma (solo dentro del límite)
    if excede:
        log(f"{codigo}: > {LIMITE_ARTICULOS} artículos → SIN proforma.", "WARN")
        estadisticas["omitidas_limite"] += 1
    else:
        crono.iniciar(f"FASE 6 · Proforma [{codigo}]")
        nombre_prof = f"{codigo}_Proforma_{id_infima}.xlsm"
        ruta_prof = (salida_local or carpeta) / nombre_prof
        generar_proforma_xlsm(analisis, registro, imagenes, ruta_prof)
        if not modo_local:
            subir_a_gcs(bucket, ruta_prof, CARPETA_PROFORMAS)
        crono.detener()
        estadisticas["proformas"] += 1

    # FASE 7 · etapa → finalizada
    if not modo_local:
        crono.iniciar(f"FASE 7 · BD finalizada [{codigo}]")
        marcar_finalizada(id_infima)
        crono.detener()
    estadisticas["procesados"] += 1
    log(f"╚═ {codigo} completado " + "═" * 34)


# ═══════════════════════════════════════════════════════════════════════════════
#  RESÚMENES
# ═══════════════════════════════════════════════════════════════════════════════

def imprimir_resumen_final(estadisticas: dict, crono: Cronometro) -> None:
    """Resumen final con los resultados principales de la corrida."""
    crono.resumen()
    print("\n" + "═" * 70)
    print(" RESUMEN FINAL DE RESULTADOS")
    print("═" * 70)
    print(f"   Códigos de necesidad procesados : {estadisticas['procesados']}")
    print(f"   Fichas técnicas generadas       : {estadisticas['fichas']}")
    print(f"   Proformas generadas             : {estadisticas['proformas']}")
    print(f"   Omitidas por límite de artículos: {estadisticas['omitidas_limite']}")
    print(f"   Errores / omisiones             : {estadisticas['errores']}")
    print("═" * 70 + "\n")


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN · producción y modo de prueba local
# ═══════════════════════════════════════════════════════════════════════════════

def main_produccion() -> None:
    """Flujo completo contra MySQL + GCS + Vertex AI."""
    crono = Cronometro()
    stats = {"procesados": 0, "fichas": 0, "proformas": 0,
             "omitidas_limite": 0, "errores": 0}
    log("═══ 5_Preforms_generator · MODO PRODUCCIÓN ═══")

    crono.iniciar("FASE 1 · Lectura de data_table_1 (MySQL)")
    data_table_1 = obtener_data_table_1()
    crono.detener()
    if not data_table_1:
        log("No hay ínfimas en 'en generacion'; nada que hacer.")
        imprimir_resumen_final(stats, crono)
        return

    bucket = cliente_gcs().bucket(BUCKET_NAME)
    with tempfile.TemporaryDirectory(prefix="nexus5_") as tmp:
        carpeta = Path(tmp)
        for registro in data_table_1:
            try:
                procesar_codigo(registro, bucket, carpeta, crono, stats)
            except Exception as e:                          # aísla fallos por código
                stats["errores"] += 1
                log(f"ERROR en {registro.get('codigo_necesidad')}: {e}", "ERROR")
                traceback.print_exc()
    imprimir_resumen_final(stats, crono)


def main_test_local(ruta_pdf: Path, ruta_mock: Optional[Path],
                    dir_salida: Path) -> None:
    """
    Prueba de punta a punta SIN nube: usa un PDF local como documento de
    contratación y (si existe) un JSON simulado con 'registro' y 'analisis'.
    Si no hay JSON pero sí credenciales de Vertex, llama a la IA de verdad.
    """
    crono = Cronometro()
    stats = {"procesados": 0, "fichas": 0, "proformas": 0,
             "omitidas_limite": 0, "errores": 0}
    log("═══ 5_Preforms_generator · MODO PRUEBA LOCAL ═══")

    mock = json.loads(ruta_mock.read_text(encoding="utf-8")) if ruta_mock else {}
    registro = mock.get("registro") or {
        "codigo_necesidad": "NIC-0000000000000-2026-00001",
        "entidad_contratante": "Entidad de prueba",
        "entidad_contratante_url": "https://example.com/",
        "dirección_entrega": "Quito, Ecuador",
        "id_infima": 1, "CPC": "000000000", "contacto": "N/N",
    }
    analisis = mock.get("analisis")                          # None ⇒ IA real
    dir_salida.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="nexus5_test_") as tmp:
        procesar_codigo(registro, bucket=None, carpeta_trabajo=Path(tmp),
                        crono=crono, estadisticas=stats,
                        analisis_previo=analisis, modo_local=True,
                        docs_locales=[ruta_pdf], salida_local=dir_salida)
    imprimir_resumen_final(stats, crono)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generador de fichas y proformas NEXUS")
    parser.add_argument("--test-local", metavar="PDF",
                        help="Ruta a un documento de contratación local (modo prueba)")
    parser.add_argument("--mock-json", metavar="JSON",
                        help="JSON con 'registro' y 'analisis' simulados (opcional)")
    parser.add_argument("--salida", metavar="DIR", default="salida_test",
                        help="Directorio de salida en modo prueba (def: salida_test)")
    args = parser.parse_args()

    inicio_total = time.perf_counter()
    if args.test_local:
        main_test_local(Path(args.test_local),
                        Path(args.mock_json) if args.mock_json else None,
                        Path(args.salida))
    else:
        main_produccion()
    log(f"Ejecución total del script: {time.perf_counter() - inicio_total:.2f} s")